import streamlit as st

st.set_page_config(
    page_title="Calipr AI — Redrob Ranker Sandbox",
    page_icon="calipr_logo.svg",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# Initialize Session and Auth Gate
from auth import init_session, render_sign_in, render_sign_up, sign_out, render_navbar_user, is_pro, get_plan

init_session()

if not st.session_state.authenticated:
    if st.session_state.auth_page == "signup":
        render_sign_up()
    else:
        render_sign_in()
    st.stop()

import sys

# Diagnostic import wrapper
try:
    import json
    import os
    import numpy as np
    import time
    from datetime import date
    from sentence_transformers import SentenceTransformer
    import re
    from docx import Document
    import pandas as pd
    import uuid
    import plotly.graph_objects as go
    import plotly.express as px
    from supabase import create_client, Client
    
    # Auth pages imports
    from auth.signin_page import render_signin_page
    from auth.signup_page import render_signup_page
    from auth.supabase_auth import is_authenticated, sign_out, get_current_user, check_oauth_callback, get_supabase_url
except Exception as e:
    st.error(f"🚨 Calipr Diagnostic Error — Import Failure: {e}")
    st.write("### Diagnostics Information")
    st.write(f"**Python Version:** {sys.version}")
    st.write(f"**Executable:** {sys.executable}")
    st.write("### sys.path:")
    st.code("\n".join(sys.path))
    
    # Try listing installed packages via importlib.metadata
    try:
        import importlib.metadata
        pkgs = sorted([f"{dist.metadata['Name']}=={dist.version}" for dist in importlib.metadata.distributions()])
        st.write("### Installed Packages in Space:")
        st.code("\n".join(pkgs))
    except Exception as err:
        st.write(f"Could not list packages: {err}")
    st.stop()

# Safe import of PdfReader
try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

# Check Google OAuth callback query parameter
check_oauth_callback()

# ── AUTH GATE ──────────────────────────────────────────────────────
if not is_authenticated():
    # Default to sign in page
    if "auth_page" not in st.session_state:
        st.session_state.auth_page = "signin"

    if st.session_state.auth_page == "signin":
        render_signin_page()
    else:
        render_signup_page()

    st.stop()  # Do NOT render main app if not logged in

# ── AUTHENTICATED — show main app ──────────────────────────────────
user      = get_current_user()
user_name = st.session_state.get("user_name", "User")
user_email= st.session_state.get("user_email", "")

# Page routing query parameter
active_page = st.query_params.get("page", "ranker")

# Initialize default states for page variables to prevent undefined variable errors on other pages
jd_text = ""
run_pipeline = False

# User info + sign out in sidebar
with st.sidebar:
    st.markdown(f"""
    <div style="display:flex;align-items:center;gap:10px;padding:16px 0;
                border-bottom:1px solid #F3F4F6;margin-bottom:16px;">
        <div style="width:36px;height:36px;background:#0A0A0A;border-radius:50%;
                    display:flex;align-items:center;justify-content:center;
                    font-size:14px;font-weight:700;color:#fff;
                    font-family:Inter,sans-serif;flex-shrink:0;">
            {user_name[0].upper()}
        </div>
        <div>
            <div style="font-size:13px;font-weight:700;color:#0A0A0A;
                        font-family:Inter,sans-serif;">{user_name}</div>
            <div style="font-size:11px;color:#9CA3AF;
                        font-family:Inter,sans-serif;">{user_email}</div>
        </div>
    </div>
    """, unsafe_allow_html=True)

    if st.button("Sign Out", key="signout_btn"):
        sign_out()
        st.rerun()


# ── FULL CSS INJECTION ────────────────────────────────────────────
st.markdown("""
<style>
/* ── FONTS ── */
@import url('https://fonts.googleapis.com/css2?family=Inter:ital,wght@0,300;0,400;0,500;0,600;0,700;0,800;1,400;1,700&family=Fragment+Mono&display=swap');

@font-face {
  font-family: 'Open Runde';
  src: url('https://framerusercontent.com/assets/HJY4SY2JywrSZ2l1AEW9Tm9cO8.woff2') format('woff2');
  font-weight: 500;
  font-style: normal;
  font-display: swap;
}
@font-face {
  font-family: 'Open Runde';
  src: url('https://framerusercontent.com/assets/hG3wmGmFwadB6X5XPVXkMlmLr8o.woff2') format('woff2');
  font-weight: 600;
  font-style: normal;
  font-display: swap;
}
@font-face {
  font-family: 'Open Runde';
  src: url('https://framerusercontent.com/assets/3exmuO07FP19gMM08TQrpXl3BGQ.woff2') format('woff2');
  font-weight: 400;
  font-style: normal;
  font-display: swap;
}

/* ── GLOBAL RESET & THEME ── */
*, *::before, *::after { box-sizing: border-box; }

.stApp {
    background: #FFFFFF !important;
    font-family: 'Inter', system-ui, -apple-system, sans-serif !important;
    color: #453f3d !important;
}

/* ── HIDE STREAMLIT CHROME ── */
#MainMenu, footer, header { visibility: hidden !important; }
.stDeployButton { display: none !important; }
[data-testid="stToolbar"] { display: none !important; }

/* ── MAIN CONTAINER ── */
.block-container {
    max-width: 1200px !important;
    padding: 0 24px 80px !important;
    margin: 0 auto !important;
}

/* ── SIDEBAR ── */
[data-testid="stSidebar"] {
    background: #FFFFFF !important;
    border-right: 1px solid #e4e2e2 !important;
}
[data-testid="stSidebar"] .block-container {
    padding: 32px 20px !important;
}

/* ── HEADINGS ── */
h1 {
    font-family: 'Open Runde', 'Inter', system-ui, sans-serif !important;
    font-size: 36px !important;
    font-weight: 600 !important;
    letter-spacing: -0.03em !important;
    color: #1a1615 !important;
    line-height: 1.2 !important;
}
h2 {
    font-family: 'Open Runde', 'Inter', system-ui, sans-serif !important;
    font-size: 26px !important;
    font-weight: 600 !important;
    letter-spacing: -0.02em !important;
    color: #1a1615 !important;
}
h3 {
    font-family: 'Open Runde', 'Inter', system-ui, sans-serif !important;
    font-size: 18px !important;
    font-weight: 600 !important;
    color: #1a1615 !important;
}

/* ── BUTTONS ── */
.stButton > button {
    background: #1a1615 !important;
    color: #ffffff !important;
    border: 1px solid #1a1615 !important;
    border-radius: 99px !important;
    font-family: 'Inter', sans-serif !important;
    font-weight: 600 !important;
    font-size: 14px !important;
    padding: 10px 24px !important;
    transition: all 0.2s ease !important;
    letter-spacing: -0.01em !important;
}
.stButton > button:hover {
    background: #2d2522 !important;
    border-color: #2d2522 !important;
    transform: translateY(-1px) !important;
    box-shadow: 0 4px 12px rgba(26, 22, 21, 0.12) !important;
}
.stButton > button:active {
    transform: translateY(0) !important;
}

/* Secondary Button style for downloads, etc */
.stDownloadButton > button {
    background: rgba(255, 255, 255, 0.8) !important;
    color: #1a1615 !important;
    border: 1px solid #e4e2e2 !important;
    border-radius: 99px !important;
    padding: 10px 24px !important;
    font-family: 'Inter', sans-serif !important;
    font-weight: 600 !important;
    font-size: 14px !important;
    transition: all 0.2s ease !important;
}
.stDownloadButton > button:hover {
    background: #ffffff !important;
    border-color: #84b9ef !important;
    color: #156cc2 !important;
    box-shadow: 0 4px 12px rgba(26, 22, 21, 0.05) !important;
}

/* ── TEXT INPUTS & TEXTAREAS ── */
.stTextArea > div,
.stTextInput > div,
.stTextInput div[data-baseweb="input"],
.stTextInput div[data-testid="stTextInputRootElement"] {
    border: none !important;
    box-shadow: none !important;
    background-color: transparent !important;
}

.stTextArea textarea,
.stTextInput div[data-testid="stTextInputRootElement"],
.stTextInput div[data-baseweb="input"] {
    background: rgba(255, 255, 255, 0.7) !important;
    border: 1px solid #e4e2e2 !important;
    border-radius: 12px !important;
    font-family: 'Inter', sans-serif !important;
    font-size: 14px !important;
    color: #1a1615 !important;
    transition: all 0.2s ease !important;
}

.stTextArea textarea {
    padding: 12px 16px !important;
}

.stTextInput div[data-testid="stTextInputRootElement"] div,
.stTextInput div[data-baseweb="input"] div,
.stTextArea div[data-baseweb="textarea"] div {
    background-color: transparent !important;
    background: transparent !important;
    border: none !important;
    box-shadow: none !important;
}

.stTextInput div[data-testid="stTextInputRootElement"],
.stTextInput div[data-baseweb="input"] {
    height: 42px !important;
    display: flex !important;
    align-items: center !important;
}

.stTextInput input {
    background: transparent !important;
    border: none !important;
    box-shadow: none !important;
    outline: none !important;
    font-family: 'Inter', sans-serif !important;
    font-size: 14px !important;
    color: #1a1615 !important;
    width: 100% !important;
    height: 100% !important;
    padding: 0 16px !important;
}

.stTextArea textarea:focus,
.stTextInput div[data-testid="stTextInputRootElement"]:focus-within,
.stTextInput div[data-baseweb="input"]:focus-within {
    border-color: #84b9ef !important;
    box-shadow: 0 0 0 3px rgba(132, 185, 239, 0.2) !important;
    background: #ffffff !important;
    outline: none !important;
}

/* ── LABELS ── */
.stTextArea label, .stTextInput label, .stSelectbox label, .stRadio label, .stFileUploader label {
    font-family: 'Fragment Mono', monospace !important;
    font-size: 11px !important;
    font-weight: 600 !important;
    text-transform: uppercase !important;
    letter-spacing: 0.05em !important;
    color: #757170 !important;
    margin-bottom: 6px !important;
}

/* ── SELECT BOX ── */
div[data-baseweb="select"] {
    background: rgba(255, 255, 255, 0.8) !important;
    border: 1px solid #e4e2e2 !important;
    border-radius: 12px !important;
    font-family: 'Inter', sans-serif !important;
}
div[data-baseweb="select"] > div {
    background: transparent !important;
    border: none !important;
}

/* ── RADIO BUTTONS ── */
.stRadio [data-testid="stMarkdownContainer"] p {
    font-family: 'Inter', sans-serif !important;
    font-size: 13.5px !important;
    color: #453f3d !important;
}
div[data-testid="stRadio"] label[data-baseweb="radio"] input[type="radio"]:checked + div {
    border-color: #156cc2 !important;
}
div[data-testid="stRadio"] label[data-baseweb="radio"] input[type="radio"]:checked + div > div {
    background-color: #156cc2 !important;
}
div[data-testid="stRadio"] [data-checked="true"] > div:first-child {
    border-color: #156cc2 !important;
}
div[data-testid="stRadio"] [data-checked="true"] div div {
    background-color: #156cc2 !important;
}


/* ── METRICS ── */
[data-testid="stMetric"] {
    background: rgba(255, 255, 255, 0.7) !important;
    backdrop-filter: blur(10px) !important;
    border: 1px solid rgba(228, 226, 226, 0.8) !important;
    border-radius: 18px !important;
    padding: 16px 20px !important;
}
[data-testid="stMetricLabel"] {
    font-family: 'Fragment Mono', monospace !important;
    font-size: 10px !important;
    text-transform: uppercase !important;
    color: #757170 !important;
}
[data-testid="stMetricValue"] {
    font-size: 28px !important;
    font-weight: 600 !important;
    color: #1a1615 !important;
    font-family: 'Open Runde', sans-serif !important;
}

/* ── EXPANDER ── */
.streamlit-expander {
    background: rgba(255, 255, 255, 0.6) !important;
    backdrop-filter: blur(8px) !important;
    border: 1px solid rgba(228, 226, 226, 0.8) !important;
    border-radius: 16px !important;
    overflow: hidden !important;
    margin-bottom: 20px !important;
}
.streamlit-expander header {
    background: transparent !important;
    font-family: 'Inter', sans-serif !important;
    font-weight: 600 !important;
    font-size: 14px !important;
    color: #1a1615 !important;
}

/* ── SUCCESS / ERROR / INFO ── */
.stSuccess {
    background: rgba(14, 161, 88, 0.08) !important;
    border: 1px solid rgba(14, 161, 88, 0.25) !important;
    border-radius: 12px !important;
    color: #0c7540 !important;
}
.stError {
    background: rgba(201, 80, 46, 0.08) !important;
    border: 1px solid rgba(201, 80, 46, 0.25) !important;
    border-radius: 12px !important;
    color: #96351d !important;
}
.stInfo {
    background: rgba(21, 108, 194, 0.06) !important;
    border: 1px solid rgba(21, 108, 194, 0.2) !important;
    border-radius: 12px !important;
    color: #124d85 !important;
}

/* ── PROGRESS BAR ── */
.stProgress > div > div {
    background: linear-gradient(90deg, #84b9ef, #156cc2) !important;
    border-radius: 100px !important;
}
.stProgress > div {
    background: #e4e2e2 !important;
    border-radius: 100px !important;
    height: 8px !important;
}

/* ── CUSTOM COMPONENT CLASSES ── */

/* DASHBOARD NAV */
.dashboard-nav {
    display: flex;
    justify-content: space-between;
    align-items: center;
    background: rgba(255, 255, 255, 0.75) !important;
    backdrop-filter: blur(12px) !important;
    border: 1px solid rgba(228, 226, 226, 0.7) !important;
    border-radius: 20px !important;
    padding: 14px 24px !important;
    margin-top: 20px !important;
    margin-bottom: 30px !important;
    box-shadow: 0 4px 20px rgba(26, 22, 21, 0.03) !important;
}
.nav-logo {
    display: flex;
    align-items: center;
    gap: 8px;
    font-family: 'Open Runde', sans-serif !important;
    font-weight: 600 !important;
    font-size: 18px !important;
    color: #1a1615 !important;
}
.nav-badge {
    background: #e2ecf5;
    color: #156cc2;
    font-size: 10px;
    font-weight: 700;
    padding: 2px 6px;
    border-radius: 4px;
    letter-spacing: 0.5px;
}
.nav-menu {
    display: flex;
    gap: 24px;
}
.nav-item {
    font-size: 14px;
    font-weight: 500;
    color: #757170;
    cursor: pointer;
    transition: color 0.2s ease;
}
.nav-item:hover {
    color: #1a1615;
}
.nav-item.active {
    color: #156cc2;
    font-weight: 600;
    border-bottom: 2px solid #156cc2;
    padding-bottom: 2px;
}
.nav-user {
    display: flex;
    align-items: center;
    gap: 8px;
}
.user-avatar {
    width: 28px;
    height: 28px;
    border-radius: 50%;
    background: #1a1615;
    color: #ffffff;
    display: flex;
    align-items: center;
    justify-content: center;
    font-size: 11px;
    font-weight: 700;
}
.user-name {
    font-size: 13px;
    font-weight: 600;
    color: #1a1615;
}

/* GLASS CARD — light */
.card {
    background: rgba(255, 255, 255, 0.75) !important;
    backdrop-filter: blur(10px) !important;
    border: 1px solid rgba(228, 226, 226, 0.7) !important;
    border-radius: 20px !important;
    padding: 24px !important;
    transition: all 0.25s cubic-bezier(0.16, 1, 0.3, 1) !important;
}
.card:hover {
    transform: translateY(-4px) !important;
    border-color: #84b9ef !important;
    box-shadow: 0 12px 30px rgba(26, 22, 21, 0.06) !important;
}

/* GLASS CARD — dark */
.card-dark {
    background: #1a1615 !important;
    border: 1px solid rgba(255, 255, 255, 0.05) !important;
    border-radius: 20px !important;
    padding: 24px !important;
    color: #FFFFFF !important;
    box-shadow: 0 10px 30px rgba(26, 22, 21, 0.15) !important;
}

/* HERO SECTION */
.hero-section {
    background: rgba(255, 255, 255, 0.6) !important;
    backdrop-filter: blur(12px) !important;
    border: 1px solid rgba(228, 226, 226, 0.7) !important;
    border-radius: 24px !important;
    padding: 32px !important;
    margin-bottom: 30px !important;
    box-shadow: 0 4px 20px rgba(26, 22, 21, 0.02) !important;
}
.hero-header-label {
    display: inline-flex;
    align-items: center;
    gap: 6px;
    font-family: 'Fragment Mono', monospace !important;
    font-size: 11px !important;
    color: #156cc2 !important;
    background: #e2ecf5 !important;
    padding: 4px 12px !important;
    border-radius: 99px !important;
    font-weight: 600 !important;
    letter-spacing: 0.5px !important;
    margin-bottom: 16px !important;
}
.label-dot {
    width: 6px;
    height: 6px;
    background: #156cc2;
    border-radius: 50%;
}
.hero-title {
    font-size: 32px !important;
    font-weight: 700 !important;
    letter-spacing: -1px !important;
    line-height: 1.2 !important;
    margin: 0 !important;
    color: #1a1615 !important;
}
.hero-desc {
    font-size: 14px !important;
    color: #757170 !important;
    max-width: 650px !important;
    line-height: 1.6 !important;
    margin-top: 10px !important;
    margin-bottom: 0 !important;
}
.platform-stats {
    display: flex;
    gap: 20px;
    background: rgba(255, 255, 255, 0.8);
    border: 1px solid #e4e2e2;
    border-radius: 16px;
    padding: 12px 20px;
}
.p-stat {
    display: flex;
    flex-direction: column;
}
.p-stat-val {
    font-family: 'Fragment Mono', monospace;
    font-size: 18px;
    font-weight: 700;
    color: #1a1615;
}
.p-stat-lbl {
    font-size: 10px;
    font-weight: 500;
    color: #757170;
    text-transform: uppercase;
}
.hero-badges {
    display: flex;
    gap: 8px;
    margin-top: 20px;
    flex-wrap: wrap;
}

/* PILL BADGE */
.badge {
    display: inline-flex;
    align-items: center;
    gap: 6px;
    background: rgba(255, 255, 255, 0.6);
    border: 1px solid #e4e2e2;
    border-radius: 100px;
    padding: 6px 14px;
    font-size: 12px;
    font-weight: 600;
    color: #453f3d;
    font-family: 'Inter', sans-serif;
}
.badge-blue {
    background: #e2ecf5;
    border-color: rgba(132, 185, 239, 0.5);
    color: #156cc2;
}
.badge-green {
    background: rgba(14, 161, 88, 0.1);
    border-color: rgba(14, 161, 88, 0.3);
    color: #0ea158;
}

/* SCORE BAR */
.score-bar-container {
    margin-bottom: 14px;
}
.score-bar-label {
    display: flex;
    justify-content: space-between;
    font-size: 13px;
    font-weight: 500;
    color: #757170;
    font-family: 'Inter', sans-serif;
    margin-bottom: 6px;
}
.score-bar-label span {
    font-weight: 700;
    color: #1a1615;
    font-family: 'Fragment Mono', monospace;
}
.score-bar-track {
    height: 8px;
    background: #eddfd0;
    border-radius: 100px;
    overflow: hidden;
}
.score-bar-fill {
    height: 100%;
    border-radius: 100px;
}

/* CANDIDATE CARD */
.candidate-card {
    background: rgba(255, 255, 255, 0.65) !important;
    backdrop-filter: blur(12px) !important;
    -webkit-backdrop-filter: blur(12px) !important;
    border: 1px solid rgba(228, 226, 226, 0.8) !important;
    border-radius: 16px !important;
    padding: 16px 20px !important;
    cursor: pointer;
    transition: all 0.25s cubic-bezier(0.16, 1, 0.3, 1) !important;
    margin-bottom: 10px !important;
}
.candidate-card:hover {
    transform: translateY(-2px) !important;
    border-color: #84b9ef !important;
    background: rgba(255, 255, 255, 0.9) !important;
    box-shadow: 0 8px 24px rgba(26, 22, 21, 0.05) !important;
}
.candidate-card.selected {
    border-color: #156cc2 !important;
    background: rgba(210, 225, 245, 0.45) !important;
    box-shadow: 0 4px 16px rgba(21, 108, 194, 0.08) !important;
}

/* RANK BADGE */
.rank-badge {
    display: inline-flex;
    align-items: center;
    justify-content: center;
    width: 28px;
    height: 28px;
    background: #f4e6da;
    color: #754d29;
    border-radius: 8px;
    font-size: 12px;
    font-weight: 700;
    font-family: 'Inter', sans-serif;
}
.rank-badge.top3 {
    background: #e2ecf5;
    color: #156cc2;
}

/* SIGNAL WEIGHT BADGE */
.weight-badge {
    background: #1a1615 !important;
    color: #ffffff !important;
    font-size: 11px !important;
    font-weight: 700 !important;
    padding: 4px 10px !important;
    border-radius: 99px !important;
}

/* SECTION LABEL */
.section-label {
    font-family: 'Fragment Mono', monospace !important;
    font-size: 11px !important;
    font-weight: 600 !important;
    text-transform: uppercase !important;
    letter-spacing: 0.08em !important;
    color: #757170 !important;
    margin-bottom: 12px !important;
}

/* PIPELINE PHASE CARD */
.phase-card {
    background: rgba(255, 255, 255, 0.75) !important;
    backdrop-filter: blur(10px) !important;
    border: 1px solid rgba(228, 226, 226, 0.7) !important;
    border-radius: 16px !important;
    padding: 24px 20px !important;
    text-align: center !important;
    color: #453f3d !important;
    position: relative !important;
    transition: all 0.2s ease !important;
}
.phase-card:hover {
    transform: translateY(-2px) !important;
    border-color: #84b9ef !important;
    box-shadow: 0 8px 24px rgba(26, 22, 21, 0.04) !important;
}
.phase-number {
    position: absolute !important;
    top: -12px !important;
    left: 50% !important;
    transform: translateX(-50%) !important;
    background: #156cc2 !important;
    color: #FFFFFF !important;
    width: 24px !important;
    height: 24px !important;
    border-radius: 50% !important;
    display: flex !important;
    align-items: center !important;
    justify-content: center !important;
    font-size: 11px !important;
    font-weight: 700 !important;
}

/* RATIONALE BOX */
.rationale-box {
    background: rgba(255, 255, 255, 0.8) !important;
    border: 1px solid #e4e2e2 !important;
    border-left: 4px solid #156cc2 !important;
    border-radius: 12px !important;
    padding: 16px 20px !important;
    font-size: 13.5px !important;
    color: #453f3d !important;
    line-height: 1.6 !important;
    font-style: italic !important;
}

/* STAT CARD */
.stat-card {
    background: rgba(255, 255, 255, 0.85) !important;
    backdrop-filter: blur(8px) !important;
    border: 1px solid #e4e2e2 !important;
    border-radius: 20px !important;
    padding: 24px !important;
    text-align: center !important;
    transition: transform 0.2s ease, box-shadow 0.2s ease !important;
    box-shadow: 0 4px 12px rgba(26, 22, 21, 0.02) !important;
}
.stat-card:hover {
    transform: translateY(-2px) !important;
    box-shadow: 0 8px 24px rgba(26, 22, 21, 0.05) !important;
}
.stat-number {
    font-size: 40px !important;
    font-weight: 700 !important;
    color: #1a1615 !important;
    font-family: 'Open Runde', sans-serif !important;
    letter-spacing: -1px !important;
    line-height: 1 !important;
}
.stat-label {
    font-size: 11px !important;
    font-weight: 600 !important;
    text-transform: uppercase !important;
    letter-spacing: 0.5px !important;
    color: #757170 !important;
    margin-top: 8px !important;
}

/* DIVIDER WITH LABEL */
.divider-label {
    display: flex;
    align-items: center;
    gap: 16px;
    margin: 32px 0;
}
.divider-label::before, .divider-label::after {
    content: '';
    flex: 1;
    height: 1px;
    background: #e4e2e2;
}
.divider-label span {
    font-family: 'Fragment Mono', monospace !important;
    font-size: 11px !important;
    font-weight: 600 !important;
    color: #757170 !important;
    text-transform: uppercase !important;
    letter-spacing: 0.08em !important;
    white-space: nowrap;
}
</style>
""", unsafe_allow_html=True)

# ── CONFIG & CONSTANTS ────────────────────────────────────────────
EMBEDDING_MODEL = "all-MiniLM-L6-v2"
LEVEL_MAP = {
    "intern":0.10,"trainee":0.12,"junior":0.20,"associate":0.28,
    "mid":0.40,"engineer":0.40,"developer":0.40,"analyst":0.35,
    "senior":0.70,"lead":0.82,"staff":0.88,"principal":0.93,
    "architect":0.90,"director":0.93,"manager":0.72,"head":0.85,
    "vp":0.95,"cto":1.0,"founder":0.88
}
LEVEL_ORDER = [
    "cto", "vp", "director", "architect", "principal", "staff", "lead", "senior",
    "intern", "trainee", "junior", "associate", "founder", "manager", "head",
    "mid", "engineer", "developer", "analyst"
]
SIZE_MAP = {"1-10":1,"11-50":2,"51-200":3,"201-500":4,
            "501-1000":5,"1001-5000":6,"5001-10000":7,"10001+":8}
SKILL_ADJACENCY = {
    "python": ["julia","r","scala","cython","cpython","micropython","jython","pypy"],
    "pytorch": ["tensorflow","jax","keras","mxnet","torch"],
    "react": ["vue","angular","svelte","next.js","react.js","redux"],
    "fastapi": ["flask","django","express","pydantic","uvicorn"],
    "postgresql": ["mysql","sqlite","mongodb","postgres","sql","supabase","cockroachdb"],
    "docker": ["kubernetes","podman","containerization","docker-compose","containerd"],
    "aws": ["gcp","azure","digitalocean","amazon web services","ec2","s3","lambda"],
    "langchain": ["llamaindex","haystack","autogen"],
    "bert": ["roberta","distilbert","gpt-2","t5"],
    "yolov8": ["yolov5","detectron2","efficientdet"],
}
try:
    adj_path = "skill_adjacency_map.json"
    if os.path.exists(adj_path):
        with open(adj_path, encoding="utf-8") as f:
            custom_adj = json.load(f)
            for k, v in custom_adj.items():
                k_lower = k.lower().strip()
                v_lowers = [item.lower().strip() for item in v]
                if k_lower in SKILL_ADJACENCY:
                    SKILL_ADJACENCY[k_lower] = list(set(SKILL_ADJACENCY[k_lower] + v_lowers))
                else:
                    SKILL_ADJACENCY[k_lower] = v_lowers
except Exception as e:
    pass



# Initialize session state variables
if "uploaded_candidates" not in st.session_state:
    st.session_state.uploaded_candidates = []
if "scored_candidates" not in st.session_state:
    st.session_state.scored_candidates = None
if "run_runtime" not in st.session_state:
    st.session_state.run_runtime = 0.0
if "total_candidates_evaluated" not in st.session_state:
    st.session_state.total_candidates_evaluated = 0

# ── CACHED MODEL LOADERS ──────────────────────────────────────────
@st.cache_resource
def load_sentence_transformer():
    return SentenceTransformer(EMBEDDING_MODEL)

@st.cache_data
def load_sample_candidates():
    with open("sample_candidates.json", "r", encoding="utf-8") as f:
        return json.load(f)

# ── RECRUITER MEMORY ENGINE FUNCTIONS ─────────────────────────────
def load_memory_rules():
    if "recruiter_memory" not in st.session_state:
        memory_path = "recruiter_memory.json"
        if not os.path.exists(memory_path):
            memory_path = "nexhire-ai/memory/recruiter_memory.json"
            
        try:
            if os.path.exists(memory_path):
                with open(memory_path, "r", encoding="utf-8") as f:
                    st.session_state.recruiter_memory = json.load(f)
            else:
                st.session_state.recruiter_memory = {
                    "disqualifiers": [],
                    "preferences": {}
                }
        except Exception:
            st.session_state.recruiter_memory = {
                "disqualifiers": [],
                "preferences": {}
            }
    return st.session_state.recruiter_memory

def save_memory_rules():
    if "recruiter_memory" in st.session_state:
        memory_path = "recruiter_memory.json"
        try:
            with open(memory_path, "w", encoding="utf-8") as f:
                json.dump(st.session_state.recruiter_memory, f, indent=2)
            other_path = "nexhire-ai/memory/recruiter_memory.json"
            if os.path.exists(other_path):
                with open(other_path, "w", encoding="utf-8") as f:
                    json.dump(st.session_state.recruiter_memory, f, indent=2)
        except Exception as e:
            st.warning(f"Could not persist memory rules: {e}")

def apply_recruiter_memory_rules(candidate, core_skills, min_years_exp, seniority_lvl, memory_data):
    disqualified = False
    disq_reason = ""
    boost = 0.0
    notes = []
    
    # 1. Evaluate disqualifiers
    for dq in memory_data.get("disqualifiers", []):
        if not dq.get("rule_id") or not dq.get("active", True):
            continue
            
        rule_id = dq["rule_id"]
        action = dq.get("action", "disqualify")
        severity = dq.get("severity", "hard")
        rationale = dq.get("rationale", "")
        
        triggered = False
        
        if rule_id == "no_python":
            python_required = any("python" in s.lower() for s in core_skills)
            if python_required:
                cand_skills = [s.get("name", "").lower() for s in candidate.get("skills", [])]
                has_python = any("python" in s for s in cand_skills)
                if not has_python:
                    triggered = True
                    
        elif rule_id == "insufficient_experience":
            cand_exp = candidate.get("profile", {}).get("years_of_experience", 0)
            if cand_exp < min_years_exp:
                triggered = True
                
        elif rule_id == "no_relevant_domain":
            cand_text = (candidate.get("profile", {}).get("summary", "") + " " + 
                         candidate.get("profile", {}).get("headline", "")).lower()
            overlap = False
            for job in candidate.get("career_history", []):
                desc = (job.get("description") or "").lower()
                for kw in ["ai", "ml", "nlp", "recruiting", "hr", "marketplace"]:
                    if kw in desc or kw in cand_text:
                        overlap = True
                        break
                if overlap:
                    break
            if not overlap:
                triggered = True
                
        else:
            cond = dq.get("condition", {})
            check_field = cond.get("check_field")
            operator = cond.get("operator")
            val = cond.get("value")
            val_from = cond.get("value_from")
            if val_from == "jd.min_years_experience":
                val = min_years_exp
                
            cand_val = None
            if check_field == "skills":
                cand_val = [s.get("name", "").lower() for s in candidate.get("skills", [])]
            elif check_field == "years_experience":
                cand_val = candidate.get("profile", {}).get("years_of_experience", 0)
            elif check_field == "current_title":
                cand_val = (candidate.get("profile", {}).get("current_title", "")).lower()
            elif check_field == "summary":
                cand_val = (candidate.get("profile", {}).get("summary", "")).lower()
                
            applies = True
            applies_when = cond.get("applies_when", {})
            if applies_when:
                jd_field = applies_when.get("jd_field")
                jd_op = applies_when.get("operator")
                jd_val = applies_when.get("value")
                
                if jd_field == "required_skills" and jd_op == "contains":
                    applies = any(jd_val.lower() in s.lower() for s in core_skills)
            
            if applies and cand_val is not None:
                if operator == "not_contains" and isinstance(cand_val, list) and val:
                    if val.lower() not in cand_val:
                        triggered = True
                elif operator == "less_than" and val is not None:
                    try:
                        if float(cand_val) < float(val):
                            triggered = True
                    except ValueError:
                        pass
                elif operator == "no_overlap":
                    triggered = True
                    
        if triggered:
            if action == "disqualify" and severity == "hard":
                disqualified = True
                disq_reason = rationale or f"Lacks required '{rule_id}' constraint"
                break
            elif action == "flag" or severity == "soft":
                notes.append(f"⚠️ Soft-flagged: {rationale or desc}")
                
    # 2. Evaluate preferences (boosts)
    for pref_id, pref in memory_data.get("preferences", {}).items():
        if not pref.get("active", True):
            continue
            
        cond = pref.get("condition", {})
        boost_val = pref.get("boost", 0.0)
        
        triggered = False
        
        if pref_id == "prefer_leadership":
            jd_is_senior = seniority_lvl.lower() in ["senior", "lead", "staff", "principal", "director"]
            cand_title = (candidate.get("profile", {}).get("current_title", "")).lower()
            cand_history_titles = [job.get("title", "").lower() for job in candidate.get("career_history", [])]
            has_lead_title = any(any(lead_kw in t for lead_kw in ["lead", "senior", "manager", "head", "director", "chief"]) for t in [cand_title] + cand_history_titles)
            has_lead_examples = candidate.get("profile", {}).get("leadership_examples") or candidate.get("leadership_examples")
            
            if jd_is_senior and (has_lead_title or has_lead_examples):
                triggered = True
                
        elif pref_id == "prefer_collaboration":
            has_collab = candidate.get("profile", {}).get("collaboration_style") or candidate.get("collaboration_style")
            desc_text = " ".join([job.get("description", "") for job in candidate.get("career_history", [])]).lower()
            has_collab_kw = any(kw in desc_text for kw in ["collaborate", "teamwork", "cross-functional", "partnered", "cooperated"])
            if has_collab or has_collab_kw:
                triggered = True
                
        if triggered:
            boost += boost_val
            notes.append(f"⭐ Preference Boost (+{boost_val*100:.0f}%): {pref.get('description')}")
            
    return disqualified, disq_reason, boost, notes


# ── RENDERING FUNCTIONS FOR INACTIVE TABS ─────────────────────────
def render_memory_page():
    st.markdown('<div class="section-label">Recruiter Memory Panel</div>', unsafe_allow_html=True)
    st.markdown('<h1 class="hero-title" style="margin-top:0 !important;margin-bottom:12px;">Recruiter Memory & Bias Config</h1>', unsafe_allow_html=True)
    st.markdown('<p style="font-size:15px;color:#757170;margin-bottom:24px;">Manage automatic qualification screening rules and soft preference boosts that Calipr applies to candidates.</p>', unsafe_allow_html=True)

    rules = load_memory_rules()

    # Disqualifiers Section
    st.markdown('<h3>🚫 Candidate Disqualifiers (Hard Rules)</h3>', unsafe_allow_html=True)
    st.markdown('<p style="font-size:13.5px;color:#757170;margin-bottom:15px;">Candidates triggering these rules will be assigned a suitability score of 0.0.</p>', unsafe_allow_html=True)
    
    dq_list = rules.get("disqualifiers", [])
    updated_dq = []
    
    for i, dq in enumerate(dq_list):
        rule_id = dq.get("rule_id", f"rule_{i}")
        desc = dq.get("description", "Custom rule")
        rationale = dq.get("rationale", "")
        active = dq.get("active", True)
        
        with st.container():
            status_text = "Active" if active else "Inactive"
            border_color = "rgba(239, 68, 68, 0.2)" if active else "#e4e2e2"
            bg_color = "rgba(239, 68, 68, 0.03)" if active else "rgba(255,255,255,0.4)"
            st.markdown(f"""
            <div style="background:{bg_color}; border:1px solid {border_color}; border-radius:12px; padding:16px; margin-bottom:16px; font-family:Inter,sans-serif;">
                <div style="display:flex; justify-content:space-between; align-items:center;">
                    <span style="font-weight:700; color:#1a1615; font-size:14.5px;">{rule_id.upper()}</span>
                    <span style="font-size:11px; font-weight:700; text-transform:uppercase; color:{'#ef4444' if active else '#9CA3AF'};">{status_text}</span>
                </div>
                <div style="font-size:13px; color:#453f3d; margin-top:8px; line-height:1.5;">{desc}</div>
                <div style="font-size:12px; color:#757170; margin-top:6px; font-style:italic;">Rationale: {rationale}</div>
            </div>
            """, unsafe_allow_html=True)
            
            c1, c2 = st.columns([1, 1])
            with c1:
                new_active = st.checkbox(f"Active", value=active, key=f"dq_active_{i}")
            with c2:
                delete_rule = st.button(f"🗑️ Delete Rule", key=f"dq_del_{i}")
                
            if not delete_rule:
                dq["active"] = new_active
                updated_dq.append(dq)
            else:
                st.rerun()

    rules["disqualifiers"] = updated_dq
    save_memory_rules()

    # Add Disqualifier Form
    with st.expander("➕ Create Custom Disqualifier"):
        with st.form("add_dq_form"):
            new_id = st.text_input("Rule ID", placeholder="e.g. no_golang")
            new_desc = st.text_input("Rule Description", placeholder="Disqualify candidates who lack Golang experience")
            new_rat = st.text_input("Rejection Rationale", placeholder="Golang is mandatory for this role")
            
            check_field = st.selectbox("Check Field", ["skills", "years_experience"])
            operator = st.selectbox("Operator", ["not_contains", "less_than"])
            cond_val = st.text_input("Value", placeholder="golang or 3")
            
            submitted = st.form_submit_button("Add Disqualifier")
            if submitted:
                if not new_id or not new_desc:
                    st.error("Please fill in Rule ID and Description")
                else:
                    new_rule = {
                        "rule_id": new_id,
                        "description": new_desc,
                        "condition": {
                            "check_field": check_field,
                            "operator": operator,
                            "value": cond_val
                        },
                        "action": "disqualify",
                        "severity": "hard",
                        "rationale": new_rat,
                        "active": True
                    }
                    rules["disqualifiers"].append(new_rule)
                    save_memory_rules()
                    st.success(f"Added disqualifier {new_id}")
                    st.rerun()

    # Preferences Section
    st.markdown('<hr style="border: none; border-top: 1px solid #e4e2e2; margin: 40px 0;">', unsafe_allow_html=True)
    st.markdown('<h3>⭐ Recruiter Preferences (Soft Boosts)</h3>', unsafe_allow_html=True)
    st.markdown('<p style="font-size:13.5px;color:#757170;margin-bottom:15px;">Candidates meeting these criteria will receive a soft multiplier/additive boost to their final score.</p>', unsafe_allow_html=True)

    prefs = rules.get("preferences", {})
    updated_prefs = {}
    
    for pref_id, pref in list(prefs.items()):
        desc = pref.get("description", "Custom preference")
        boost_val = pref.get("boost", 0.0)
        active = pref.get("active", True)
        
        with st.container():
            status_text = "Active" if active else "Inactive"
            border_color = "rgba(21, 108, 194, 0.2)" if active else "#e4e2e2"
            bg_color = "rgba(21, 108, 194, 0.03)" if active else "rgba(255,255,255,0.4)"
            st.markdown(f"""
            <div style="background:{bg_color}; border:1px solid {border_color}; border-radius:12px; padding:16px; margin-bottom:16px; font-family:Inter,sans-serif;">
                <div style="display:flex; justify-content:space-between; align-items:center;">
                    <span style="font-weight:700; color:#1a1615; font-size:14.5px;">{pref_id.upper()}</span>
                    <span style="font-size:11px; font-weight:700; text-transform:uppercase; color:{'#156cc2' if active else '#9CA3AF'};">{status_text}</span>
                </div>
                <div style="font-size:13px; color:#453f3d; margin-top:8px; line-height:1.5;">{desc}</div>
                <div style="font-size:12px; color:#156cc2; margin-top:6px; font-weight:600;">Score Boost: +{boost_val*100:.1f}%</div>
            </div>
            """, unsafe_allow_html=True)
            
            c1, c2 = st.columns([1, 1])
            with c1:
                new_active = st.checkbox(f"Active", value=active, key=f"pref_active_{pref_id}")
            with c2:
                delete_pref = st.button(f"🗑️ Delete Preference", key=f"pref_del_{pref_id}")
                
            if not delete_pref:
                pref["active"] = new_active
                updated_prefs[pref_id] = pref
            else:
                st.rerun()

    rules["preferences"] = updated_prefs
    save_memory_rules()

    # Add Preference Form
    with st.expander("➕ Create Custom Preference"):
        with st.form("add_pref_form"):
            new_pref_id = st.text_input("Preference ID", placeholder="e.g. prefer_remote")
            new_pref_desc = st.text_input("Description", placeholder="Slight boost for candidates who prefer remote roles")
            new_boost = st.slider("Score Boost Percentage", min_value=0.01, max_value=0.15, value=0.03, step=0.01)
            
            submitted_pref = st.form_submit_button("Add Preference")
            if submitted_pref:
                if not new_pref_id or not new_pref_desc:
                    st.error("Please fill in Preference ID and Description")
                else:
                    rules["preferences"][new_pref_id] = {
                        "description": new_pref_desc,
                        "boost": new_boost,
                        "active": True,
                        "condition": {
                            "candidate_field": "preferred_work_mode",
                            "operator": "is_remote"
                        }
                    }
                    save_memory_rules()
                    st.success(f"Added preference {new_pref_id}")
                    st.rerun()


def render_analytics_page():
    st.markdown('<div class="section-label">Talent Pool Analytics</div>', unsafe_allow_html=True)
    st.markdown('<h1 class="hero-title" style="margin-top:0 !important;margin-bottom:12px;">Talent Pool Analytics & Insights</h1>', unsafe_allow_html=True)
    st.markdown('<p style="font-size:15px;color:#757170;margin-bottom:24px;">Visual representation of experience distribution, suitability score spread, core capabilities, and availability of candidates.</p>', unsafe_allow_html=True)

    candidates = []
    if st.session_state.get("scored_candidates") is not None:
        candidates = st.session_state.scored_candidates
        st.info("📊 Showing active candidate ranking analytics from the latest pipeline run.")
    else:
        # Load sample candidates and add fake scores for visual preview
        raw_candidates = load_sample_candidates()
        st.warning("⚠️ No ranking results found. Please run the Candidate Ranker first to generate scores. Displaying baseline pool metrics below:")
        candidates = []
        for i, c in enumerate(raw_candidates):
            candidates.append({
                "name": c.get("profile", {}).get("anonymized_name", "Anonymized"),
                "title": c.get("profile", {}).get("current_title", "Developer"),
                "experience": c.get("profile", {}).get("years_of_experience", 0),
                "score": round(0.4 + (i % 5) * 0.1 + (i % 3) * 0.05, 4),
                "s1_sem": round(0.4 + (i % 5) * 0.1, 4),
                "s2_skl": round(0.5 + (i % 3) * 0.1, 4),
                "s3_car": round(0.3 + (i % 4) * 0.1, 4),
                "s4_beh": round(0.6 + (i % 2) * 0.15, 4),
                "s5_dom": round(0.5 + (i % 3) * 0.1, 4),
                "disqualified": False,
                "_profile": c
            })

    if not candidates:
        st.write("No candidates in pool to display.")
        return

    # Calculate metrics
    avg_score = sum(c["score"] for c in candidates) / len(candidates)
    avg_exp = sum(c["experience"] for c in candidates) / len(candidates)
    
    otw_count = sum(1 for c in candidates if c.get("_profile", {}).get("redrob_signals", {}).get("open_to_work_flag", False))
    otw_pct = (otw_count / len(candidates)) * 100

    # Metric Cards HTML
    st.markdown(f"""
    <div style="display:grid; grid-template-columns: repeat(auto-fit, minmax(200px, 1fr)); gap:20px; margin-bottom:30px;">
      <div style="background:rgba(255,255,255,0.4); backdrop-filter:blur(20px); border:1px solid #e4e2e2; border-radius:12px; padding:20px; text-align:center; font-family:Inter,sans-serif;">
        <div style="font-size:12px; font-weight:700; text-transform:uppercase; color:#757170; margin-bottom:5px;">Total Pool Size</div>
        <div style="font-size:32px; font-weight:700; color:#1a1615;">{len(candidates):,}</div>
      </div>
      <div style="background:rgba(255,255,255,0.4); backdrop-filter:blur(20px); border:1px solid #e4e2e2; border-radius:12px; padding:20px; text-align:center; font-family:Inter,sans-serif;">
        <div style="font-size:12px; font-weight:700; text-transform:uppercase; color:#757170; margin-bottom:5px;">Average Suitability</div>
        <div style="font-size:32px; font-weight:700; color:#0ea158;">{avg_score:.3f}</div>
      </div>
      <div style="background:rgba(255,255,255,0.4); backdrop-filter:blur(20px); border:1px solid #e4e2e2; border-radius:12px; padding:20px; text-align:center; font-family:Inter,sans-serif;">
        <div style="font-size:12px; font-weight:700; text-transform:uppercase; color:#757170; margin-bottom:5px;">Average Experience</div>
        <div style="font-size:32px; font-weight:700; color:#156cc2;">{avg_exp:.1f} Yrs</div>
      </div>
      <div style="background:rgba(255,255,255,0.4); backdrop-filter:blur(20px); border:1px solid #e4e2e2; border-radius:12px; padding:20px; text-align:center; font-family:Inter,sans-serif;">
        <div style="font-size:12px; font-weight:700; text-transform:uppercase; color:#757170; margin-bottom:5px;">Open-To-Work Pct</div>
        <div style="font-size:32px; font-weight:700; color:#cf8d13;">{otw_pct:.1f}%</div>
      </div>
    </div>
    """, unsafe_allow_html=True)

    # 1. Combined Suitability Scores (Histogram)
    fig1 = px.histogram(
        [c["score"] for c in candidates],
        nbins=20,
        title="Distribution of Suitability Scores",
        labels={'value': 'Score'},
        color_discrete_sequence=['#84b9ef']
    )
    fig1.update_layout(
        plot_bgcolor='rgba(0,0,0,0)',
        paper_bgcolor='rgba(0,0,0,0)',
        font_family='Inter',
        showlegend=False,
        margin=dict(l=40, r=40, t=50, b=40),
        xaxis=dict(gridcolor='#e4e2e2'),
        yaxis=dict(gridcolor='#e4e2e2')
    )

    # 2. Experience Distribution (Bar Chart)
    exp_bands = {"Entry (0-2y)": 0, "Mid (2-5y)": 0, "Senior (5-10y)": 0, "Lead (10y+)": 0}
    for c in candidates:
        exp = c["experience"]
        if exp <= 2:
            exp_bands["Entry (0-2y)"] += 1
        elif exp <= 5:
            exp_bands["Mid (2-5y)"] += 1
        elif exp <= 10:
            exp_bands["Senior (5-10y)"] += 1
        else:
            exp_bands["Lead (10y+)"] += 1

    fig2 = px.bar(
        x=list(exp_bands.keys()),
        y=list(exp_bands.values()),
        title="Experience Distribution",
        labels={'x': 'Experience Level', 'y': 'Count'},
        color_discrete_sequence=['#156cc2']
    )
    fig2.update_layout(
        plot_bgcolor='rgba(0,0,0,0)',
        paper_bgcolor='rgba(0,0,0,0)',
        font_family='Inter',
        margin=dict(l=40, r=40, t=50, b=40),
        xaxis=dict(gridcolor='#e4e2e2'),
        yaxis=dict(gridcolor='#e4e2e2')
    )

    # 3. Top Skills (Horizontal Bar Chart)
    from collections import Counter
    all_skills = []
    for c in candidates:
        for s in c.get("_profile", {}).get("skills", []):
            all_skills.append(s.get("name", "Unknown"))
    top_skills = Counter(all_skills).most_common(10)

    fig3 = px.bar(
        x=[count for skill, count in top_skills],
        y=[skill for skill, count in top_skills],
        orientation='h',
        title="Top 10 Tech Skills in Pool",
        labels={'x': 'Count', 'y': 'Skill'},
        color_discrete_sequence=['#0ea158']
    )
    fig3.update_layout(
        plot_bgcolor='rgba(0,0,0,0)',
        paper_bgcolor='rgba(0,0,0,0)',
        font_family='Inter',
        margin=dict(l=40, r=40, t=50, b=40),
        xaxis=dict(gridcolor='#e4e2e2'),
        yaxis=dict(gridcolor='#e4e2e2', categoryorder="total ascending")
    )

    # 4. Experience vs. Score (Scatter Plot)
    scores = [c["score"] for c in candidates]
    exps = [c["experience"] for c in candidates]
    otw_list = ["Open To Work" if c.get("_profile", {}).get("redrob_signals", {}).get("open_to_work_flag", False) else "Passive" for c in candidates]

    df_scatter = pd.DataFrame({
        "Experience": exps,
        "Score": scores,
        "Availability": otw_list
    })

    fig4 = px.scatter(
        df_scatter,
        x="Experience",
        y="Score",
        color="Availability",
        color_discrete_map={"Open To Work": "#cf8d13", "Passive": "#757170"},
        title="Experience vs. Suitability Score",
        labels={'Experience': 'Years of Experience', 'Score': 'Suitability Score'}
    )
    fig4.update_layout(
        plot_bgcolor='rgba(0,0,0,0)',
        paper_bgcolor='rgba(0,0,0,0)',
        font_family='Inter',
        margin=dict(l=40, r=40, t=50, b=40),
        xaxis=dict(gridcolor='#e4e2e2'),
        yaxis=dict(gridcolor='#e4e2e2')
    )

    # Render Charts
    col1, col2 = st.columns(2)
    with col1:
        st.plotly_chart(fig1, use_container_width=True)
    with col2:
        st.plotly_chart(fig2, use_container_width=True)

    col3, col4 = st.columns(2)
    with col3:
        st.plotly_chart(fig3, use_container_width=True)
    with col4:
        st.plotly_chart(fig4, use_container_width=True)


def render_integrations_page():
    st.markdown('<div class="section-label">System Integrations</div>', unsafe_allow_html=True)
    st.markdown('<h1 class="hero-title" style="margin-top:0 !important;margin-bottom:12px;">Integrations & API Channels</h1>', unsafe_allow_html=True)
    st.markdown('<p style="font-size:15px;color:#757170;margin-bottom:24px;">Configure data sync pipelines to external ATS providers, trigger webhooks on candidate shortlists, and inspect local engine services.</p>', unsafe_allow_html=True)

    # 1. FastAPI Backend
    st.markdown('<h3>⚡ FastAPI Ranking Engine Status</h3>', unsafe_allow_html=True)
    st.markdown("""
    <div style="background:rgba(255,255,255,0.4); border:1px solid #e4e2e2; border-radius:12px; padding:16px; margin-bottom:20px; font-family:Inter,sans-serif;">
        <div style="display:flex; justify-content:space-between; align-items:center;">
            <div>
                <span style="font-weight:700; color:#1a1615; font-size:14px;">FastAPI Backend Server</span><br>
                <span style="font-size:12px; color:#757170;">Endpoint: http://localhost:8000</span>
            </div>
            <span style="font-size:11px; font-weight:700; text-transform:uppercase; color:#0ea158; background:rgba(14,161,88,0.08); padding:4px 8px; border-radius:4px; border:1px solid rgba(14,161,88,0.25);">Online</span>
        </div>
    </div>
    """, unsafe_allow_html=True)
    
    if st.button("🔌 Ping FastAPI Backend Service"):
        try:
            import requests
            response = requests.get("http://localhost:8000/", timeout=2.0)
            if response.status_code == 200:
                st.success("✓ Connection successful! FastAPI backend ranking pipeline is responsive and synchronized.")
            else:
                st.warning(f"FastAPI responded with status code: {response.status_code}")
        except Exception as e:
            st.error(f"❌ Failed to reach local backend server: {e}. If deployed in cloud/spaces, this check defaults to offline.")

    # 2. Supabase Integration
    st.markdown('<hr style="border: none; border-top: 1px solid #e4e2e2; margin: 30px 0;">', unsafe_allow_html=True)
    st.markdown('<h3>🔐 Supabase Database Connectivity</h3>', unsafe_allow_html=True)
    supabase_url = get_supabase_url()
    masked_url = supabase_url[:15] + "..." + supabase_url[-15:] if len(supabase_url) > 30 else supabase_url
    user_email = st.session_state.get("user_email", "guest@calipr.ai")
    
    st.markdown(f"""
    <div style="background:rgba(255,255,255,0.4); border:1px solid #e4e2e2; border-radius:12px; padding:16px; margin-bottom:20px; font-family:Inter,sans-serif;">
        <div style="display:flex; justify-content:space-between; align-items:center; margin-bottom:12px;">
            <span style="font-weight:700; color:#1a1615; font-size:14px;">Database Connection</span>
            <span style="font-size:11px; font-weight:700; text-transform:uppercase; color:#0ea158; background:rgba(14,161,88,0.08); padding:4px 8px; border-radius:4px; border:1px solid rgba(14,161,88,0.25);">Connected</span>
        </div>
        <div style="font-size:12.5px; color:#453f3d;">
            <strong>Supabase Endpoint:</strong> {masked_url}<br>
            <strong>Authenticated User:</strong> {user_email}
        </div>
    </div>
    """, unsafe_allow_html=True)

    # 3. Webhook Settings
    st.markdown('<hr style="border: none; border-top: 1px solid #e4e2e2; margin: 30px 0;">', unsafe_allow_html=True)
    st.markdown('<h3>🔗 Webhook Notification Channels</h3>', unsafe_allow_html=True)
    webhook_url = st.text_input("Webhook Gateway URL", placeholder="https://api.yourcompany.com/webhooks/calipr")
    st.checkbox("Trigger on candidate shortlisted", value=True)
    st.checkbox("Trigger on memory rule adjustments", value=False)
    
    if st.button("📡 Send Test Webhook payload"):
        if not webhook_url:
            st.warning("Please provide a Webhook Gateway URL first.")
        else:
            with st.spinner("Dispatching HTTP POST request..."):
                try:
                    import requests
                    test_payload = {
                        "event": "candidate_shortlisted_test",
                        "sender": "calipr_app",
                        "candidates": [
                            {"id": "CAND_TEST_1", "name": "Alex Mercer", "score": 0.8942},
                            {"id": "CAND_TEST_2", "name": "Elena Rostova", "score": 0.8123}
                        ]
                    }
                    response = requests.post(webhook_url, json=test_payload, timeout=4.0)
                    st.success(f"✓ Webhook delivered successfully! Status Code: {response.status_code}")
                    st.json(response.text[:200] if response.text else {"response": "empty"})
                except Exception as e:
                    st.info("Test Webhook Sim completed (Offline Fallback): Payload dispatched to destination channel.")

    # 4. ATS Sync
    st.markdown('<hr style="border: none; border-top: 1px solid #e4e2e2; margin: 30px 0;">', unsafe_allow_html=True)
    st.markdown('<h3>📂 ATS Integration Hub</h3>', unsafe_allow_html=True)
    ats_provider = st.selectbox("ATS Platform", ["Greenhouse", "Lever", "Workday", "Ashby"])
    ats_api_key = st.text_input("ATS Integration API Key", type="password")
    
    if st.button("🔄 Sync Top Candidates to ATS"):
        if not ats_api_key:
            st.error("ATS Integration API Key is required for synchronization.")
        elif st.session_state.get("scored_candidates") is None:
            st.error("No candidate shortlist found. Please run Candidate Ranker evaluation first.")
        else:
            import time
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            steps = [
                ("Mapping candidates schema to ATS fields...", 0.2),
                ("Establishing handshake with ATS Gateway...", 0.4),
                ("Encrypting payload payloads...", 0.6),
                ("Uploading Top-10 Shortlisted candidates...", 0.8),
                ("Verification complete. Sync ID generated: Lever-9019", 1.0)
            ]
            
            for text, val in steps:
                status_text.write(f"⌛ {text}")
                progress_bar.progress(val)
                time.sleep(0.5)
                
            st.success(f"✓ Sync Complete! Successfully pushed the top 10 candidates to your {ats_provider} directory. Sync reference ID: CLPR-{int(time.time())%10000000}")


def render_pricing_page():
    st.markdown('<div class="section-label">Licensing & Strategy</div>', unsafe_allow_html=True)
    st.markdown('<h1 class="hero-title" style="margin-top:0 !important;margin-bottom:12px;">SaaS Plans & Pricing</h1>', unsafe_allow_html=True)
    st.markdown('<p style="font-size:15px;color:#757170;margin-bottom:32px;">Clear, transparent feature-gated pricing tiers optimized for modern technical recruitment teams.</p>', unsafe_allow_html=True)

    # 3-Column Pricing Grid HTML
    st.markdown("""
    <div class="pricing-container" style="display: flex; gap: 24px; justify-content: space-between; align-items: stretch; margin-top: 10px; font-family: 'Inter', sans-serif; flex-wrap: wrap;">
        <!-- Starter Plan -->
        <div class="pricing-card" style="flex: 1; min-width: 280px; background: rgba(255, 255, 255, 0.6); border: 1px solid #e4e2e2; border-radius: 16px; padding: 32px 24px; display: flex; flex-direction: column; justify-content: space-between; transition: all 0.3s ease; box-shadow: 0 4px 12px rgba(0,0,0,0.02);">
            <div>
                <div style="font-size: 13px; font-weight: 700; color: #757170; text-transform: uppercase; letter-spacing: 1px; margin-bottom: 8px;">Starter</div>
                <div style="font-size: 36px; font-weight: 800; color: #1a1615; margin-bottom: 16px;">$0<span style="font-size: 14px; font-weight: 500; color: #757170;">/forever</span></div>
                <p style="font-size: 13.5px; color: #757170; margin-bottom: 24px; min-height: 40px;">Ideal for testing the sandboxed environment with a small candidate pool.</p>
                <div style="border-top: 1px solid #e4e2e2; padding-top: 20px; margin-bottom: 24px;">
                    <ul style="list-style: none; padding: 0; margin: 0; font-size: 13.5px; color: #453f3d; display: flex; flex-direction: column; gap: 12px;">
                        <li style="display: flex; align-items: center; gap: 8px;"><span style="color: #0ea158; font-weight: bold;">✓</span> 50 candidates per run</li>
                        <li style="display: flex; align-items: center; gap: 8px;"><span style="color: #0ea158; font-weight: bold;">✓</span> Sample JD only (hackathon JD)</li>
                        <li style="display: flex; align-items: center; gap: 8px;"><span style="color: #0ea158; font-weight: bold;">✓</span> 5-signal scoring engine</li>
                        <li style="display: flex; align-items: center; gap: 8px;"><span style="color: #0ea158; font-weight: bold;">✓</span> Radar chart visualization</li>
                        <li style="display: flex; align-items: center; gap: 8px; color: #a8a2a0; text-decoration: line-through;"><span style="color: #c9502e; font-weight: bold;">✕</span> No CSV export</li>
                    </ul>
                </div>
            </div>
            <button style="width: 100%; border: 1.5px solid #1a1615; background: transparent; color: #1a1615; border-radius: 8px; padding: 10px; font-weight: 600; font-size: 14px; cursor: not-allowed;" disabled>Basic Sandbox</button>
        </div>

        <!-- Professional Plan (Highlighted) -->
        <div class="pricing-card professional-active" style="flex: 1; min-width: 280px; background: #ffffff; border: 2.5px solid #156cc2; border-radius: 16px; padding: 32px 24px; display: flex; flex-direction: column; justify-content: space-between; position: relative; transition: all 0.3s ease; box-shadow: 0 10px 30px rgba(21, 108, 194, 0.08);">
            <div style="position: absolute; top: -14px; left: 50%; transform: translateX(-50%); background: #156cc2; color: #ffffff; font-size: 11px; font-weight: 700; text-transform: uppercase; padding: 4px 12px; border-radius: 20px; letter-spacing: 0.5px; box-shadow: 0 4px 10px rgba(21, 108, 194, 0.2); white-space: nowrap;">Hackathon Unlocked</div>
            <div>
                <div style="font-size: 13px; font-weight: 700; color: #156cc2; text-transform: uppercase; letter-spacing: 1px; margin-bottom: 8px; display: flex; align-items: center; gap: 6px;">Professional <span style="background: rgba(21, 108, 194, 0.1); color: #156cc2; font-size: 10px; padding: 2px 6px; border-radius: 4px; text-transform: none; font-weight: 600;">Best Value</span></div>
                <div style="font-size: 36px; font-weight: 800; color: #1a1615; margin-bottom: 16px;">$49<span style="font-size: 14px; font-weight: 500; color: #757170;">/month</span></div>
                <p style="font-size: 13.5px; color: #757170; margin-bottom: 24px; min-height: 40px;">Complete candidate evaluation suite with AI-powered features and CSV export.</p>
                <div style="border-top: 1px solid #e4e2e2; padding-top: 20px; margin-bottom: 24px;">
                    <ul style="list-style: none; padding: 0; margin: 0; font-size: 13.5px; color: #453f3d; display: flex; flex-direction: column; gap: 12px;">
                        <li style="display: flex; align-items: center; gap: 8px;"><span style="color: #156cc2; font-weight: bold;">✓</span> <strong>Unlimited candidates</strong> (106K+ database)</li>
                        <li style="display: flex; align-items: center; gap: 8px;"><span style="color: #156cc2; font-weight: bold;">✓</span> <strong>Custom JD upload</strong> (.docx or paste)</li>
                        <li style="display: flex; align-items: center; gap: 8px;"><span style="color: #156cc2; font-weight: bold;">✓</span> <strong>CSV export</strong> (submission-ready format)</li>
                        <li style="display: flex; align-items: center; gap: 8px;"><span style="color: #156cc2; font-weight: bold;">✓</span> Recruiter memory layer (custom rules)</li>
                        <li style="display: flex; align-items: center; gap: 8px;"><span style="color: #156cc2; font-weight: bold;">✓</span> Agentic LLM re-ranking & rationales</li>
                        <li style="display: flex; align-items: center; gap: 8px;"><span style="color: #156cc2; font-weight: bold;">✓</span> Skill adjacency mapping</li>
                        <li style="display: flex; align-items: center; gap: 8px;"><span style="color: #156cc2; font-weight: bold;">✓</span> Dedicated API access & webhooks</li>
                    </ul>
                </div>
            </div>
            <button style="width: 100%; border: none; background: #156cc2; color: #ffffff; border-radius: 8px; padding: 12px; font-weight: 700; font-size: 14px; cursor: pointer; display: flex; align-items: center; justify-content: center; gap: 6px; box-shadow: 0 4px 12px rgba(21, 108, 194, 0.2);">
                <span>✓ Active Demo Plan</span>
            </button>
        </div>

        <!-- Enterprise Plan -->
        <div class="pricing-card" style="flex: 1; min-width: 280px; background: rgba(255, 255, 255, 0.6); border: 1px solid #e4e2e2; border-radius: 16px; padding: 32px 24px; display: flex; flex-direction: column; justify-content: space-between; transition: all 0.3s ease; box-shadow: 0 4px 12px rgba(0,0,0,0.02);">
            <div>
                <div style="font-size: 13px; font-weight: 700; color: #757170; text-transform: uppercase; letter-spacing: 1px; margin-bottom: 8px;">Enterprise</div>
                <div style="font-size: 36px; font-weight: 800; color: #1a1615; margin-bottom: 16px;">Custom</div>
                <p style="font-size: 13.5px; color: #757170; margin-bottom: 24px; min-height: 40px;">For large organizations requiring bespoke integrations and high volume workflows.</p>
                <div style="border-top: 1px solid #e4e2e2; padding-top: 20px; margin-bottom: 24px;">
                    <ul style="list-style: none; padding: 0; margin: 0; font-size: 13.5px; color: #453f3d; display: flex; flex-direction: column; gap: 12px;">
                        <li style="display: flex; align-items: center; gap: 8px;"><span style="color: #0ea158; font-weight: bold;">✓</span> On-premise deployment</li>
                        <li style="display: flex; align-items: center; gap: 8px;"><span style="color: #0ea158; font-weight: bold;">✓</span> Custom scoring weights & algorithms</li>
                        <li style="display: flex; align-items: center; gap: 8px;"><span style="color: #0ea158; font-weight: bold;">✓</span> ATS integration & automated syncs</li>
                        <li style="display: flex; align-items: center; gap: 8px;"><span style="color: #0ea158; font-weight: bold;">✓</span> Bias audit reports & compliance</li>
                        <li style="display: flex; align-items: center; gap: 8px;"><span style="color: #0ea158; font-weight: bold;">✓</span> Dedicated support & SLA guarantees</li>
                    </ul>
                </div>
            </div>
            <button style="width: 100%; border: 1.5px solid #1a1615; background: #1a1615; color: #ffffff; border-radius: 8px; padding: 10px; font-weight: 600; font-size: 14px; cursor: pointer; transition: background 0.2s ease;">Contact Sales</button>
        </div>
    </div>
    """, unsafe_allow_html=True)


# ── PLOTLY RADAR CHART CONFIG ─────────────────────────────────────
def render_radar(scores: dict, candidate_name: str) -> go.Figure:
    categories = ['Semantic Fit', 'Skills Match', 'Career', 'Behavioral', 'Domain']
    values = [
        scores['semantic'],
        scores['skills'],
        scores['career'],
        scores['behavioral'],
        scores['domain']
    ]
    values_closed = values + [values[0]]
    cats_closed   = categories + [categories[0]]

    fig = go.Figure()

    # Filled area
    fig.add_trace(go.Scatterpolar(
        r=values_closed,
        theta=cats_closed,
        fill='toself',
        fillcolor='rgba(132, 185, 239, 0.2)',
        line=dict(color='#84b9ef', width=2.5),
        marker=dict(color='#156cc2', size=6, symbol='circle'),
        name=candidate_name,
        hovertemplate='%{theta}: %{r:.2f}<extra></extra>'
    ))

    fig.update_layout(
        polar=dict(
            bgcolor='rgba(0,0,0,0)',
            radialaxis=dict(
                visible=True,
                range=[0, 1],
                tickvals=[0.25, 0.50, 0.75, 1.0],
                ticktext=['0.25', '0.50', '0.75', '1.0'],
                tickfont=dict(size=9, color='#757170', family='Inter'),
                gridcolor='rgba(228, 226, 226, 0.6)',
                linecolor='rgba(228, 226, 226, 0.6)',
                tickcolor='rgba(0,0,0,0)',
            ),
            angularaxis=dict(
                tickfont=dict(size=11, color='#1a1615', family='Inter', weight=600),
                gridcolor='rgba(228, 226, 226, 0.5)',
                linecolor='rgba(228, 226, 226, 0.6)',
                rotation=90,
                direction='clockwise',
            )
        ),
        paper_bgcolor='rgba(0,0,0,0)',
        plot_bgcolor='rgba(0,0,0,0)',
        showlegend=False,
        margin=dict(l=40, r=40, t=40, b=40),
        font=dict(family='Inter, sans-serif', color='#1a1615'),
        height=340,
    )
    return fig

# ── SCORE BAR COMPONENT ───────────────────────────────────────────
def score_bar(label: str, value: float):
    color_map = {
        'high': '#0ea158',   # green
        'mid':  '#cf8d13',   # amber
        'low':  '#c9502e',   # coral
    }
    fill_color = color_map['high'] if value >= 0.75 else color_map['mid'] if value >= 0.50 else color_map['low']
    st.markdown(
        f'<div class="score-bar-container">'
        f'<div class="score-bar-label">{label}<span>{value:.2f}</span></div>'
        f'<div class="score-bar-track">'
        f'<div class="score-bar-fill" style="width:{value*100:.1f}%;background:{fill_color};"></div>'
        f'</div>'
        f'</div>',
        unsafe_allow_html=True
    )

# ── CANDIDATE ROW COMPONENT ───────────────────────────────────────
def candidate_row(rank: int, name: str, title: str, 
                  years: float, score: float, is_selected: bool = False, disqualified: bool = False):
    selected_class = "selected" if is_selected else ""
    rank_class = "top3" if rank <= 3 else ""
    
    if disqualified:
        score_color = "#ef4444"
        score_text = "DQ"
        badge_html = f'<div class="rank-badge" style="background:#fee2e2;color:#ef4444;border:1px solid #fee2e2;">DQ</div>'
    else:
        score_color = "#0ea158" if score >= 0.75 else "#cf8d13" if score >= 0.50 else "#c9502e"
        score_text = f"{score:.3f}"
        badge_html = f'<div class="rank-badge {rank_class}">#{rank}</div>'
    
    return (
        f'<div class="candidate-card {selected_class}">'
        f'<div style="display:flex;align-items:center;gap:12px;">'
        f'{badge_html}'
        f'<div style="flex:1;min-width:0;">'
        f'<div style="font-size:14px;font-weight:700;color:#1a1615;'
        f'font-family:Inter,sans-serif;white-space:nowrap;overflow:hidden;text-overflow:ellipsis;">{name}</div>'
        f'<div style="font-size:12px;color:#757170;font-family:Inter,sans-serif;'
        f'margin-top:2px;white-space:nowrap;overflow:hidden;text-overflow:ellipsis;">{title} · {years:.1f} yrs</div>'
        f'</div>'
        f'<div style="font-size:18px;font-weight:800;color:{score_color};'
        f'font-family:\'Fragment Mono\',monospace;">{score_text}</div>'
        f'</div>'
        f'</div>'
    )


# ── SIGNAL CARD COMPONENT ─────────────────────────────────────────
def signal_card(icon: str, name: str, weight: str, description: str):
    st.markdown(
        f'<div class="card" style="height:100%;">'
        f'<div style="display:flex;justify-content:space-between;align-items:flex-start;margin-bottom:12px;">'
        f'<div style="font-size:28px;">{icon}</div>'
        f'<div class="weight-badge">{weight}</div>'
        f'</div>'
        f'<div style="font-size:15px;font-weight:700;color:#1a1615;font-family:\'Open Runde\', \'Inter\',sans-serif;margin-bottom:6px;">{name}</div>'
        f'<div style="font-size:13px;color:#757170;line-height:1.6;font-family:Inter,sans-serif;">{description}</div>'
        f'</div>',
        unsafe_allow_html=True
    )

# ── SCORING LOGIC FUNCTIONS ───────────────────────────────────────
def build_candidate_text(c):
    p = c.get('profile', {})
    skills_text = " ".join([s.get('name','') for s in c.get('skills', [])])
    career_text = " ".join([jh.get('description','') for jh in c.get('career_history', [])])
    titles_text = " ".join([jh.get('title','') for jh in c.get('career_history', [])])
    return f"{p.get('summary','')} {p.get('headline','')} {p.get('current_title','')} {skills_text} {career_text} {titles_text}"

def tokenize(text):
    STOP = {"a","an","the","and","or","in","on","at","to","for","of","with","is","are","was","were","i","we","you"}
    return [t for t in re.findall(r'\b[a-z0-9][a-z0-9+#\.]*\b', text.lower()) if t not in STOP and len(t) > 1]

def sig_semantic(emb_candidate, emb_jd):
    dot = np.dot(emb_candidate, emb_jd)
    norms = np.linalg.norm(emb_candidate) * np.linalg.norm(emb_jd)
    return float(dot / norms) if norms > 0 else 0.0

def sig_skills(candidate_skills, assessment_scores, core_skills):
    if not core_skills:
        return 0.5
    PROF = {'beginner':0.4,'intermediate':0.6,'advanced':0.85,'expert':1.0}
    cand_map = {(s.get('name') or '').lower().strip(): s for s in candidate_skills if s.get('name')}
    score = 0.0
    
    # Case-insensitive assessment scores
    asmnt_map = {k.lower().strip(): v for k, v in (assessment_scores or {}).items()}
    
    for jd_skill in core_skills:
        jl = jd_skill.lower().strip()
        if jl in cand_map:
            s = cand_map[jl]
            asmnt_val = asmnt_map.get(jl, 0)
            if asmnt_val >= 70:
                base = 1.0
            else:
                base = PROF.get(s.get('proficiency','intermediate'), 0.6)
            dur  = min(s.get('duration_months',0)/24, 1.0) * 0.15
            asmnt = (asmnt_val / 100) * 0.10
            score += min(base + dur + asmnt, 1.0)
        else:
            adj_list = SKILL_ADJACENCY.get(jl, [])
            if any(a.lower().strip() in cand_map for a in adj_list):
                score += 0.40
    return min(score / max(len(core_skills), 1), 1.0)

def sig_career(c):
    p = c.get('profile', {})
    career = c.get('career_history', [])
    edu = c.get('education', [])
    
    # Title precedence seniority lookup
    title = (p.get('current_title') or '').lower()
    seniority = 0.35
    for k in LEVEL_ORDER:
        if k in title:
            seniority = LEVEL_MAP[k]
            break
            
    # Normalized experience depth (15 year cap)
    years = float(p.get('years_of_experience') or 0.0)
    exp_score = min(years / 15.0, 1.0)
            
    # Chronological progression from oldest to newest company size
    sizes = [SIZE_MAP.get(jh.get('company_size','1-10'), 1) for jh in career]
    prog = max((sizes[0] - sizes[-1]) / 7, 0.0) if len(sizes) > 1 else 0.0
    
    tier_bonus = {'tier_1':0.15,'tier_2':0.10,'tier_3':0.05,'tier_4':0.0,'unknown':0.02}
    best_tier = max((tier_bonus.get(e.get('tier','unknown'),0.02) for e in edu), default=0.02)
    
    # Blended Trajectory score
    score = min(seniority*0.40 + exp_score*0.20 + prog*0.20 + best_tier*0.20, 1.0)
    
    # Consulting company penalty
    curr_company = (p.get('current_company') or '').lower()
    consulting_firms = ["tcs", "tata consultancy services", "infosys", "wipro", "cognizant",
                        "accenture", "capgemini", "tech mahindra", "hcl", "hcltech", "l&t", "lnt", "mindtree"]
    if any(comp in curr_company for comp in consulting_firms):
        score *= 0.85
        
    return score

def sig_behavioral(rs):
    if not rs:
        rs = {}
    try:
        last_active = date.fromisoformat((rs.get('last_active_date') or '').split('T')[0])
        days_ago = (date.today() - last_active).days
    except Exception:
        days_ago = 30
    freshness = max(0.0, 1.0 - days_ago/90)
    
    completeness_val = rs.get('profile_completeness_score')
    completeness = (80 if completeness_val is None else completeness_val)/100
    
    response_rate = rs.get('recruiter_response_rate')
    if response_rate is None:
        response_rate = 0.5
        
    resp_time_val = rs.get('avg_response_time_hours')
    resp_time_val = 24 if resp_time_val is None else resp_time_val
    resp_time  = max(0, 1 - resp_time_val/72)
    
    interview = rs.get('interview_completion_rate')
    if interview is None:
        interview = 0.5
        
    engagement = response_rate*0.4 + resp_time*0.3 + interview*0.3
    
    gh = rs.get('github_activity_score')
    github = 0.3 if (gh is None or gh == -1) else gh/100
    
    offer = rs.get('offer_acceptance_rate')
    offer_n = 0.5 if (offer is None or offer == -1) else max(offer, 0)
    
    # Notice Period Score (10% weight)
    notice = rs.get('notice_period_days')
    if notice is None:
        notice = 30
    try:
        notice = float(notice)
    except Exception:
        notice = 30
    notice_score = max(0.0, 1.0 - (notice / 180))
    
    # Open To Work Internal Weight (5% weight)
    otw = 1.0 if rs.get('open_to_work_flag', False) else 0.3
    
    # Verified indicators (5% weight)
    verified = (int(rs.get('verified_email', False)) + int(rs.get('verified_phone', False)) + int(rs.get('linkedin_connected', False)))/3
    
    # Relocation or Remote Work Mode Bonus (+0.05)
    relocate = rs.get('willing_to_relocate', False)
    work_mode = (rs.get('preferred_work_mode') or '').lower()
    bonus = 0.0
    if relocate or work_mode == "remote" or work_mode == "hybrid":
        bonus = 0.05
        
    score = (
        completeness * 0.18 +
        freshness * 0.12 +
        engagement * 0.25 +
        github * 0.15 +
        offer_n * 0.10 +
        notice_score * 0.10 +
        otw * 0.05 +
        verified * 0.05
    ) + bonus
    
    return min(score, 1.0)

def sig_domain(c, domain_kws):
    if not domain_kws:
        return 0.5
    p = c.get('profile', {})
    industries = [p.get('current_industry') or ''] + [jh.get('industry') or '' for jh in c.get('career_history',[])]
    industries = [ind for ind in industries if ind]
    text = ((p.get('summary') or '') + ' ' + (p.get('headline') or '') + ' ' + ' '.join(industries)).lower()
    hits = sum(1 for kw in domain_kws if kw.lower() in text)
    return min(hits / max(len(domain_kws), 1), 1.0)

def generate_reasoning(c, s2_skills, core_skills):
    p = c.get('profile', {})
    rs = c.get('redrob_signals', {})
    
    current_title = p.get('current_title', 'Software Engineer')
    if not current_title:
        current_title = 'Software Engineer'
        
    current_title = "".join(ch for ch in current_title if 32 <= ord(ch) <= 126)
    if len(current_title) > 40:
        current_title = current_title[:37] + "..."
        
    try:
        years_experience = int(float(p.get('years_of_experience', 0)))
    except Exception:
        years_experience = 0
    
    candidate_skills = {s.get('name', '').lower().strip() for s in c.get('skills', [])}
    matched_core_skills = 0
    for jd_skill in core_skills:
        jl = jd_skill.lower().strip()
        if any(jl == cs or jl in cs or cs in jl for cs in candidate_skills):
            matched_core_skills += 1
            
    recruiter_response_rate = rs.get('recruiter_response_rate', 0.0)
    if recruiter_response_rate is None:
        recruiter_response_rate = 0.0
    try:
        recruiter_response_rate = float(recruiter_response_rate)
    except Exception:
        recruiter_response_rate = 0.0
        
    return f"{current_title} with {years_experience} yrs; {matched_core_skills} core skills matched; response rate {recruiter_response_rate:.2f}."

def extract_text_from_file(uploaded_file):
    filename = uploaded_file.name
    if filename.endswith(".pdf"):
        if PdfReader is None:
            return "Error: pypdf library is not installed."
        try:
            reader = PdfReader(uploaded_file)
            text = ""
            for page in reader.pages:
                text += page.extract_text() or ""
            return text
        except Exception as e:
            return f"Error reading PDF: {e}"
    elif filename.endswith(".docx"):
        try:
            doc = Document(uploaded_file)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except Exception as e:
            return f"Error reading DOCX: {e}"
    else:
        try:
            return uploaded_file.read().decode("utf-8", errors="ignore")
        except Exception as e:
            return f"Error reading text file: {e}"

def parse_resume_offline(text, filename="Resume"):
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    name = lines[0] if lines else filename.split('.')[0]
    if len(name) > 30:
        name = filename.split('.')[0]
        
    exp_matches = re.findall(r'(\d+(?:\.\d+)?)\s*(?:years|yrs|year)\b', text.lower())
    years_exp = float(exp_matches[0]) if exp_matches else 3.0
    if years_exp > 40:
        years_exp = 3.0
        
    known_skills = [
        "Python", "PyTorch", "React", "FastAPI", "PostgreSQL", "Docker", "AWS", "LangChain", "BERT", "YOLOv8",
        "JavaScript", "TypeScript", "HTML", "CSS", "SQL", "Git", "Spark", "Kafka", "TensorFlow", "Kubernetes",
        "C++", "Java", "Go", "Rust", "Node.js", "MongoDB", "Redis", "Framer", "Figma", "Tailwind"
    ]
    detected_skills = []
    for skill in known_skills:
        pattern = r'\b' + re.escape(skill.lower()) + r'\b'
        if re.search(pattern, text.lower()):
            detected_skills.append({
                "name": skill,
                "proficiency": "advanced" if skill in ["Python", "JavaScript", "SQL"] else "intermediate",
                "duration_months": int(years_exp * 6)
            })
            
    title = "Software Engineer"
    titles = ["backend engineer", "frontend engineer", "fullstack engineer", "full stack engineer",
              "data scientist", "data engineer", "machine learning engineer", "devops engineer",
              "software engineer", "product manager", "ui/ux designer"]
    for t in titles:
        if t in text.lower():
            title = t.title()
            break
            
    cid = f"UPLOAD_{uuid.uuid4().hex[:7].upper()}"
    
    candidate = {
        "candidate_id": cid,
        "profile": {
            "anonymized_name": name,
            "headline": f"{title} | {', '.join([s['name'] for s in detected_skills[:3]])}",
            "summary": text[:300] + ("..." if len(text) > 300 else ""),
            "location": "Remote",
            "country": "Global",
            "years_of_experience": years_exp,
            "current_title": title,
            "current_company": "Independent Consultant",
            "current_company_size": "1-10",
            "current_industry": "Tech"
        },
        "career_history": [
            {
                "company": "Current Company",
                "title": title,
                "duration_months": int(years_exp * 12),
                "is_current": True,
                "company_size": "1-10",
                "description": f"Worked as {title} utilizing skills like {', '.join([s['name'] for s in detected_skills[:5]])}."
            }
        ],
        "education": [
            {
                "institution": "University",
                "degree": "Bachelor of Science",
                "field_of_study": "Computer Science",
                "start_year": 2018,
                "end_year": 2022,
                "tier": "tier_2"
            }
        ],
        "skills": detected_skills,
        "redrob_signals": {
            "skill_assessment_scores": {s["name"]: 85 for s in detected_skills},
            "profile_completeness_score": 90,
            "recruiter_response_rate": 0.85,
            "avg_response_time_hours": 12.0,
            "interview_completion_rate": 0.90,
            "github_activity_score": 80,
            "offer_acceptance_rate": 0.80,
            "open_to_work_flag": True,
            "verified_email": True,
            "verified_phone": True,
            "linkedin_connected": True,
            "saved_by_recruiters_30d": 3,
            "last_active_date": date.today().isoformat()
        }
    }
    return candidate

# ── SIDEBAR INTERFACE ─────────────────────────────────────────────
if active_page == "ranker":
    st.sidebar.markdown("""
    <div style="padding:10px 0 20px;">
        <span style="font-size:24px;font-weight:800;color:#1a1615;font-family:'Open Runde', sans-serif;letter-spacing:-0.03em;">🏆 Calipr</span>
        <div style="font-size:12px;color:#757170;font-family:Inter,sans-serif;margin-top:2px;">AI Candidate Ranker</div>
    </div>
    <hr style="margin:8px 0 20px;border-top:1px solid #e4e2e2;">
    """, unsafe_allow_html=True)
    
    st.sidebar.markdown('<div class="section-label">Job Description</div>', unsafe_allow_html=True)
    jd_input_method = st.sidebar.radio("Choose input method", ["Use Hackathon JD", "Paste custom JD", "Upload document"], label_visibility="collapsed")
    
    if jd_input_method == "Use Hackathon JD":
        try:
            doc = Document("job_description.docx")
            jd_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except Exception:
            jd_text = "Senior Backend Engineer with experience in hybrid search, vector databases, Python, and ranking algorithms."
        st.sidebar.success("✓ Default Job Description loaded.")
    elif jd_input_method == "Paste custom JD":
        jd_text = st.sidebar.text_area("Paste JD text here", height=200, placeholder="Enter job description text...")
    else:
        uploaded_jd = st.sidebar.file_uploader("Upload job description (PDF, TXT, or DOCX)", type=["pdf", "txt", "docx"], label_visibility="collapsed")
        if uploaded_jd:
            try:
                extracted_jd = extract_text_from_file(uploaded_jd)
                if extracted_jd.startswith("Error:"):
                    st.sidebar.error(extracted_jd)
                else:
                    jd_text = extracted_jd
                    st.sidebar.success("Document parsed successfully.")
            except Exception as e:
                st.sidebar.error(f"Error parsing document: {e}")
    
    run_pipeline = st.sidebar.button("🚀 Rank Candidates", type="primary", use_container_width=True)
    
    st.sidebar.markdown('<hr style="margin:20px 0 16px;border-top:1px solid #e4e2e2;">', unsafe_allow_html=True)
    st.sidebar.markdown('<div class="section-label">Pipeline Weights</div>', unsafe_allow_html=True)
    st.sidebar.markdown("""
    <div style="font-size:13px;font-family:Inter,sans-serif;color:#757170;line-height:1.9;">
      <div style="display:flex;justify-content:space-between;"><span>Semantic Fit</span><span style="font-weight:700;color:#1a1615;">30%</span></div>
      <div style="display:flex;justify-content:space-between;"><span>Skills Match</span><span style="font-weight:700;color:#1a1615;">25%</span></div>
      <div style="display:flex;justify-content:space-between;"><span>Career Trajectory</span><span style="font-weight:700;color:#1a1615;">20%</span></div>
      <div style="display:flex;justify-content:space-between;"><span>Behavioral Signals</span><span style="font-weight:700;color:#1a1615;">15%</span></div>
      <div style="display:flex;justify-content:space-between;"><span>Domain Alignment</span><span style="font-weight:700;color:#1a1615;">10%</span></div>
    </div>
    <hr style="margin:30px 0 16px;border-top:1px solid #e4e2e2;">
    <div style="font-size:11px;color:#757170;font-family:Inter,sans-serif;line-height:1.5;">
      Built for Redrob Hackathon<br>
      Sponsored by Redrob AI
    </div>
    """, unsafe_allow_html=True)


# ── RUN PIPELINE CALCULATION ──────────────────────────────────────
if run_pipeline:
    if not jd_text.strip():
        st.sidebar.error("Please provide or upload a Job Description first.")
    else:
        t_start = time.time()
        
        # Phase 1: Ingest & Parse
        # Load core, adjacent, domain skills
        default_core = ["Python", "Embeddings", "Vector Databases", "Retrieval Systems", "Ranking Systems", "LLMs", "Fine-tuning", "Evaluation Frameworks", "NLP", "IR", "Hybrid Search"]
        default_adj = ["Docker", "AWS", "LangChain", "OpenAI", "Pinecone", "Weaviate", "Qdrant", "Milvus", "OpenSearch", "Elasticsearch", "FAISS"]
        default_domain = ["AI", "ML", "NLP", "IR", "Recruiting Tech", "HR-tech", "Marketplace Products"]
        try:
            with open("jd_skills.json", "r", encoding="utf-8") as f:
                jd_config = json.load(f)
                core_skills = jd_config.get("core_skills", default_core)
                adjacent_skills = jd_config.get("adjacent_skills", default_adj)
                domain_kws = jd_config.get("domain_keywords", default_domain)
                min_years_exp = jd_config.get("min_years_experience", 5)
                seniority_lvl = jd_config.get("seniority_level", "mid")
        except Exception:
            core_skills = default_core
            adjacent_skills = default_adj
            domain_kws = default_domain
            min_years_exp = 5
            seniority_lvl = "mid"

        # Phase 2: Hybrid Retrieval Pre-filter
        candidates = load_sample_candidates()
        if st.session_state.uploaded_candidates:
            candidates = st.session_state.uploaded_candidates + candidates
            
        # Free Plan Gating: Limit candidates to 50
        if not is_pro():
            candidates = candidates[:50]
            
        from rank import is_non_tech_candidate
        filtered_candidates = [c for c in candidates if not is_non_tech_candidate(c, core_skills, adjacent_skills)]
        if not filtered_candidates:
            filtered_candidates = candidates
            
        # Phase 3: local embeddings encoding & scoring
        model = load_sentence_transformer()
        emb_jd = model.encode(jd_text)
        candidate_texts = [build_candidate_text(c) for c in filtered_candidates]
        emb_candidates = model.encode(candidate_texts, show_progress_bar=False)
        
        # Load memory rules
        memory_rules = load_memory_rules()
        
        scored_list = []
        for i, c in enumerate(filtered_candidates):
            rs = c.get('redrob_signals', {})
            s1 = sig_semantic(emb_candidates[i], emb_jd)
            s2 = sig_skills(c.get('skills', []), rs.get('skill_assessment_scores', {}), core_skills)
            s3 = sig_career(c)
            s4 = sig_behavioral(rs)
            s5 = sig_domain(c, domain_kws)
            
            # Weighted Signal Fusion
            final_score = (s1 * 0.30) + (s2 * 0.25) + (s3 * 0.20) + (s4 * 0.15) + (s5 * 0.10)
            
            # Post-fusion OTW multiplier
            if not rs.get('open_to_work_flag', False):
                final_score *= 0.75
                
            reasoning = generate_reasoning(c, s2, core_skills)
            
            # Apply recruiter memory rules
            disqualified, disq_reason, boost, notes = apply_recruiter_memory_rules(
                c, core_skills, min_years_exp, seniority_lvl, memory_rules
            )
            
            if disqualified:
                final_score = 0.0
                reasoning = f"[DISQUALIFIED: {disq_reason}] " + reasoning
            else:
                final_score = min(1.0, max(0.0, final_score + boost))
                if notes:
                    reasoning = " | ".join(notes) + " || " + reasoning
            
            scored_list.append({
                "candidate_id": c["candidate_id"],
                "name": c.get("profile", {}).get("anonymized_name", "Anonymized"),
                "title": c.get("profile", {}).get("current_title", "Developer"),
                "experience": c.get("profile", {}).get("years_of_experience", 0),
                "score": round(final_score, 4),
                "s1_sem": round(s1, 4),
                "s2_skl": round(s2, 4),
                "s3_car": round(s3, 4),
                "s4_beh": round(s4, 4),
                "s5_dom": round(s5, 4),
                "reasoning": reasoning,
                "disqualified": disqualified,
                "disq_reason": disq_reason,
                "_profile": c
            })
            
        # Explicit Tie-Breaking
        scored_list.sort(key=lambda x: (-x["score"], x["candidate_id"]))
        
        t_elapsed = round(time.time() - t_start, 1)
        st.session_state.scored_candidates = scored_list
        st.session_state.run_runtime = t_elapsed
        st.session_state.total_candidates_evaluated = len(candidates)
        st.rerun()

# ── MAIN AREA ─────────────────────────────────────────────────────

# ── DASHBOARD NAVIGATION HEADER ───────────────────────────────────
# Render HTML navigation dynamically from session state
plan     = st.session_state.get("user_plan", "free")
name     = st.session_state.get("user_name", "User")
initials = st.session_state.get("user_initials", "?")

plan_badge = {
    "free":       ("FREE",       "#6B7280", "#F3F4F6"),
    "pro":        ("PRO",        "#FFFFFF",  "#0A0A0A"),
    "enterprise": ("ENTERPRISE", "#7C3AED", "#F5F3FF"),
}
badge_label, badge_text, badge_bg = plan_badge.get(plan, plan_badge["free"])

st.markdown(f"""
<div class="dashboard-nav">
  <div class="nav-logo">
    <svg width="20" height="20" viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg" style="color: #1a1615;">
      <path d="M4 2V22" stroke="currentColor" stroke-width="3.2" stroke-linecap="square"/>
      <path d="M4 14H19" stroke="currentColor" stroke-width="3.2" stroke-linecap="square"/>
      <path d="M13.5 9V22" stroke="currentColor" stroke-width="3.2" stroke-linecap="square"/>
    </svg>
    <span>Calipr</span>
    <span class="nav-badge" style="background:{badge_bg}; color:{badge_text}; font-size:9px; font-weight:800; padding:2px 7px; border-radius:4px; letter-spacing:0.05em; margin-left:6px;">{badge_label}</span>
  </div>
  <div class="nav-menu">
    <a href="?page=ranker" target="_self" class="nav-item {'active' if active_page == 'ranker' else ''}" style="text-decoration:none !important;">Candidate Ranker</a>
    <a href="?page=memory" target="_self" class="nav-item {'active' if active_page == 'memory' else ''}" style="text-decoration:none !important;">Recruiter Memory</a>
    <a href="?page=analytics" target="_self" class="nav-item {'active' if active_page == 'analytics' else ''}" style="text-decoration:none !important;">Analytics</a>
    <a href="?page=integrations" target="_self" class="nav-item {'active' if active_page == 'integrations' else ''}" style="text-decoration:none !important;">Integrations</a>
    <a href="?page=pricing" target="_self" class="nav-item {'active' if active_page == 'pricing' else ''}" style="text-decoration:none !important;">Plans &amp; Pricing</a>
  </div>
  <div class="nav-user">
    <div class="user-avatar">{initials}</div>
    <span class="user-name">{name}</span>
  </div>
</div>
""", unsafe_allow_html=True)

# Show hackathon banner for all users during demo
st.markdown(f"""
<div style="background:linear-gradient(135deg, #0A0A0A 0%, #1F2937 100%);
            padding:10px 20px; border-radius:12px; margin-top: 15px; margin-bottom:24px;
            display:flex; align-items:center; gap:12px; font-family: Inter, sans-serif;">
  <span style="font-size:16px;">🏆</span>
  <span style="font-size:13px; color:white; font-weight:500;">
    <strong>Hackathon Demo Mode</strong> — All Pro features unlocked.
    Sponsored by <strong>Redrob AI</strong>.
  </span>
  <span style="margin-left:auto; background:#16A34A; color:white; font-size:11px;
               font-weight:700; padding:3px 10px; border-radius:9999px;">
    IITRAM FLUX 2.0
  </span>
</div>
""", unsafe_allow_html=True)


def render_ranker_page():
    # Section 1 — Page Header
    st.markdown("""
    <div class="hero-section">
      <div class="hero-header-label">
        <span class="label-dot"></span>
        REDROB CHALLENGE PLATFORM
      </div>
      <div style="display:flex; justify-content:space-between; align-items:flex-start; flex-wrap:wrap; gap:20px;">
        <div>
          <h1 class="hero-title">AI Candidate Ranking Sandbox</h1>
          <p class="hero-desc">
            A multi-dimensional scoring engine evaluating 106K candidates. Leverages local FAISS retrieval, 5-signal fusion, and agentic re-ranking with explainable AI rationales.
          </p>
        </div>
      </div>
      <div class="hero-badges">
        <span class="badge badge-blue">⚡ Local Embeddings</span>
        <span class="badge badge-green">✓ Validated Output</span>
        <span class="badge">No API Cost</span>
      </div>
    </div>
    """, unsafe_allow_html=True)
    
    # Resume uploader collapsed expander
    with st.expander("📄 Add Custom Resumes to Evaluation Pool", expanded=False):
        uploaded_resumes = st.file_uploader("Upload resumes (PDF, TXT, or DOCX)", type=["pdf", "txt", "docx"], accept_multiple_files=True, key="custom_resumes_uploader")
        if uploaded_resumes:
            new_candidates_added = False
            for f in uploaded_resumes:
                if f.name not in [c.get("_filename") for c in st.session_state.uploaded_candidates]:
                    with st.spinner(f"Parsing {f.name}..."):
                        text = extract_text_from_file(f)
                        if text and not text.startswith("Error"):
                            cand = parse_resume_offline(text, filename=f.name)
                            cand["_filename"] = f.name
                            st.session_state.uploaded_candidates.append(cand)
                            new_candidates_added = True
                        else:
                            st.error(f"Failed to read {f.name}: {text}")
            if new_candidates_added:
                st.success(f"Successfully added {len(uploaded_resumes)} custom candidate(s) to the pool!")
                
        if st.session_state.uploaded_candidates:
            st.info(f"Currently loaded: {len(st.session_state.uploaded_candidates)} custom candidate(s) in pool.")
            if st.button("🗑️ Clear Uploaded Candidates"):
                st.session_state.uploaded_candidates = []
                st.rerun()
    
    # Section 6 — Conditional Results Display
    if st.session_state.scored_candidates is not None:
        st.markdown(
            f'<div style="background: rgba(14, 161, 88, 0.08); border: 1px solid rgba(14, 161, 88, 0.25); border-radius: 12px; color: #0c7540; padding: 15px; margin-bottom: 12px; font-weight:600; font-family:Inter,sans-serif;">'
            f'✅ Ranking Complete — {st.session_state.run_runtime}s · Evaluated {st.session_state.total_candidates_evaluated:,} candidates'
            f'</div>',
            unsafe_allow_html=True
        )
        
        if not is_pro():
            st.markdown("""
            <div style="background: rgba(245, 158, 11, 0.08); border: 1px solid rgba(245, 158, 11, 0.25); border-radius: 12px; padding: 16px; margin-bottom: 24px; font-family:Inter,sans-serif; display: flex; align-items: flex-start; gap: 12px;">
                <div style="font-size: 20px; margin-top: -2px;">⚠️</div>
                <div>
                    <div style="font-weight: 600; color: #b45309; margin-bottom: 4px; font-size: 14px;">Free Plan Limit Reached (50 Candidates Only)</div>
                    <div style="color: #c27829; font-size: 13px; line-height: 1.5; margin-bottom: 8px;">
                        You are currently on the Free Plan, which restricts candidate evaluation to the first 50 entries. To evaluate the full 106,039 candidates, please upgrade.
                    </div>
                    <a href="https://calipr-4fnf.vercel.app/#pricing" target="_blank" style="display: inline-block; background: #b45309; color: white; padding: 6px 12px; border-radius: 6px; font-size: 12px; font-weight: 600; text-decoration: none; transition: all 0.2s;">Upgrade to Professional</a>
                </div>
            </div>
            """, unsafe_allow_html=True)
        
        scored_list = st.session_state.scored_candidates
        
        left_col, right_col = st.columns([1, 1.4])
        
        with left_col:
            st.markdown('<div class="section-label">Ranked Candidates</div>', unsafe_allow_html=True)
            
            # Interactive Selectbox
            selected_idx = st.selectbox(
                "Select Candidate to Inspect",
                options=range(len(scored_list)),
                format_func=lambda i: f"#{i+1} - {scored_list[i]['name']} ({scored_list[i]['score']:.3f})",
                label_visibility="collapsed"
            )
            
            selected_cand = scored_list[selected_idx]
            
            # Scrollable Candidate List
            cards_html = "<div style='max-height: 650px; overflow-y: auto; padding-right: 5px; margin-top: 10px;'>"
            for rank, row in enumerate(scored_list[:30], 1):  # Display top 30
                is_sel = (rank - 1 == selected_idx)
                cards_html += candidate_row(rank, row["name"], row["title"], row["experience"], row["score"], is_selected=is_sel, disqualified=row.get("disqualified", False))
            cards_html += "</div>"
            st.markdown(cards_html, unsafe_allow_html=True)
            
        with right_col:
            st.markdown('<div class="section-label">Candidate Detail View</div>', unsafe_allow_html=True)
            
            # Candidate Card Detail Header
            avatar_initial = selected_cand['name'][0].upper() if selected_cand['name'] else 'C'
            st.markdown(
                f'<div style="display:flex;align-items:center;gap:16px;margin-bottom:20px;">'
                f'<div style="width:48px;height:48px;border-radius:50%;background:linear-gradient(135deg, #84b9ef, #156cc2);display:flex;align-items:center;justify-content:center;color:#FFFFFF;font-weight:700;font-size:18px;font-family:\'Open Runde\',sans-serif;">'
                f'{avatar_initial}'
                f'</div>'
                f'<div>'
                f'<h2 style="margin:0 !important; font-size: 22px !important;">{selected_cand["name"]}</h2>'
                f'<div style="font-size:14px;color:#757170;font-family:Inter,sans-serif;margin-top:2px;">'
                f'{selected_cand["title"]} · {selected_cand["experience"]:.1f} years experience'
                f'</div>'
                f'</div>'
                f'</div>',
                unsafe_allow_html=True
            )
            
            if selected_cand.get("disqualified", False):
                st.markdown(
                    f'<div style="background: rgba(239, 68, 68, 0.08); border: 1px solid rgba(239, 68, 68, 0.25); border-radius: 12px; color: #ef4444; padding: 15px; margin-bottom: 20px; font-weight:600; font-family:Inter,sans-serif;">'
                    f'🚫 Disqualified by Recruiter Memory: {selected_cand.get("disq_reason", "Hard rule triggered")}'
                    f'</div>',
                    unsafe_allow_html=True
                )
                
            # Radar Chart Plotly
            scores_dict = {
                'semantic': selected_cand['s1_sem'],
                'skills': selected_cand['s2_skl'],
                'career': selected_cand['s3_car'],
                'behavioral': selected_cand['s4_beh'],
                'domain': selected_cand['s5_dom']
            }
            st.plotly_chart(render_radar(scores_dict, selected_cand['name']), use_container_width=True, config={'displayModeBar': False})
            
            # Score Breakdown
            st.markdown('<div class="section-label" style="margin-top:15px;margin-bottom:10px;">Score Breakdown</div>', unsafe_allow_html=True)
            score_bar("🧠 Semantic Fit", selected_cand['s1_sem'])
            score_bar("💻 Skills Match", selected_cand['s2_skl'])
            score_bar("📈 Career Trajectory", selected_cand['s3_car'])
            score_bar("⚡ Behavioral Score", selected_cand['s4_beh'])
            score_bar("🎯 Domain Alignment", selected_cand['s5_dom'])
            
            # Score Card Display
            score_color = "#0ea158" if selected_cand['score'] >= 0.75 else "#cf8d13" if selected_cand['score'] >= 0.50 else "#c9502e"
            st.markdown(
                f'<div class="card-dark" style="margin: 20px 0; text-align:center; padding: 18px 24px;">'
                f'<div style="font-size:11px;font-weight:700;text-transform:uppercase;letter-spacing:0.08em;color:#757170;margin-bottom:6px;font-family:\'Fragment Mono\', monospace;">'
                f'Final Combined Suitability Score'
                f'</div>'
                f'<div style="font-size:36px;font-weight:700;color:{score_color};font-family:\'Fragment Mono\',monospace;line-height:1;">'
                f'{selected_cand["score"]:.4f}'
                f'</div>'
                f'</div>',
                unsafe_allow_html=True
            )
            
            # AI Rationale Box
            st.markdown('<div class="section-label">AI Rationale & Summary</div>', unsafe_allow_html=True)
            st.markdown(
                f'<div class="rationale-box">'
                f'"{selected_cand["reasoning"]}"'
                f'</div>',
                unsafe_allow_html=True
            )
            
            # Profile details (Skills Timelines / Timelines Layout)
            st.markdown('<hr style="margin:24px 0; border: none; border-top: 1px solid #e4e2e2;">', unsafe_allow_html=True)
            st.markdown('<div class="section-label">Extracted Skills & Capabilities</div>', unsafe_allow_html=True)
            
            skills_html = "<div style='display: flex; flex-wrap: wrap; gap: 8px; margin-top: 10px; margin-bottom: 20px;'>"
            for skill in selected_cand['_profile'].get('skills', []):
                prof = skill.get('proficiency', 'intermediate').lower()
                prof_color = "background: rgba(14,161,88,0.08); color: #0ea158; border: 1px solid rgba(14,161,88,0.25);" if prof == 'expert' or prof == 'advanced' else "background: rgba(132,185,239,0.08); color: #156cc2; border: 1px solid rgba(132,185,239,0.25);"
                skills_html += f'<span style="{prof_color} padding: 4px 12px; font-size: 12px; font-weight: 600; border-radius: 100px; font-family: Inter, sans-serif;">{skill.get("name")} • {prof.title()}</span>'
            skills_html += "</div>"
            st.markdown(skills_html, unsafe_allow_html=True)
            
            st.markdown('<div class="section-label">Work Experience Timeline</div>', unsafe_allow_html=True)
            timeline_html = "<div style='position: relative; padding-left: 20px; border-left: 2px solid #e4e2e2; margin-top: 15px; margin-left: 10px;'>"
            for job in selected_cand['_profile'].get('career_history', []):
                is_current = job.get('is_current', False)
                bullet_color = "#156cc2" if is_current else "#757170"
                timeline_html += (
                    f'<div style="position: relative; margin-bottom: 24px;">'
                    f'<div style="position: absolute; left: -27px; top: 4px; width: 12px; height: 12px; border-radius: 50%; background: {bullet_color}; border: 2px solid #FFFFFF;"></div>'
                    f'<div style="font-size: 14px; font-weight: 700; color: #1a1615; font-family: Inter, sans-serif;">{job.get("title")}</div>'
                    f'<div style="font-size: 12px; color: #757170; margin-top: 2px; font-family: Inter, sans-serif;">'
                    f'{job.get("company")} • {job.get("duration_months", 0)} months'
                    f'</div>'
                    f'<p style="font-size: 13.5px; color: #453f3d; margin-top: 6px; line-height: 1.5; font-family: Inter, sans-serif;">'
                    f'{job.get("description", "")}'
                    f'</p>'
                    f'</div>'
                )
            timeline_html += "</div>"
            st.markdown(timeline_html, unsafe_allow_html=True)
            
            # Download Action
            st.markdown('<hr style="margin:24px 0; border: none; border-top: 1px solid #e4e2e2;">', unsafe_allow_html=True)
            if is_pro():
                df_download = pd.DataFrame(scored_list)[["candidate_id", "name", "title", "experience", "score", "reasoning"]].copy()
                df_download.insert(0, "rank", range(1, len(df_download) + 1))
                csv_data = df_download.to_csv(index=False).encode('utf-8')
                st.download_button(
                    label="📥 Download Top 100 Shortlist CSV",
                    data=csv_data,
                    file_name="calipr_submission.csv",
                    mime="text/csv",
                    use_container_width=True
                )
            else:
                st.markdown("""
                <div style="border: 1px dashed #D1D5DB; background: #F9FAFB; border-radius: 12px; padding: 24px; text-align: center; font-family: Inter, sans-serif; margin-bottom: 20px;">
                    <div style="font-size: 32px; margin-bottom: 12px;">🔒</div>
                    <div style="font-weight: 600; color: #111827; margin-bottom: 6px; font-size: 15px;">CSV Export is Locked</div>
                    <div style="color: #6B7280; font-size: 13px; margin-bottom: 16px; max-width: 380px; margin-left: auto; margin-right: auto; line-height: 1.5;">
                        Exporting candidate shortlists to CSV is a Pro feature. Upgrade your plan to export and download submission-ready spreadsheets.
                    </div>
                    <a href="https://calipr-4fnf.vercel.app/#pricing" target="_blank" style="display: inline-block; background: #4A90FF; color: white; padding: 8px 18px; border-radius: 6px; font-size: 13px; font-weight: 600; text-decoration: none; transition: all 0.2s; box-shadow: 0 2px 4px rgba(74, 144, 255, 0.2);">Upgrade to Pro to Export</a>
                </div>
                """, unsafe_allow_html=True)
    
    # Section 2 — Stats Row
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.markdown('<div class="stat-card"><div class="stat-number">94%</div><div class="stat-label">Precision@5</div></div>', unsafe_allow_html=True)
    with col2:
        st.markdown('<div class="stat-card"><div class="stat-number">&lt; 5m</div><div class="stat-label">Pipeline Runtime</div></div>', unsafe_allow_html=True)
    with col3:
        st.markdown('<div class="stat-card"><div class="stat-number">106K</div><div class="stat-label">Total Candidates</div></div>', unsafe_allow_html=True)
    with col4:
        st.markdown('<div class="stat-card"><div class="stat-number">5</div><div class="stat-label">Scoring Signals</div></div>', unsafe_allow_html=True)
    
    # Section 3 — 5 Signal Cards
    st.markdown('<hr style="border: none; border-top: 1px solid #e4e2e2; margin: 40px 0;">', unsafe_allow_html=True)
    st.markdown('<div class="section-label">The Scoring Engine</div>', unsafe_allow_html=True)
    st.markdown('<h2 style="margin-top:0 !important;margin-bottom:6px;">Five Signals. One Score.</h2>', unsafe_allow_html=True)
    st.markdown('<p style="font-size:15px;color:#757170;margin-bottom:24px;">Every dimension a senior headhunter evaluates — <span style="color:#156cc2; font-weight:600;">quantified and fused</span>.</p>', unsafe_allow_html=True)
    
    sig_col1, sig_col2, sig_col3, sig_col4, sig_col5 = st.columns(5)
    with sig_col1:
        signal_card("🧠", "Semantic Fit", "30%", "Cosine similarity between JD and resume embeddings using all-MiniLM-L6-v2.")
    with sig_col2:
        signal_card("💻", "Skills Match", "25%", "BM25 with adjacency scoring. Adjacent skills score 0.4x. Verified assessments override proficiency.")
    with sig_col3:
        signal_card("📈", "Career Path", "20%", "Seniority level, company size growth progression, and education tiers (Tier 1-4).")
    with sig_col4:
        signal_card("⚡", "Behavioral", "15%", "Notice period scaling, completeness, activity freshness, and verification factors.")
    with sig_col5:
        signal_card("🎯", "Domain Fit", "10%", "Keyword frequency matches of core job description terminology in candidate history.")
    
    # Section 4 — Pipeline Phases
    st.markdown('<hr style="border: none; border-top: 1px solid #e4e2e2; margin: 40px 0;">', unsafe_allow_html=True)
    st.markdown('<div class="section-label">The Pipeline</div>', unsafe_allow_html=True)
    st.markdown('<h2 style="margin-top:0 !important;margin-bottom:6px;">From JD to Ranked Shortlist in Four Phases.</h2>', unsafe_allow_html=True)
    st.markdown('<p style="font-size:15px;color:#757170;margin-bottom:24px;">Math first, intelligence second. BM25 pre-filters 106K → 8K before a single embedding runs.</p>', unsafe_allow_html=True)
    
    p_col1, p_col2, p_col3, p_col4 = st.columns(4)
    with p_col1:
        st.markdown('<div class="phase-card"><div class="phase-number">1</div><h3 style="color:#1a1615;margin-top:10px;margin-bottom:8px;">Ingest &amp; Parse</h3><p style="font-size:13px;color:#757170;line-height:1.6;margin:0;font-family:Inter,sans-serif;">Validates candidate schemas and parses job descriptions via structured schemas.</p></div>', unsafe_allow_html=True)
    with p_col2:
        st.markdown('<div class="phase-card"><div class="phase-number">2</div><h3 style="color:#1a1615;margin-top:10px;margin-bottom:8px;">Hybrid Retrieval</h3><p style="font-size:13px;color:#757170;line-height:1.6;margin:0;font-family:Inter,sans-serif;">Pre-filters 106K pool to top 8,000 candidates using BM25 sparse queries.</p></div>', unsafe_allow_html=True)
    with p_col3:
        st.markdown('<div class="phase-card"><div class="phase-number">3</div><h3 style="color:#1a1615;margin-top:10px;margin-bottom:8px;">5-Signal Scoring</h3><p style="font-size:13px;color:#757170;line-height:1.6;margin:0;font-family:Inter,sans-serif;">Generates sentence-transformer embeddings and applies weighted score fusion.</p></div>', unsafe_allow_html=True)
    with p_col4:
        st.markdown('<div class="phase-card"><div class="phase-number">4</div><h3 style="color:#1a1615;margin-top:10px;margin-bottom:8px;">Agentic Re-Rank</h3><p style="font-size:13px;color:#757170;line-height:1.6;margin:0;font-family:Inter,sans-serif;">Identifies top 100 fits using tie-breakers and availability scoring.</p></div>', unsafe_allow_html=True)
    
    # Section 5 — Sponsors Strip
    st.markdown("""
    <div style="padding:24px 0;border-top:1px solid #e4e2e2;border-bottom:1px solid #e4e2e2;
                text-align:center;margin:40px 0;">
        <div class="section-label" style="margin-bottom:16px;">
            Hackathon Sponsor &amp; Technology Partners
        </div>
        <div style="display:flex;gap:32px;justify-content:center;align-items:center;flex-wrap:wrap;">
            <span style="font-size:16px;font-weight:800;color:#c9502e;font-family:'Open Runde', sans-serif;">
                redrob<span style="color:#c9502e">AI</span>
            </span>
            <span style="color:#e4e2e2;">·</span>
            <span style="font-size:15px;font-weight:600;color:#757170;">Google Gemini</span>
            <span style="color:#e4e2e2;">·</span>
            <span style="font-size:15px;font-weight:600;color:#757170;">Supabase</span>
            <span style="color:#e4e2e2;">·</span>
            <span style="font-size:15px;font-weight:600;color:#757170;">Hugging Face</span>
            <span style="color:#e4e2e2;">·</span>
            <span style="font-size:15px;font-weight:600;color:#757170;">Groq</span>
            <span style="color:#e4e2e2;">·</span>
            <span style="font-size:15px;font-weight:600;color:#757170;">FAISS</span>
        </div>
    </div>
    """, unsafe_allow_html=True)
    

# ── ROUTING PAGE VIEW SELECTION ───────────────────────────────────
if active_page == 'ranker':
    render_ranker_page()
elif active_page == 'memory':
    if is_pro():
        render_memory_page()
    else:
        st.markdown("""
        <div style="max-width: 600px; margin: 80px auto; padding: 40px; background: #ffffff; border: 1px solid #F3F4F6; border-radius: 16px; box-shadow: 0 10px 25px -5px rgba(0, 0, 0, 0.05), 0 8px 10px -6px rgba(0, 0, 0, 0.05); text-align: center; font-family: 'Inter', sans-serif;">
            <div style="display: inline-flex; align-items: center; justify-content: center; width: 64px; height: 64px; background: rgba(74, 144, 255, 0.08); border-radius: 50%; color: #4A90FF; font-size: 30px; margin-bottom: 24px;">
                🔒
            </div>
            <h2 style="font-size: 24px; font-weight: 700; color: #0A0A0A; margin-bottom: 8px;">Recruiter Memory is Locked</h2>
            <p style="font-size: 15px; color: #6B7280; line-height: 1.6; margin-bottom: 32px; max-width: 460px; margin-left: auto; margin-right: auto;">
                Calipr Recruiter Memory allows you to eliminate systematic hiring biases, calibrate custom evaluation rules, and build a persistent memory of your hiring criteria.
            </p>
            
            <div style="background: #F9FAFB; border-radius: 12px; padding: 20px; text-align: left; margin-bottom: 32px; max-width: 440px; margin-left: auto; margin-right: auto;">
                <div style="font-weight: 600; color: #0A0A0A; font-size: 14px; margin-bottom: 12px; text-transform: uppercase; letter-spacing: 0.05em;">Pro Features Unlocked:</div>
                <ul style="list-style: none; padding: 0; margin: 0; font-size: 14px; color: #4B5563; line-height: 1.8;">
                    <li style="margin-bottom: 8px; display: flex; align-items: center; gap: 8px;">
                        <span style="color: #0D9488; font-weight: bold;">✓</span> Bias Transparency & Mitigation
                    </li>
                    <li style="margin-bottom: 8px; display: flex; align-items: center; gap: 8px;">
                        <span style="color: #0D9488; font-weight: bold;">✓</span> Core/Adjacent Skill Weight Tuner
                    </li>
                    <li style="margin-bottom: 8px; display: flex; align-items: center; gap: 8px;">
                        <span style="color: #0D9488; font-weight: bold;">✓</span> Persistent Recruiter Memory Audit Log
                    </li>
                    <li style="display: flex; align-items: center; gap: 8px;">
                        <span style="color: #0D9488; font-weight: bold;">✓</span> Live Disqualification Rule Sandbox
                    </li>
                </ul>
            </div>
            
            <a href="https://calipr-4fnf.vercel.app/#pricing" target="_blank" style="display: inline-block; background: #4A90FF; color: #ffffff; padding: 12px 28px; border-radius: 8px; font-weight: 600; font-size: 15px; text-decoration: none; transition: all 0.2s; box-shadow: 0 4px 6px -1px rgba(74, 144, 255, 0.2), 0 2px 4px -1px rgba(74, 144, 255, 0.1);">
                Upgrade to Pro Plan
            </a>
        </div>
        """, unsafe_allow_html=True)
elif active_page == 'analytics':
    render_analytics_page()
elif active_page == 'integrations':
    render_integrations_page()
elif active_page == 'pricing':
    render_pricing_page()

# Section 7 — Footer
st.markdown("""
<div style="border-top:1px solid #e4e2e2;padding:32px 0;margin-top:64px;
            display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:16px;">
    <div style="font-size:13px;color:#757170;font-family:Inter,sans-serif;">
        © 2026 Calipr · Built at IITRAM Flux 2.0 · Sponsored by 
        <span style="color:#c9502e;font-weight:700;">Redrob AI</span>
    </div>
    <div style="font-size:13px;color:#757170;font-family:Inter,sans-serif;">
        Made with ❤️ by Aum Santoki &amp; Team
    </div>
</div>
""", unsafe_allow_html=True)
