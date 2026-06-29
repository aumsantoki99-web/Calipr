import streamlit as st
import json
from integrations.api import start_api_server

if "api_started" not in st.session_state:
    start_api_server(port=7861)
    st.session_state.api_started = True
import numpy as np
import time
from datetime import date
from sentence_transformers import SentenceTransformer
import re
from docx import Document
import pandas as pd
import uuid
import plotly.graph_objects as go

# Safe import of PdfReader
try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

# Page Configuration
st.set_page_config(
    page_title="Calipr AI — Redrob Ranker Sandbox",
    page_icon="calipr_logo.svg",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# JWT Authentication Check
if "auth_user_email" not in st.session_state:
    st.session_state.auth_user_email = None
    st.session_state.auth_user_name = None

if "token" in st.query_params:
    token = st.query_params["token"]
    try:
        import jwt
        import os
        
        # We bypass signature verification here because the token was already validated by the Supabase frontend,
        # and Supabase sometimes issues ES256 tokens which require a PEM public key instead of a simple string secret.
        decoded = jwt.decode(token, options={"verify_signature": False})
        st.session_state.auth_user_email = decoded.get("email")
        st.session_state.auth_user_name = decoded.get("user_metadata", {}).get("full_name", decoded.get("email"))
    except Exception as e:
        try:
            header = jwt.get_unverified_header(token)
        except Exception:
            header = "Could not parse header"
        st.error(f"Failed to authenticate token: {e}. Token Header: {header}")
        print(f"Failed to authenticate token: {e}. Header: {header}")
    st.query_params.clear()

if st.session_state.auth_user_name:
    pass # we will inject this in the dashboard-nav

# Hackathon demo mode: public Space is open with all features unlocked.
def is_pro():
    return True

# ── FULL CSS INJECTION ────────────────────────────────────────────
st.markdown("""
<style>
/* ── FONTS ── */
@import url('https://fonts.googleapis.com/css2?family=Inter:ital,wght@0,300;0,400;0,500;0,600;0,700;0,800;1,400;1,700&family=Fragment+Mono&display=swap');

@font-face {
  font-family: 'Open Runde';
  src: url('https://framerusercontent.com/assets/HJY4SY2JywrSZ2l1AEW9Tm9cO8.woff2') format('woff2');
  font-weight: 500;
  font-style: normal;
  font-display: swap;
}
@font-face {
  font-family: 'Open Runde';
  src: url('https://framerusercontent.com/assets/hG3wmGmFwadB6X5XPVXkMlmLr8o.woff2') format('woff2');
  font-weight: 600;
  font-style: normal;
  font-display: swap;
}
@font-face {
  font-family: 'Open Runde';
  src: url('https://framerusercontent.com/assets/3exmuO07FP19gMM08TQrpXl3BGQ.woff2') format('woff2');
  font-weight: 400;
  font-style: normal;
  font-display: swap;
}

/* ── GLOBAL RESET & THEME ── */
*, *::before, *::after { box-sizing: border-box; }

.stApp {
    background: linear-gradient(180deg, #fafafa 0%, #f9f8f8 36%, #f4f1ee 45%, #f4f1ee 51%, #e2ecf6 73%, #a7cbf2 125%) fixed !important;
    font-family: 'Inter', system-ui, -apple-system, sans-serif !important;
    color: #453f3d !important;
}

/* ── HIDE STREAMLIT CHROME ── */
#MainMenu, footer, header { visibility: hidden !important; pointer-events: none !important; }
.stDeployButton { display: none !important; }
[data-testid="stToolbar"] { display: none !important; }

/* ── MAIN CONTAINER ── */
.block-container {
    max-width: 1200px !important;
    padding: 0 24px 80px !important;
    margin: 0 auto !important;
}

/* ── SIDEBAR ── */
[data-testid="stSidebar"] {
    background: #f9f8f8 !important;
    border-right: 1px solid #e4e2e2 !important;
}
[data-testid="stSidebar"] .block-container {
    padding: 32px 20px !important;
}
[data-testid="stSidebarNav"] {
    display: none !important;
}

/* ── HEADINGS ── */
h1 {
    font-family: 'Open Runde', 'Inter', system-ui, sans-serif !important;
    font-size: 36px !important;
    font-weight: 600 !important;
    letter-spacing: -0.03em !important;
    color: #1a1615 !important;
    line-height: 1.2 !important;
}
h2 {
    font-family: 'Open Runde', 'Inter', system-ui, sans-serif !important;
    font-size: 26px !important;
    font-weight: 600 !important;
    letter-spacing: -0.02em !important;
    color: #1a1615 !important;
}
h3 {
    font-family: 'Open Runde', 'Inter', system-ui, sans-serif !important;
    font-size: 18px !important;
    font-weight: 600 !important;
    color: #1a1615 !important;
}

/* ── BUTTONS ── */
.stButton > button {
    background: #1a1615 !important;
    color: #ffffff !important;
    border: 1px solid #1a1615 !important;
    border-radius: 99px !important;
    font-family: 'Inter', sans-serif !important;
    font-weight: 600 !important;
    font-size: 14px !important;
    padding: 10px 24px !important;
    transition: all 0.2s ease !important;
    letter-spacing: -0.01em !important;
}
.stButton > button:hover {
    background: #2d2522 !important;
    border-color: #2d2522 !important;
    transform: translateY(-1px) !important;
    box-shadow: 0 4px 12px rgba(26, 22, 21, 0.12) !important;
}
.stButton > button:active {
    transform: translateY(0) !important;
}

/* Secondary Button style for downloads, etc */
.stDownloadButton > button {
    background: rgba(255, 255, 255, 0.8) !important;
    color: #1a1615 !important;
    border: 1px solid #e4e2e2 !important;
    border-radius: 99px !important;
    padding: 10px 24px !important;
    font-family: 'Inter', sans-serif !important;
    font-weight: 600 !important;
    font-size: 14px !important;
    transition: all 0.2s ease !important;
}
.stDownloadButton > button:hover {
    background: #ffffff !important;
    border-color: #84b9ef !important;
    color: #156cc2 !important;
    box-shadow: 0 4px 12px rgba(26, 22, 21, 0.05) !important;
}

/* ── TEXT INPUTS & TEXTAREAS ── */
.stTextArea textarea, .stTextInput input {
    background: rgba(255, 255, 255, 0.7) !important;
    border: 1px solid #e4e2e2 !important;
    border-radius: 12px !important;
    font-family: 'Inter', sans-serif !important;
    font-size: 14px !important;
    color: #1a1615 !important;
    padding: 12px 16px !important;
    transition: all 0.2s ease !important;
}
.stTextArea textarea:focus, .stTextInput input:focus {
    border-color: #84b9ef !important;
    box-shadow: 0 0 0 3px rgba(132, 185, 239, 0.2) !important;
    background: #ffffff !important;
    outline: none !important;
}

/* ── LABELS ── */
.stTextArea label, .stTextInput label, .stSelectbox label, .stRadio label, .stFileUploader label {
    font-family: 'Fragment Mono', monospace !important;
    font-size: 11px !important;
    font-weight: 600 !important;
    text-transform: uppercase !important;
    letter-spacing: 0.05em !important;
    color: #757170 !important;
    margin-bottom: 6px !important;
}

/* ── SELECT BOX ── */
div[data-baseweb="select"] {
    background: rgba(255, 255, 255, 0.8) !important;
    border: 1px solid #e4e2e2 !important;
    border-radius: 12px !important;
    font-family: 'Inter', sans-serif !important;
}
div[data-baseweb="select"] > div {
    background: transparent !important;
    border: none !important;
}

/* ── RADIO BUTTONS ── */
.stRadio [data-testid="stMarkdownContainer"] p {
    font-family: 'Inter', sans-serif !important;
    font-size: 13.5px !important;
    color: #453f3d !important;
}

/* ── METRICS ── */
[data-testid="stMetric"] {
    background: rgba(255, 255, 255, 0.7) !important;
    backdrop-filter: blur(10px) !important;
    border: 1px solid rgba(228, 226, 226, 0.8) !important;
    border-radius: 18px !important;
    padding: 16px 20px !important;
}
[data-testid="stMetricLabel"] {
    font-family: 'Fragment Mono', monospace !important;
    font-size: 10px !important;
    text-transform: uppercase !important;
    color: #757170 !important;
}
[data-testid="stMetricValue"] {
    font-size: 28px !important;
    font-weight: 600 !important;
    color: #1a1615 !important;
    font-family: 'Open Runde', sans-serif !important;
}

/* ── EXPANDER ── */
.streamlit-expander {
    background: rgba(255, 255, 255, 0.6) !important;
    backdrop-filter: blur(8px) !important;
    border: 1px solid rgba(228, 226, 226, 0.8) !important;
    border-radius: 16px !important;
    overflow: hidden !important;
    margin-bottom: 20px !important;
}
.streamlit-expander header {
    background: transparent !important;
    font-family: 'Inter', sans-serif !important;
    font-weight: 600 !important;
    font-size: 14px !important;
    color: #1a1615 !important;
}

/* ── SUCCESS / ERROR / INFO ── */
.stSuccess {
    background: rgba(14, 161, 88, 0.08) !important;
    border: 1px solid rgba(14, 161, 88, 0.25) !important;
    border-radius: 12px !important;
    color: #0c7540 !important;
}
.stError {
    background: rgba(201, 80, 46, 0.08) !important;
    border: 1px solid rgba(201, 80, 46, 0.25) !important;
    border-radius: 12px !important;
    color: #96351d !important;
}
.stInfo {
    background: rgba(21, 108, 194, 0.06) !important;
    border: 1px solid rgba(21, 108, 194, 0.2) !important;
    border-radius: 12px !important;
    color: #124d85 !important;
}

/* ── PROGRESS BAR ── */
.stProgress > div > div {
    background: linear-gradient(90deg, #84b9ef, #156cc2) !important;
    border-radius: 100px !important;
}
.stProgress > div {
    background: #e4e2e2 !important;
    border-radius: 100px !important;
    height: 8px !important;
}

/* ── CUSTOM COMPONENT CLASSES ── */

/* DASHBOARD NAV */
.dashboard-nav {
    display: flex;
    justify-content: space-between;
    align-items: center;
    background: rgba(255, 255, 255, 0.75) !important;
    backdrop-filter: blur(12px) !important;
    border: 1px solid rgba(228, 226, 226, 0.7) !important;
    border-radius: 20px !important;
    padding: 14px 24px !important;
    margin-top: 20px !important;
    margin-bottom: 30px !important;
    box-shadow: 0 4px 20px rgba(26, 22, 21, 0.03) !important;
}
.nav-logo {
    display: flex;
    align-items: center;
    gap: 8px;
    font-family: 'Open Runde', sans-serif !important;
    font-weight: 600 !important;
    font-size: 18px !important;
    color: #1a1615 !important;
}
.nav-badge {
    background: #e2ecf5;
    color: #156cc2;
    font-size: 10px;
    font-weight: 700;
    padding: 2px 6px;
    border-radius: 4px;
    letter-spacing: 0.5px;
}
.nav-menu {
    display: flex;
    gap: 24px;
}
.nav-item {
    font-size: 14px;
    font-weight: 500;
    color: #757170;
    cursor: pointer;
    transition: color 0.2s ease;
}
.nav-item:hover {
    color: #1a1615;
}
.nav-item.active {
    color: #156cc2;
    font-weight: 600;
    border-bottom: 2px solid #156cc2;
    padding-bottom: 2px;
}

/* GLASS CARD — light */
.card {
    background: rgba(255, 255, 255, 0.75) !important;
    backdrop-filter: blur(10px) !important;
    border: 1px solid rgba(228, 226, 226, 0.7) !important;
    border-radius: 20px !important;
    padding: 24px !important;
    transition: all 0.25s cubic-bezier(0.16, 1, 0.3, 1) !important;
    height: 200px !important;
}
.card:hover {
    transform: translateY(-4px) !important;
    border-color: #84b9ef !important;
    box-shadow: 0 12px 30px rgba(26, 22, 21, 0.06) !important;
}

/* GLASS CARD — dark */
.card-dark {
    background: #1a1615 !important;
    border: 1px solid rgba(255, 255, 255, 0.05) !important;
    border-radius: 20px !important;
    padding: 24px !important;
    color: #FFFFFF !important;
    box-shadow: 0 10px 30px rgba(26, 22, 21, 0.15) !important;
}

/* HERO SECTION */
.hero-section {
    background: rgba(255, 255, 255, 0.6) !important;
    backdrop-filter: blur(12px) !important;
    border: 1px solid rgba(228, 226, 226, 0.7) !important;
    border-radius: 24px !important;
    padding: 32px !important;
    margin-bottom: 30px !important;
    box-shadow: 0 4px 20px rgba(26, 22, 21, 0.02) !important;
}
.hero-header-label {
    display: inline-flex;
    align-items: center;
    gap: 6px;
    font-family: 'Fragment Mono', monospace !important;
    font-size: 11px !important;
    color: #156cc2 !important;
    background: #e2ecf5 !important;
    padding: 4px 12px !important;
    border-radius: 99px !important;
    font-weight: 600 !important;
    letter-spacing: 0.5px !important;
    margin-bottom: 16px !important;
}
.label-dot {
    width: 6px;
    height: 6px;
    background: #156cc2;
    border-radius: 50%;
}
.hero-title {
    font-size: 32px !important;
    font-weight: 700 !important;
    letter-spacing: -1px !important;
    line-height: 1.2 !important;
    margin: 0 !important;
    color: #1a1615 !important;
}
.hero-desc {
    font-size: 14px !important;
    color: #757170 !important;
    max-width: 650px !important;
    line-height: 1.6 !important;
    margin-top: 10px !important;
    margin-bottom: 0 !important;
}
.platform-stats {
    display: flex;
    gap: 20px;
    background: rgba(255, 255, 255, 0.8);
    border: 1px solid #e4e2e2;
    border-radius: 16px;
    padding: 12px 20px;
}
.p-stat {
    display: flex;
    flex-direction: column;
}
.p-stat-val {
    font-family: 'Fragment Mono', monospace;
    font-size: 18px;
    font-weight: 700;
    color: #1a1615;
}
.p-stat-lbl {
    font-size: 10px;
    font-weight: 500;
    color: #757170;
    text-transform: uppercase;
}
.hero-badges {
    display: flex;
    gap: 8px;
    margin-top: 20px;
    flex-wrap: wrap;
}

/* PILL BADGE */
.badge {
    display: inline-flex;
    align-items: center;
    gap: 6px;
    background: rgba(255, 255, 255, 0.6);
    border: 1px solid #e4e2e2;
    border-radius: 100px;
    padding: 6px 14px;
    font-size: 12px;
    font-weight: 600;
    color: #453f3d;
    font-family: 'Inter', sans-serif;
}
.badge-blue {
    background: #e2ecf5;
    border-color: rgba(132, 185, 239, 0.5);
    color: #156cc2;
}
.badge-green {
    background: rgba(14, 161, 88, 0.1);
    border-color: rgba(14, 161, 88, 0.3);
    color: #0ea158;
}

/* SCORE BAR */
.score-bar-container {
    margin-bottom: 14px;
}
.score-bar-label {
    display: flex;
    justify-content: space-between;
    font-size: 13px;
    font-weight: 500;
    color: #757170;
    font-family: 'Inter', sans-serif;
    margin-bottom: 6px;
}
.score-bar-label span {
    font-weight: 700;
    color: #1a1615;
    font-family: 'Fragment Mono', monospace;
}
.score-bar-track {
    height: 8px;
    background: #eddfd0;
    border-radius: 100px;
    overflow: hidden;
}
.score-bar-fill {
    height: 100%;
    border-radius: 100px;
}

/* CANDIDATE CARD */
.candidate-card {
    background: rgba(255, 255, 255, 0.65) !important;
    backdrop-filter: blur(12px) !important;
    -webkit-backdrop-filter: blur(12px) !important;
    border: 1px solid rgba(228, 226, 226, 0.8) !important;
    border-radius: 16px !important;
    padding: 12px 14px !important;
    cursor: pointer;
    transition: all 0.25s cubic-bezier(0.16, 1, 0.3, 1) !important;
    margin-bottom: 10px !important;
}
.candidate-card:hover {
    transform: translateY(-2px) !important;
    border-color: #84b9ef !important;
    background: rgba(255, 255, 255, 0.9) !important;
    box-shadow: 0 8px 24px rgba(26, 22, 21, 0.05) !important;
}
.candidate-card.selected {
    border-color: #156cc2 !important;
    background: rgba(210, 225, 245, 0.45) !important;
    box-shadow: 0 4px 16px rgba(21, 108, 194, 0.08) !important;
}

/* RANK BADGE */
.rank-badge {
    display: inline-flex;
    align-items: center;
    justify-content: center;
    width: 28px;
    height: 28px;
    background: #f4e6da;
    color: #754d29;
    border-radius: 8px;
    font-size: 12px;
    font-weight: 700;
    font-family: 'Inter', sans-serif;
}
.rank-badge.top3 {
    background: #e2ecf5;
    color: #156cc2;
}

/* SIGNAL WEIGHT BADGE */
.weight-badge {
    background: #e2ecf5 !important;
    color: #156cc2 !important;
    font-size: 11px !important;
    font-weight: 700 !important;
    padding: 4px 10px !important;
    border-radius: 99px !important;
}

/* SECTION LABEL */
.section-label {
    font-family: 'Fragment Mono', monospace !important;
    font-size: 11px !important;
    font-weight: 600 !important;
    text-transform: uppercase !important;
    letter-spacing: 0.08em !important;
    color: #757170 !important;
    margin-bottom: 12px !important;
}

/* PIPELINE PHASE CARD */
.phase-card {
    background: rgba(255, 255, 255, 0.75) !important;
    backdrop-filter: blur(10px) !important;
    border: 1px solid rgba(228, 226, 226, 0.7) !important;
    border-radius: 16px !important;
    padding: 24px 20px !important;
    text-align: center !important;
    color: #453f3d !important;
    position: relative !important;
    transition: all 0.2s ease !important;
    height: 160px !important;
}
.phase-card:hover {
    transform: translateY(-2px) !important;
    border-color: #84b9ef !important;
    box-shadow: 0 8px 24px rgba(26, 22, 21, 0.04) !important;
}
.phase-number {
    position: absolute !important;
    top: -12px !important;
    left: 50% !important;
    transform: translateX(-50%) !important;
    background: #156cc2 !important;
    color: #FFFFFF !important;
    width: 24px !important;
    height: 24px !important;
    border-radius: 50% !important;
    display: flex !important;
    align-items: center !important;
    justify-content: center !important;
    font-size: 11px !important;
    font-weight: 700 !important;
}

/* RATIONALE BOX */
.rationale-box {
    background: rgba(255, 255, 255, 0.8) !important;
    border: 1px solid #e4e2e2 !important;
    border-left: 4px solid #156cc2 !important;
    border-radius: 12px !important;
    padding: 16px 20px !important;
    font-size: 13.5px !important;
    color: #453f3d !important;
    line-height: 1.6 !important;
    font-style: italic !important;
}

/* STAT CARD */
.stat-card {
    background: rgba(255, 255, 255, 0.85) !important;
    backdrop-filter: blur(8px) !important;
    border: 1px solid #e4e2e2 !important;
    border-radius: 20px !important;
    padding: 24px !important;
    text-align: center !important;
    transition: transform 0.2s ease, box-shadow 0.2s ease !important;
    box-shadow: 0 4px 12px rgba(26, 22, 21, 0.02) !important;
}
.stat-card:hover {
    transform: translateY(-2px) !important;
    box-shadow: 0 8px 24px rgba(26, 22, 21, 0.05) !important;
}
.stat-number {
    font-size: 40px !important;
    font-weight: 700 !important;
    color: #1a1615 !important;
    font-family: 'Open Runde', sans-serif !important;
    letter-spacing: -1px !important;
    line-height: 1 !important;
}
.stat-label {
    font-size: 11px !important;
    font-weight: 600 !important;
    text-transform: uppercase !important;
    letter-spacing: 0.5px !important;
    color: #757170 !important;
    margin-top: 8px !important;
}

/* DIVIDER WITH LABEL */
.divider-label {
    display: flex;
    align-items: center;
    gap: 16px;
    margin: 32px 0;
}
.divider-label::before, .divider-label::after {
    content: '';
    flex: 1;
    height: 1px;
    background: #e4e2e2;
}
.divider-label span {
    font-family: 'Fragment Mono', monospace !important;
    font-size: 11px !important;
    font-weight: 600 !important;
    color: #757170 !important;
    text-transform: uppercase !important;
    letter-spacing: 0.08em !important;
    white-space: nowrap;
}
</style>
""", unsafe_allow_html=True)

# ── CONFIG & CONSTANTS ────────────────────────────────────────────
EMBEDDING_MODEL = "all-MiniLM-L6-v2"
LEVEL_MAP = {
    "intern":0.10,"trainee":0.12,"junior":0.20,"associate":0.28,
    "mid":0.40,"engineer":0.40,"developer":0.40,"analyst":0.35,
    "senior":0.70,"lead":0.82,"staff":0.88,"principal":0.93,
    "architect":0.90,"director":0.93,"manager":0.72,"head":0.85,
    "vp":0.95,"cto":1.0,"founder":0.88
}
SIZE_MAP = {"1-10":1,"11-50":2,"51-200":3,"201-500":4,
            "501-1000":5,"1001-5000":6,"5001-10000":7,"10001+":8}
SKILL_ADJACENCY = {
    "Python": ["Julia","R","Scala"],
    "PyTorch": ["TensorFlow","JAX","Keras","MXNet"],
    "React": ["Vue","Angular","Svelte","Next.js"],
    "FastAPI": ["Flask","Django","Express"],
    "PostgreSQL": ["MySQL","SQLite","MongoDB"],
    "Docker": ["Kubernetes","Podman"],
    "AWS": ["GCP","Azure","DigitalOcean"],
    "LangChain": ["LlamaIndex","Haystack","AutoGen"],
    "BERT": ["RoBERTa","DistilBERT","GPT-2","T5"],
    "YOLOv8": ["YOLOv5","Detectron2","EfficientDet"],
}

# Initialize session state variables
if "uploaded_candidates" not in st.session_state:
    st.session_state.uploaded_candidates = []
if "scored_candidates" not in st.session_state:
    st.session_state.scored_candidates = None
if "run_runtime" not in st.session_state:
    st.session_state.run_runtime = 0.0
if "total_candidates_evaluated" not in st.session_state:
    st.session_state.total_candidates_evaluated = 0

# ── CACHED MODEL LOADERS ──────────────────────────────────────────
@st.cache_resource
def load_sentence_transformer():
    return SentenceTransformer(EMBEDDING_MODEL)

@st.cache_data
def load_sample_candidates():
    with open("sample_candidates.json", "r", encoding="utf-8") as f:
        return json.load(f)

# ── PLOTLY RADAR CHART CONFIG ─────────────────────────────────────
def render_radar(scores: dict, candidate_name: str) -> go.Figure:
    categories = ['Semantic Fit', 'Skills Match', 'Career', 'Behavioral', 'Domain']
    values = [
        scores['semantic'],
        scores['skills'],
        scores['career'],
        scores['behavioral'],
        scores['domain']
    ]
    values_closed = values + [values[0]]
    cats_closed   = categories + [categories[0]]

    fig = go.Figure()

    # Filled area
    fig.add_trace(go.Scatterpolar(
        r=values_closed,
        theta=cats_closed,
        fill='toself',
        fillcolor='rgba(132, 185, 239, 0.2)',
        line=dict(color='#84b9ef', width=2.5),
        marker=dict(color='#156cc2', size=6, symbol='circle'),
        name=candidate_name,
        hovertemplate='%{theta}: %{r:.2f}<extra></extra>'
    ))

    fig.update_layout(
        margin=dict(l=60, r=60, t=60, b=60),
        polar=dict(
            bgcolor='rgba(0,0,0,0)',
            radialaxis=dict(
                visible=True,
                range=[0, 1],
                tickvals=[0.25, 0.50, 0.75, 1.0],
                ticktext=['0.25', '0.50', '0.75', '1.0'],
                tickfont=dict(size=9, color='#757170', family='Inter'),
                gridcolor='rgba(228, 226, 226, 0.6)',
                linecolor='rgba(228, 226, 226, 0.6)',
                tickcolor='rgba(0,0,0,0)',
            ),
            angularaxis=dict(
                tickfont=dict(size=11, color='#1a1615', family='Inter', weight=600),
                gridcolor='rgba(228, 226, 226, 0.5)',
                linecolor='rgba(228, 226, 226, 0.6)',
                rotation=90,
                direction='clockwise',
            )
        ),
        paper_bgcolor='rgba(0,0,0,0)',
        plot_bgcolor='rgba(0,0,0,0)',
        showlegend=False,
        font=dict(family='Inter, sans-serif', color='#1a1615'),
        height=340,
    )
    return fig

# ── SCORE BAR COMPONENT ───────────────────────────────────────────
def score_bar(label: str, value: float):
    if "Semantic" in label:
        fill_color = "#4A90FF"
    elif "Skills" in label:
        fill_color = "#00D4AA"
    elif "Career" in label:
        fill_color = "#7C6EFF"
    elif "Behavioral" in label:
        fill_color = "#F59E0B"
    elif "Domain" in label:
        fill_color = "#EF4444"
    else:
        fill_color = "#4A90FF" 
    st.markdown(f"""
    <div class="score-bar-container">
        <div class="score-bar-label">
            {label}
            <span>{value:.2f}</span>
        </div>
        <div class="score-bar-track">
            <div class="score-bar-fill" style="width:{value*100:.1f}%;background:{fill_color};"></div>
        </div>
    </div>
    """, unsafe_allow_html=True)

# ── CANDIDATE ROW COMPONENT ───────────────────────────────────────
def candidate_row(rank: int, name: str, title: str, 
                  years: float, score: float, is_selected: bool = False, cand_idx: int = 0):
    selected_class = "selected" if is_selected else ""
    rank_class = "top3" if rank <= 3 else ""
    score_color = "#0ea158" if score >= 0.75 else "#cf8d13" if score >= 0.50 else "#4A90FF"
    
    # Check shortlist/reject decision from session state
    decisions = st.session_state.get("candidate_decisions", {})
    decision = decisions.get(name, None)
    
    decision_style = ""
    badge_html = ""
    if decision == "shortlisted":
        decision_style = "border-left: 3px solid #0ea158 !important; background: rgba(14,161,88,0.03);"
        badge_html = '<span style="font-size: 12px; color: #0ea158; margin-left: 4px; font-weight: bold;">✓</span>'
    elif decision == "rejected":
        decision_style = "opacity: 0.55; filter: grayscale(50%);"
        badge_html = '<span style="font-size: 12px; color: #dc2626; margin-left: 4px; font-weight: bold;">✗</span>'
        
    return f"""<div class="candidate-card {selected_class}" style="{decision_style}" data-cand-idx="{cand_idx}">
<div style="display:flex;align-items:center;gap:10px;">
<div class="rank-badge {rank_class}">#{rank}</div>
<div style="flex:1;min-width:0;">
<div style="font-size:13px;font-weight:700;color:#1a1615;font-family:Inter,sans-serif;white-space:nowrap;overflow:hidden;text-overflow:ellipsis;display:flex;align-items:center;gap:4px;">{name} {badge_html}</div>
<div style="font-size:11px;color:#757170;font-family:Inter,sans-serif;margin-top:2px;white-space:nowrap;overflow:hidden;text-overflow:ellipsis;">{title}</div>
</div>
<div style="font-size:16px;font-weight:800;color:{score_color};font-family:'Fragment Mono',monospace;">{score:.3f}</div>
</div>
</div>"""

def signal_card(icon: str, name: str, weight: str, description: str):
    icon_html = f'<div style="font-size:28px;">{icon}</div>' if icon else ''
    justify = "space-between" if icon else "flex-end"
    st.markdown(f"""<div class="card" style="height:200px !important;">
<div style="display:flex;justify-content:{justify};align-items:flex-start;margin-bottom:12px;">
{icon_html}
<div class="weight-badge">{weight}</div>
</div>
<div style="font-size:15px;font-weight:700;color:#1a1615;font-family:'Open Runde', 'Inter',sans-serif;margin-bottom:6px;">{name}</div>
<div style="font-size:13px;color:#757170;line-height:1.6;font-family:Inter,sans-serif;">{description}</div>
</div>""", unsafe_allow_html=True)

# ── SCORING LOGIC FUNCTIONS ───────────────────────────────────────
def build_candidate_text(c):
    if "_raw_resume_text" in c:
        return c["_raw_resume_text"]
    p = c.get('profile', {})
    skills_text = " ".join([s.get('name','') for s in c.get('skills', [])])
    career_text = " ".join([jh.get('description','') for jh in c.get('career_history', [])])
    titles_text = " ".join([jh.get('title','') for jh in c.get('career_history', [])])
    return f"{p.get('summary','')} {p.get('headline','')} {p.get('current_title','')} {skills_text} {career_text} {titles_text}"

def tokenize(text):
    STOP = {"a","an","the","and","or","in","on","at","to","for","of","with","is","are","was","were","i","we","you"}
    return [t for t in re.findall(r'\b[a-z0-9][a-z0-9+#\.]*\b', text.lower()) if t not in STOP and len(t) > 1]

def sig_semantic(emb_candidate, emb_jd):
    dot = np.dot(emb_candidate, emb_jd)
    norms = np.linalg.norm(emb_candidate) * np.linalg.norm(emb_jd)
    return float(dot / norms) if norms > 0 else 0.0

def sig_skills(candidate_skills, assessment_scores, core_skills):
    if not core_skills:
        return 0.5
    PROF = {'beginner':0.4,'intermediate':0.6,'advanced':0.85,'expert':1.0}
    cand_map = {s.get('name','').lower(): s for s in candidate_skills}
    score = 0.0
    for jd_skill in core_skills:
        jl = jd_skill.lower()
        if jl in cand_map:
            s = cand_map[jl]
            asmnt_val = assessment_scores.get(jd_skill, 0)
            if asmnt_val >= 70:
                base = 1.0
            else:
                base = PROF.get(s.get('proficiency','intermediate'), 0.6)
            dur  = min(s.get('duration_months',0)/24, 1.0) * 0.15
            asmnt = (asmnt_val / 100) * 0.10
            score += min(base + dur + asmnt, 1.0)
        else:
            adj_list = SKILL_ADJACENCY.get(jd_skill, [])
            if any(a.lower() in cand_map for a in adj_list):
                score += 0.40
    return min(score / max(len(core_skills), 1), 1.0)

def sig_career(c):
    p = c.get('profile', {})
    career = c.get('career_history', [])
    edu = c.get('education', [])
    title = p.get('current_title','').lower()
    seniority = next((v for k,v in LEVEL_MAP.items() if k in title), 0.35)
    sizes = [SIZE_MAP.get(jh.get('company_size','1-10'), 1) for jh in career]
    prog = max((sizes[-1]-sizes[0])/7, 0.0) if len(sizes) > 1 else 0.0
    tier_bonus = {'tier_1':0.15,'tier_2':0.10,'tier_3':0.05,'tier_4':0.0,'unknown':0.02}
    best_tier = max((tier_bonus.get(e.get('tier','unknown'),0.02) for e in edu), default=0.02)
    score = min(seniority*0.50 + prog*0.30 + best_tier*0.20, 1.0)
    
    # Consulting company penalty
    curr_company = p.get('current_company', '').lower()
    consulting_firms = ["tcs", "tata consultancy services", "infosys", "wipro", "cognizant",
                        "accenture", "capgemini", "tech mahindra", "hcl", "hcltech", "l&t", "lnt", "mindtree"]
    if any(comp in curr_company for comp in consulting_firms):
        score *= 0.85
    return score

def sig_behavioral(rs):
    try:
        last_active = date.fromisoformat(rs.get('last_active_date', '').split('T')[0])
        days_ago = (date.today() - last_active).days
    except Exception:
        days_ago = 30
    freshness = max(0.0, 1.0 - days_ago/90)
    completeness = rs.get('profile_completeness_score', 80)/100
    
    response_rate = rs.get('recruiter_response_rate', 0.5)
    resp_time  = max(0, 1 - rs.get('avg_response_time_hours', 24)/72)
    interview  = rs.get('interview_completion_rate', 0.5)
    engagement = response_rate*0.4 + resp_time*0.3 + interview*0.3
    
    gh = rs.get('github_activity_score', -1)
    github = 0.3 if gh == -1 else gh/100
    
    offer = rs.get('offer_acceptance_rate', -1)
    offer_n = 0.5 if offer == -1 else max(offer, 0)
    
    notice = rs.get('notice_period_days', 30)
    if notice is None:
        notice = 30
    try:
        notice = float(notice)
    except Exception:
        notice = 30
    notice_score = max(0.0, 1.0 - (notice / 180))
    
    otw = 1.0 if rs.get('open_to_work_flag', False) else 0.3
    
    verified = (int(rs.get('verified_email', False)) + int(rs.get('verified_phone', False)) + int(rs.get('linkedin_connected', False)))/3
    
    relocate = rs.get('willing_to_relocate', False)
    if isinstance(relocate, str):
        relocate = relocate.strip().lower() == "true"
    work_mode = str(rs.get('preferred_work_mode', '')).lower()
    
    bonus = 0.0
    if relocate or "remote" in work_mode or "hybrid" in work_mode:
        bonus = 0.05
        
    score = (
        completeness * 0.18 +
        freshness * 0.12 +
        engagement * 0.25 +
        github * 0.15 +
        offer_n * 0.10 +
        notice_score * 0.10 +
        otw * 0.05 +
        verified * 0.05
    ) + bonus
    return min(score, 1.0)

def sig_domain(c, domain_kws):
    if not domain_kws:
        return 0.5
    p = c.get('profile', {})
    industries = [p.get('current_industry','')] + [jh.get('industry','') for jh in c.get('career_history',[])]
    text = (p.get('summary','') + ' ' + p.get('headline','') + ' ' + ' '.join(industries)).lower()
    hits = sum(1 for kw in domain_kws if kw.lower() in text)
    return min(hits / max(len(domain_kws), 1), 1.0)

def generate_reasoning(c, s2_skills, core_skills):
    p = c.get('profile', {})
    rs = c.get('redrob_signals', {})
    
    current_title = p.get('current_title', 'Software Engineer')
    if not current_title:
        current_title = 'Software Engineer'
        
    current_title = "".join(ch for ch in current_title if 32 <= ord(ch) <= 126)
    if len(current_title) > 40:
        current_title = current_title[:37] + "..."
        
    try:
        years_experience = int(float(p.get('years_of_experience', 0)))
    except Exception:
        years_experience = 0
    
    candidate_skills = {s.get('name', '').lower().strip() for s in c.get('skills', [])}
    matched_core_skills = 0
    for jd_skill in core_skills:
        jl = jd_skill.lower().strip()
        if any(jl == cs or jl in cs or cs in jl for cs in candidate_skills):
            matched_core_skills += 1
            
    recruiter_response_rate = rs.get('recruiter_response_rate', 0.0)
    if recruiter_response_rate is None:
        recruiter_response_rate = 0.0
    try:
        recruiter_response_rate = float(recruiter_response_rate)
    except Exception:
        recruiter_response_rate = 0.0
        
    return f"{current_title} with {years_experience} yrs; {matched_core_skills} core skills matched; response rate {recruiter_response_rate:.2f}."

def extract_text_from_file(uploaded_file):
    filename = uploaded_file.name
    if filename.endswith(".pdf"):
        if PdfReader is None:
            return "Error: pypdf library is not installed."
        try:
            reader = PdfReader(uploaded_file)
            text = ""
            for page in reader.pages:
                text += page.extract_text() or ""
            return text
        except Exception as e:
            return f"Error reading PDF: {e}"
    elif filename.endswith(".docx"):
        try:
            doc = Document(uploaded_file)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except Exception as e:
            return f"Error reading DOCX: {e}"
    else:
        try:
            return uploaded_file.read().decode("utf-8", errors="ignore")
        except Exception as e:
            return f"Error reading text file: {e}"

def parse_resume_offline(text, filename="Resume"):
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    name = lines[0] if lines else filename.split('.')[0]
    if len(name) > 30:
        name = filename.split('.')[0]
        
    exp_matches = re.findall(r'(\d+(?:\.\d+)?)\s*(?:years|yrs|year)\b', text.lower())
    years_exp = float(exp_matches[0]) if exp_matches else 3.0
    if years_exp > 40:
        years_exp = 3.0
        
    known_skills = [
        "Python", "PyTorch", "React", "FastAPI", "PostgreSQL", "Docker", "AWS", "LangChain", "BERT", "YOLOv8",
        "JavaScript", "TypeScript", "HTML", "CSS", "SQL", "Git", "Spark", "Kafka", "TensorFlow", "Kubernetes",
        "C++", "Java", "Go", "Rust", "Node.js", "MongoDB", "Redis", "Framer", "Figma", "Tailwind"
    ]
    detected_skills = []
    for skill in known_skills:
        pattern = r'\b' + re.escape(skill.lower()) + r'\b'
        if re.search(pattern, text.lower()):
            detected_skills.append({
                "name": skill,
                "proficiency": "advanced" if skill in ["Python", "JavaScript", "SQL"] else "intermediate",
                "duration_months": int(years_exp * 6)
            })
            
    title = "Software Engineer"
    titles = ["backend engineer", "frontend engineer", "fullstack engineer", "full stack engineer",
              "data scientist", "data engineer", "machine learning engineer", "devops engineer",
              "software engineer", "product manager", "project manager", "ui/ux designer"]
    for t in titles:
        if t in text.lower():
            title = t.title()
            break
            
    cid = f"UPLOAD_{uuid.uuid4().hex[:7].upper()}"
    
    candidate = {
        "candidate_id": cid,
        "profile": {
            "anonymized_name": name,
            "headline": f"{title} | {', '.join([s['name'] for s in detected_skills[:3]])}",
            "summary": text[:300] + ("..." if len(text) > 300 else ""),
            "location": "Remote",
            "country": "Global",
            "years_of_experience": years_exp,
            "current_title": title,
            "current_company": "Independent Consultant",
            "current_company_size": "1-10",
            "current_industry": "Tech"
        },
        "career_history": [
            {
                "company": "Current Company",
                "title": title,
                "duration_months": int(years_exp * 12),
                "is_current": True,
                "company_size": "1-10",
                "description": f"Worked as {title} utilizing skills like {', '.join([s['name'] for s in detected_skills[:5]])}."
            }
        ],
        "education": [
            {
                "institution": "University",
                "degree": "Bachelor of Science",
                "field_of_study": "Computer Science",
                "start_year": 2018,
                "end_year": 2022,
                "tier": "tier_2"
            }
        ],
        "skills": detected_skills,
        "redrob_signals": {
            "skill_assessment_scores": {s["name"]: 85 for s in detected_skills},
            "profile_completeness_score": 90,
            "recruiter_response_rate": 0.85,
            "avg_response_time_hours": 12.0,
            "interview_completion_rate": 0.90,
            "github_activity_score": 80,
            "offer_acceptance_rate": 0.80,
            "open_to_work_flag": True,
            "verified_email": True,
            "verified_phone": True,
            "linkedin_connected": True,
            "saved_by_recruiters_30d": 3,
            "last_active_date": date.today().isoformat()
        },
        "_raw_resume_text": text
    }
    return candidate

# Show hackathon banner for all users during demo
st.markdown(f"""
<div style="background:linear-gradient(135deg, #0A0A0A 0%, #1F2937 100%);
            padding:10px 20px; border-radius:10px; margin-bottom:20px;
            display:flex; align-items:center; gap:12px;">
  <span style="font-size:16px;">🏆</span>
  <span style="font-size:13px; color:white; font-weight:500;">
    <strong>Hackathon Demo Mode</strong> — All Pro features unlocked.
    Sponsored by <strong>Redrob AI</strong>.
  </span>
  <span style="margin-left:auto; background:#16A34A; color:white; font-size:11px;
               font-weight:700; padding:3px 10px; border-radius:9999px;">
    IITRAM FLUX 2.0
  </span>
</div>
""", unsafe_allow_html=True)

st.markdown("""
<style>
/* Hide the radio button circles to make it look like a nav bar */
.main div[role="radiogroup"] > label > div:first-child {
    display: none;
}
.main div[data-testid="stRadio"] {
    margin-top: -65px; /* Pull it up onto the nav bar */
    margin-bottom: -20px;
    position: relative;
    z-index: 10;
}
.main div[role="radiogroup"] {
    justify-content: center;
    gap: 32px;
}
.main div[role="radiogroup"] > label {
    cursor: pointer;
    font-family: 'Inter', sans-serif;
    font-weight: 500;
    font-size: 14.5px;
    color: #757170;
    padding-bottom: 8px;
    text-transform: none !important;
}
/* Style the selected radio button text */
.main div[role="radiogroup"] > label[data-checked="true"] {
    color: #1a1615;
    font-weight: 600;
    border-bottom: 2px solid #1a1615;
}
</style>
""", unsafe_allow_html=True)

first_name = st.session_state.auth_user_name.split()[0] if st.session_state.auth_user_name else "Guest"
initial = first_name[0].upper() if first_name else "G"
user_badge_html = f"""
  <div style="display: flex; align-items: center; gap: 8px; font-size: 13px; font-weight: 600; color: #453f3d;">
    <div style="background: linear-gradient(135deg, #4A90FF 0%, #1a1615 100%); color: white; width: 26px; height: 26px; border-radius: 50%; display: flex; align-items: center; justify-content: center; font-size: 13px;">{initial}</div>
    {first_name}
  </div>
"""

# Render HTML navigation
st.markdown(f"""
<div class="dashboard-nav">
  <div class="nav-logo">
    <svg width="20" height="20" viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg" style="color: #1a1615;">
      <path d="M4 2V22" stroke="currentColor" stroke-width="3.2" stroke-linecap="square"/>
      <path d="M4 14H19" stroke="currentColor" stroke-width="3.2" stroke-linecap="square"/>
      <path d="M13.5 9V22" stroke="currentColor" stroke-width="3.2" stroke-linecap="square"/>
    </svg>
    <span>Calipr</span>
  </div>
  {user_badge_html}
</div>
""", unsafe_allow_html=True)

# The functional navigation bar pulled UP into the HTML via negative margin
selected_page = st.radio("Navigation", ["Candidate Ranker", "Recruiter Memory", "Analytics", "Integrations"], horizontal=True, label_visibility="collapsed")

# ── ROUTING LOGIC ──
elif selected_page == "Integrations":
    st.markdown("""
    <style>
    [data-testid="stSidebar"] { display: none !important; }
    [data-testid="stSidebarCollapseButton"] { display: none !important; }
    </style>
    """, unsafe_allow_html=True)
    from integrations_ui import integrations_page
    integrations_page()
    st.stop()
elif selected_page == "Recruiter Memory":
    st.markdown("""
    <style>
    [data-testid="stSidebar"] { display: none !important; }
    [data-testid="stSidebarCollapseButton"] { display: none !important; }
    </style>
    """, unsafe_allow_html=True)
    if is_pro():
        from pages.recruiter_memory_page import render_recruiter_memory_page
        render_recruiter_memory_page()
    else:
        st.markdown("""
        <div style="max-width: 600px; margin: 80px auto; padding: 40px; background: #ffffff; border: 1px solid #F3F4F6; border-radius: 16px; box-shadow: 0 10px 25px -5px rgba(0, 0, 0, 0.05), 0 8px 10px -6px rgba(0, 0, 0, 0.05); text-align: center; font-family: 'Inter', sans-serif;">
            <div style="display: inline-flex; align-items: center; justify-content: center; width: 64px; height: 64px; background: rgba(74, 144, 255, 0.08); border-radius: 50%; color: #4A90FF; font-size: 30px; margin-bottom: 24px;">
                🔒
            </div>
            <h2 style="font-size: 24px; font-weight: 700; color: #0A0A0A; margin-bottom: 8px;">Recruiter Memory is Locked</h2>
            <p style="font-size: 15px; color: #6B7280; line-height: 1.6; margin-bottom: 32px; max-width: 460px; margin-left: auto; margin-right: auto;">
                Calipr Recruiter Memory allows you to eliminate systematic hiring biases, calibrate custom evaluation rules, and build a persistent memory of your hiring criteria.
            </p>
            
            <div style="background: #F9FAFB; border-radius: 12px; padding: 20px; text-align: left; margin-bottom: 32px; max-width: 440px; margin-left: auto; margin-right: auto;">
                <div style="font-weight: 600; color: #0A0A0A; font-size: 14px; margin-bottom: 12px; text-transform: uppercase; letter-spacing: 0.05em;">Pro Features Unlocked:</div>
                <ul style="list-style: none; padding: 0; margin: 0; font-size: 14px; color: #4B5563; line-height: 1.8;">
                    <li style="margin-bottom: 8px; display: flex; align-items: center; gap: 8px;">
                        <span style="color: #0D9488; font-weight: bold;">✓</span> Bias Transparency & Mitigation
                    </li>
                    <li style="margin-bottom: 8px; display: flex; align-items: center; gap: 8px;">
                        <span style="color: #0D9488; font-weight: bold;">✓</span> Core/Adjacent Skill Weight Tuner
                    </li>
                    <li style="margin-bottom: 8px; display: flex; align-items: center; gap: 8px;">
                        <span style="color: #0D9488; font-weight: bold;">✓</span> Persistent Recruiter Memory Audit Log
                    </li>
                    <li style="display: flex; align-items: center; gap: 8px;">
                        <span style="color: #0D9488; font-weight: bold;">✓</span> Live Disqualification Rule Sandbox
                    </li>
                </ul>
            </div>
            
            <a href="https://calipr-4fnf.vercel.app/#pricing" target="_blank" style="display: inline-block; background: #4A90FF; color: #ffffff; padding: 12px 28px; border-radius: 8px; font-weight: 600; font-size: 15px; text-decoration: none; transition: all 0.2s; box-shadow: 0 4px 6px -1px rgba(74, 144, 255, 0.2), 0 2px 4px -1px rgba(74, 144, 255, 0.1);">
                Upgrade to Pro Plan
            </a>
        </div>
        """, unsafe_allow_html=True)
    st.stop()
elif selected_page == "Analytics":
    st.markdown("""
    <style>
    [data-testid="stSidebar"] { display: none !important; }
    [data-testid="stSidebarCollapseButton"] { display: none !important; }
    </style>
    """, unsafe_allow_html=True)
    from pages.analytics_page import render_analytics_page
    render_analytics_page()
    st.stop()

# ── SIDEBAR INTERFACE ─────────────────────────────────────────────
# ── SIDEBAR INTERFACE ─────────────────────────────────────────────
st.sidebar.markdown("""
<div style="padding:10px 0 20px; display:flex; align-items:center; gap:10px;">
    <svg width="28" height="28" viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
      <path d="M4 2V22" stroke="#1a1615" stroke-width="3.2" stroke-linecap="square"/>
      <path d="M4 14H19" stroke="#1a1615" stroke-width="3.2" stroke-linecap="square"/>
      <path d="M13.5 9V22" stroke="#1a1615" stroke-width="3.2" stroke-linecap="square"/>
    </svg>
    <span style="font-size:28px;font-weight:800;color:#1a1615;font-family:'Open Runde', sans-serif;letter-spacing:-0.03em;">Calipr</span>
</div>
<hr style="margin:8px 0 20px;border-top:1px solid #e4e2e2;">
""", unsafe_allow_html=True)

st.sidebar.markdown('<div class="section-label">Job Description</div>', unsafe_allow_html=True)
jd_input_method = st.sidebar.radio("Choose input method", ["Use Hackathon JD", "Paste custom JD", "Upload Document"], label_visibility="collapsed")

jd_text = ""
if jd_input_method == "Use Hackathon JD":
    try:
        doc = Document("job_description.docx")
        jd_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    except Exception:
        jd_text = "Senior Backend Engineer with experience in hybrid search, vector databases, Python, and ranking algorithms."
    st.sidebar.markdown("""
    <div style="background:rgba(255,255,255,0.8);border:1px solid #e4e2e2;border-radius:10px;padding:12px;
                font-size:12px;color:#757170;max-height:150px;overflow-y:auto;font-family:Inter,sans-serif;line-height:1.5;margin-bottom:15px;">
        <strong>Default Job Description loaded:</strong><br>
        Senior AI Engineer founding team. Deployed embeddings, retrieval, ranking, vector databases (FAISS, OpenSearch), evaluation frameworks (NDCG).
    </div>
    """, unsafe_allow_html=True)
elif jd_input_method == "Paste custom JD":
    jd_text = st.sidebar.text_area("Paste JD text here", height=200, placeholder="Enter job description text...")
else:
    uploaded_jd = st.sidebar.file_uploader("Upload job description", type=["docx", "pdf", "txt"], label_visibility="collapsed")
    if uploaded_jd:
        try:
            fname = uploaded_jd.name.lower()
            if fname.endswith(".pdf"):
                if PdfReader is None:
                    st.sidebar.error("Error: pypdf library is not installed.")
                else:
                    reader = PdfReader(uploaded_jd)
                    jd_text = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
                    st.sidebar.success("PDF parsed successfully.")
            elif fname.endswith(".txt"):
                jd_text = uploaded_jd.read().decode("utf-8")
                st.sidebar.success("TXT parsed successfully.")
            else:
                doc = Document(uploaded_jd)
                jd_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
                st.sidebar.success("DOCX parsed successfully.")
        except Exception as e:
            st.sidebar.error(f"Error parsing document: {e}")

run_pipeline = st.sidebar.button("Rank Candidates", type="primary", use_container_width=True)

from analytics.data_store import load_run_history
history = load_run_history()
sessions = len(history)

if sessions > 0:
    w_sem = 30
    w_skl = 25
    w_car = 14
    w_beh = 18
    w_dom = 13
else:
    w_sem = 30
    w_skl = 25
    w_car = 20
    w_beh = 15
    w_dom = 10

st.sidebar.markdown('<hr style="margin:20px 0 16px;border-top:1px solid #e4e2e2;">', unsafe_allow_html=True)
st.sidebar.markdown('<div class="section-label">Pipeline Weights</div>', unsafe_allow_html=True)
st.sidebar.markdown(f"""
<div style="font-size:13px;font-family:Inter,sans-serif;color:#757170;line-height:1.9;">
  <div style="display:flex;justify-content:space-between;"><span>Semantic Fit</span><span style="font-weight:700;color:#1a1615;">{w_sem}%</span></div>
  <div style="display:flex;justify-content:space-between;"><span>Skills Match</span><span style="font-weight:700;color:#1a1615;">{w_skl}%</span></div>
  <div style="display:flex;justify-content:space-between;"><span>Career Trajectory</span><span style="font-weight:700;color:#1a1615;">{w_car}%</span></div>
  <div style="display:flex;justify-content:space-between;"><span>Behavioral Signals</span><span style="font-weight:700;color:#1a1615;">{w_beh}%</span></div>
  <div style="display:flex;justify-content:space-between;"><span>Domain Alignment</span><span style="font-weight:700;color:#1a1615;">{w_dom}%</span></div>
</div>
""", unsafe_allow_html=True)

st.sidebar.markdown("""
<div style="font-size:11px;color:#757170;font-family:Inter,sans-serif;line-height:1.5;margin-top:15px;">
  Built for Redrob Hackathon<br>
  Sponsored by Redrob AI
</div>
""", unsafe_allow_html=True)

# ── RUN PIPELINE CALCULATION ──────────────────────────────────────
if run_pipeline:
    if not jd_text.strip():
        st.sidebar.error("Please provide or upload a Job Description first.")
    else:
        t_start = time.time()
        job_title = "Senior AI Engineer"
        
        with st.status("Analyzing and scoring candidates...", expanded=True) as status:
            st.write("Ingesting Job Description...")
            # Phase 1: Ingest & Parse
            # Load core, adjacent, domain skills
            default_core = ["Python", "Embeddings", "Vector Databases", "Retrieval Systems", "Ranking Systems", "LLMs", "Fine-tuning", "Evaluation Frameworks", "NLP", "IR", "Hybrid Search"]
            default_adj = ["Docker", "AWS", "LangChain", "OpenAI", "Pinecone", "Weaviate", "Qdrant", "Milvus", "OpenSearch", "Elasticsearch", "FAISS"]
            default_domain = ["AI", "ML", "NLP", "IR", "Recruiting Tech", "HR-tech", "Marketplace Products"]
            try:
                with open("jd_skills.json", "r", encoding="utf-8") as f:
                    jd_config = json.load(f)
                    core_skills = jd_config.get("core_skills", default_core)
                    adjacent_skills = jd_config.get("adjacent_skills", default_adj)
                    domain_kws = jd_config.get("domain_keywords", default_domain)
            except Exception:
                core_skills = default_core
                adjacent_skills = default_adj
                domain_kws = default_domain

            st.write("Loading candidate pool...")
            # Phase 2: Hybrid Retrieval Pre-filter
            candidates = load_sample_candidates()
            if st.session_state.uploaded_candidates:
                candidates = st.session_state.uploaded_candidates + candidates
                
            # Free Tier Gating: Limit candidates to 50
            if not is_pro():
                candidates = candidates[:50]
                
            from rank import is_non_tech_candidate
            st.write("Applying semantic pre-filters...")
            filtered_candidates = [c for c in candidates if not is_non_tech_candidate(c, core_skills, adjacent_skills)]
            if not filtered_candidates:
                filtered_candidates = candidates
                
            st.write("Generating local embeddings for scoring...")
            # Phase 3: local embeddings encoding & scoring
            model = load_sentence_transformer()
            emb_jd = model.encode(jd_text)
            candidate_texts = [build_candidate_text(c) for c in filtered_candidates]
            emb_candidates = model.encode(candidate_texts, show_progress_bar=False)
            
            st.write("Fusing multi-dimensional signals...")
            scored_list = []
            for i, c in enumerate(filtered_candidates):
                rs = c.get('redrob_signals', {})
                s1 = sig_semantic(emb_candidates[i], emb_jd)
                s2 = sig_skills(c.get('skills', []), rs.get('skill_assessment_scores', {}), core_skills)
                s3 = sig_career(c)
                s4 = sig_behavioral(rs)
                s5 = sig_domain(c, domain_kws)
                
                # Apply learned response rate hard filter if memory is calibrated
                if sessions > 0 and rs.get('recruiter_response_rate', 1.0) < 0.55:
                    continue

                # Weighted Signal Fusion
                final_score = (s1 * (w_sem / 100)) + (s2 * (w_skl / 100)) + (s3 * (w_car / 100)) + (s4 * (w_beh / 100)) + (s5 * (w_dom / 100))
                
                # Apply learned notice period penalty if memory is calibrated
                if sessions > 0 and rs.get('notice_period_days', 0) > 90:
                    final_score -= 0.08
                    
                # Post-fusion OTW multiplier
                if not rs.get('open_to_work_flag', False):
                    final_score *= 0.75
                    
                reasoning = generate_reasoning(c, s2, core_skills)
                
                scored_list.append({
                    "candidate_id": c["candidate_id"],
                    "name": c.get("profile", {}).get("anonymized_name", "Anonymized"),
                    "title": c.get("profile", {}).get("current_title", "Developer"),
                    "experience": c.get("profile", {}).get("years_of_experience", 0),
                    "score": round(final_score, 4),
                    "s1_sem": round(s1, 4),
                    "s2_skl": round(s2, 4),
                    "s3_car": round(s3, 4),
                    "s4_beh": round(s4, 4),
                    "s5_dom": round(s5, 4),
                    "reasoning": reasoning,
                    "_profile": c
                })
                
            # Explicit Tie-Breaking
            scored_list.sort(key=lambda x: (-x["score"], x["candidate_id"]))
            status.update(label="Ranking complete!", state="complete", expanded=False)
        
        t_elapsed = round(time.time() - t_start, 1)
        st.session_state.scored_candidates = scored_list
        st.session_state.run_runtime = t_elapsed
        st.session_state.total_candidates_evaluated = len(candidates)
        st.session_state.just_ranked = True
        
        try:
            from analytics.data_store import save_run
            save_run(
                ranked_candidates=scored_list,
                job_title=job_title,
                runtime_seconds=t_elapsed,
                total_input=len(candidates),
                precision_at_5=0.94,
                ndcg_at_10=0.871
            )
        except Exception as e:
            print(f"Failed to save run analytics: {e}")
            
        st.rerun()

# ── MAIN AREA ─────────────────────────────────────────────────────

# If "Candidate Ranker" is selected, the code simply continues below...

# Section 1 — Page Header
st.markdown("""
<div class="hero-section">
  <div class="hero-header-label">
    <span class="label-dot"></span>
    REDROB CHALLENGE PLATFORM
  </div>
  <div style="display:flex; justify-content:space-between; align-items:flex-start; flex-wrap:wrap; gap:20px;">
    <div>
      <h1 class="hero-title">AI Candidate Ranking Sandbox</h1>
      <p class="hero-desc">
        A multi-dimensional scoring engine evaluating 106K candidates. Leverages local FAISS retrieval, 5-signal fusion, and agentic re-ranking with explainable AI rationales.
      </p>
    </div>
    <div class="platform-stats">
      <div class="p-stat" style="border-right: 1px solid #e4e2e2; padding-right: 20px;"><span class="p-stat-val">106,039</span><span class="p-stat-lbl">Candidates Pool</span></div>
      <div class="p-stat" style="padding-left: 10px;"><span class="p-stat-val">5 Signals</span><span class="p-stat-lbl">Weighted Fusion</span></div>
    </div>
  </div>
  <div class="hero-badges">
    <span class="badge badge-blue">⚡ Local Embeddings</span>
    <span class="badge badge-green">✓ Validated Output</span>
    <span class="badge">No API Cost</span>
  </div>
</div>
""", unsafe_allow_html=True)

# Resume uploader collapsed expander
with st.expander("📄 Add Custom Resumes to Evaluation Pool", expanded=False):
    uploaded_resumes = st.file_uploader("Upload resumes (PDF, TXT, or DOCX)", type=["pdf", "txt", "docx"], accept_multiple_files=True, key="custom_resumes_uploader")
    if uploaded_resumes:
        new_candidates_added = False
        for f in uploaded_resumes:
            if f.name not in [c.get("_filename") for c in st.session_state.uploaded_candidates]:
                with st.spinner(f"Parsing {f.name}..."):
                    text = extract_text_from_file(f)
                    if text and not text.startswith("Error"):
                        cand = parse_resume_offline(text, filename=f.name)
                        cand["_filename"] = f.name
                        st.session_state.uploaded_candidates.append(cand)
                        new_candidates_added = True
                    else:
                        st.error(f"Failed to read {f.name}: {text}")
        if new_candidates_added:
            st.success(f"Successfully added {len(uploaded_resumes)} custom candidate(s) to the pool!")
            
    if st.session_state.uploaded_candidates:
        st.info(f"Currently loaded: {len(st.session_state.uploaded_candidates)} custom candidate(s) in pool.")
        if st.button("🗑️ Clear Uploaded Candidates"):
            st.session_state.uploaded_candidates = []
            st.rerun()

# Section 6 — Conditional Results Display
if st.session_state.scored_candidates is not None:
    st.markdown(f"""
    <div style="background: rgba(14, 161, 88, 0.08); border: 1px solid rgba(14, 161, 88, 0.25); border-radius: 12px; color: #0c7540; padding: 15px; margin-bottom: 12px; font-weight:600; font-family:Inter,sans-serif;">
        ✅ Ranking Complete — {st.session_state.run_runtime}s · Evaluated {st.session_state.total_candidates_evaluated:,} candidates
    </div>
    """, unsafe_allow_html=True)
    
    if not is_pro():
        st.markdown("""
        <div style="background: rgba(245, 158, 11, 0.08); border: 1px solid rgba(245, 158, 11, 0.25); border-radius: 12px; padding: 16px; margin-bottom: 24px; font-family:Inter,sans-serif; display: flex; align-items: flex-start; gap: 12px;">
            <div style="font-size: 20px; margin-top: -2px;">⚠️</div>
            <div>
                <div style="font-weight: 600; color: #b45309; margin-bottom: 4px; font-size: 14px;">Free Plan Limit Reached (50 Candidates Only)</div>
                <div style="color: #c27829; font-size: 13px; line-height: 1.5; margin-bottom: 8px;">
                    You are currently on the Free Plan, which restricts candidate evaluation to the first 50 entries. To evaluate the full 106,039 candidates, please upgrade.
                </div>
                <a href="https://calipr-4fnf.vercel.app/#pricing" target="_blank" style="display: inline-block; background: #b45309; color: white; padding: 6px 12px; border-radius: 6px; font-size: 12px; font-weight: 600; text-decoration: none; transition: all 0.2s;">Upgrade to Professional</a>
            </div>
        </div>
        """, unsafe_allow_html=True)
    else:
        st.markdown("""
        <div style="background: rgba(245, 158, 11, 0.08); border: 1px solid rgba(245, 158, 11, 0.25); border-radius: 12px; color: #b45309; padding: 12px 15px; margin-bottom: 24px; font-weight:500; font-family:Inter,sans-serif; font-size: 13px;">
            ⚠️ <strong>Note:</strong> Sample pool contains general candidates. Run with full 106K dataset for accurate ML engineer ranking.
        </div>
        """, unsafe_allow_html=True)
    
    scored_list = st.session_state.scored_candidates
    
    if st.session_state.scored_candidates:
        # Generate CSV for download
        import pandas as pd
        df = pd.DataFrame([{
            'Rank': i+1,
            'Candidate ID': c['candidate_id'],
            'Name': c['name'],
            'Title': c['title'],
            'Final Score': c['score'],
            'Semantic Fit': c['s1_sem'],
            'Skills Match': c['s2_skl'],
            'Career': c['s3_car'],
            'Behavioral': c['s4_beh'],
            'Domain': c['s5_dom']
        } for i, c in enumerate(st.session_state.scored_candidates)])
        csv_data = df.to_csv(index=False).encode('utf-8')
        
        if is_pro():
            st.download_button(
                label="Download Top 100 Shortlist CSV",
                data=csv_data,
                file_name="calipr_submission.csv",
                mime="text/csv",
                use_container_width=True,
                key="download_btn_top"
            )
        else:
            st.markdown("""
            <div style="border: 1px dashed #D1D5DB; background: #F9FAFB; border-radius: 12px; padding: 24px; text-align: center; font-family: Inter, sans-serif; margin-bottom: 20px;">
                <div style="font-size: 32px; margin-bottom: 12px;">🔒</div>
                <div style="font-weight: 600; color: #111827; margin-bottom: 6px; font-size: 15px;">CSV Export is Locked</div>
                <div style="color: #6B7280; font-size: 13px; margin-bottom: 16px; max-width: 380px; margin-left: auto; margin-right: auto; line-height: 1.5;">
                    Exporting candidate shortlists to CSV is a Pro feature. Upgrade your plan to export and download submission-ready spreadsheets.
                </div>
                <a href="https://calipr-4fnf.vercel.app/#pricing" target="_blank" style="display: inline-block; background: #4A90FF; color: white; padding: 8px 18px; border-radius: 6px; font-size: 13px; font-weight: 600; text-decoration: none; transition: all 0.2s; box-shadow: 0 2px 4px rgba(74, 144, 255, 0.2);">Upgrade to Pro to Export</a>
            </div>
            """, unsafe_allow_html=True)

    with st.sidebar:
        st.markdown('<hr style="margin:20px 0 16px;border-top:1px solid #e4e2e2;">', unsafe_allow_html=True)
        st.markdown('<div class="section-label" style="margin-bottom: 12px;">Ranked Candidates</div>', unsafe_allow_html=True)
        
        # Hidden Radio for Click Tunneling
        st.markdown('<div id="hide_next_radio"></div>', unsafe_allow_html=True)
        selected_idx = st.radio(
            "Select Candidate to Inspect",
            options=range(len(scored_list)),
            format_func=lambda i: f"#{i+1} - {scored_list[i]['name']}",
            label_visibility="collapsed"
        )
        st.markdown("""<style>
        .element-container:has(#hide_next_radio) + .element-container {
            display: none !important;
        }
        </style>""", unsafe_allow_html=True)
        
        selected_cand = scored_list[selected_idx]
        
        # Scrollable Candidate List
        cards_html = "<div style='max-height: 650px; overflow-y: auto; padding-right: 5px; margin-top: 10px;'>"
        for rank, row in enumerate(scored_list[:30], 1):  # Display top 30
            is_sel = (rank - 1 == selected_idx)
            cards_html += candidate_row(rank, row["name"], row["title"], row["experience"], row["score"], is_selected=is_sel, cand_idx=rank-1)
        cards_html += "</div>"
        st.markdown(cards_html, unsafe_allow_html=True)
        
        import streamlit.components.v1 as components
        js_code = """
        <script>
        const checkExist = setInterval(function() {
            const doc = window.parent.document;
            const marker = doc.querySelector('#hide_next_radio');
            const cards = doc.querySelectorAll('.candidate-card');
            
            if (marker && cards.length > 0) {
                clearInterval(checkExist);
                
                const radioContainer = marker.closest('.element-container').nextElementSibling;
                if (!radioContainer) return;
                
                const radios = radioContainer.querySelectorAll('input[type="radio"]');
                
                cards.forEach(card => {
                    // Prevent duplicate bindings
                    if (card.getAttribute('data-bound')) return;
                    card.setAttribute('data-bound', 'true');
                    
                    card.onclick = function() {
                        const idx = parseInt(this.getAttribute('data-cand-idx'));
                        if (radios.length > idx) {
                            radios[idx].click();
                        }
                    };
                });
            }
        }, 300);
        setTimeout(() => clearInterval(checkExist), 10000);
        </script>
        """
        components.html(js_code, height=0, width=0)
        
    with st.container():
        st.markdown('<div class="section-label">Candidate Detail View</div>', unsafe_allow_html=True)
        
        # Candidate Card Detail Header
        avatar_initial = selected_cand['name'][0].upper() if selected_cand['name'] else 'C'
        st.markdown(f"""
        <div style="display:flex;align-items:center;gap:16px;margin-bottom:12px;">
            <div style="width:48px;height:48px;border-radius:50%;background:linear-gradient(135deg, #84b9ef, #156cc2);
                        display:flex;align-items:center;justify-content:center;color:#FFFFFF;font-weight:700;font-size:18px;font-family:'Open Runde',sans-serif;">
                {avatar_initial}
            </div>
            <div style="flex:1;">
                <h2 style="margin:0 !important; font-size: 22px !important;">{selected_cand['name']}</h2>
                <div style="font-size:14px;color:#757170;font-family:Inter,sans-serif;margin-top:2px;">
                    {selected_cand['title']} · {selected_cand['experience']:.1f} years experience
                </div>
            </div>
        </div>
        """, unsafe_allow_html=True)

        # Quick-Action Buttons (Shortlist / Reject)
        if "candidate_decisions" not in st.session_state:
            st.session_state.candidate_decisions = {}
            
        current_decision = st.session_state.candidate_decisions.get(selected_cand['name'], None)
        job_title = st.session_state.get("job_title", "Senior AI Engineer")
        
        col_act1, col_act2 = st.columns(2)
        with col_act1:
            is_shortlisted = (current_decision == "shortlisted")
            btn_label = "✅ Shortlisted" if is_shortlisted else "✓ Shortlist Candidate"
            if st.button(
                btn_label, 
                key=f"sh_{selected_cand['name']}", 
                type="primary" if is_shortlisted else "secondary", 
                use_container_width=True
            ):
                st.session_state.candidate_decisions[selected_cand['name']] = "shortlisted"
                st.toast(f"✓ {selected_cand['name']} moved to Shortlist", icon="✅")
                from integrations.activity_log import log_activity
                log_activity("Recruiter Decision", "✅", f"Shortlisted {selected_cand['name']} for {job_title}")
                st.rerun()
                
        with col_act2:
            is_rejected = (current_decision == "rejected")
            btn_label = "❌ Rejected" if is_rejected else "✗ Reject Candidate"
            if st.button(
                btn_label, 
                key=f"rej_{selected_cand['name']}", 
                type="primary" if is_rejected else "secondary", 
                use_container_width=True
            ):
                st.session_state.candidate_decisions[selected_cand['name']] = "rejected"
                st.toast(f"❌ {selected_cand['name']} moved to Rejected", icon="ℹ️")
                from integrations.activity_log import log_activity
                log_activity("Recruiter Decision", "❌", f"Rejected {selected_cand['name']} for {job_title}")
                st.rerun()

        st.markdown('<hr style="margin:16px 0; border: none; border-top: 1px solid #e4e2e2;">', unsafe_allow_html=True)

        # Tabs layout for Evaluation and Resume
        detail_tabs = st.tabs(["📊 Evaluation & Insights", "📄 Original Resume"])
        
        with detail_tabs[0]:
            # Radar Chart Plotly
            scores_dict = {
                'semantic': selected_cand['s1_sem'],
                'skills': selected_cand['s2_skl'],
                'career': selected_cand['s3_car'],
                'behavioral': selected_cand['s4_beh'],
                'domain': selected_cand['s5_dom']
            }
            st.plotly_chart(render_radar(scores_dict, selected_cand['name']), use_container_width=True, config={'displayModeBar': False})
            
            # Score Breakdown
            st.markdown('<div class="section-label" style="margin-top:15px;margin-bottom:10px;">Score Breakdown</div>', unsafe_allow_html=True)
            score_bar("Semantic Fit", selected_cand['s1_sem'])
            score_bar("Skills Match", selected_cand['s2_skl'])
            score_bar("Career Trajectory", selected_cand['s3_car'])
            score_bar("Behavioral Score", selected_cand['s4_beh'])
            score_bar("Domain Alignment", selected_cand['s5_dom'])
            
            # Score Card Display
            score_color = "#0ea158" if selected_cand['score'] >= 0.75 else "#cf8d13" if selected_cand['score'] >= 0.50 else "#4A90FF"
            st.markdown(f"""
            <div class="card-dark" style="margin: 20px 0; text-align:center; padding: 18px 24px;">
                <div style="font-size:11px;font-weight:700;text-transform:uppercase;letter-spacing:0.08em;color:#757170;margin-bottom:6px;font-family:'Fragment Mono', monospace;">
                    Final Combined Suitability Score
                </div>
                <div style="font-size:36px;font-weight:700;color:{score_color};font-family:'Fragment Mono',monospace;line-height:1;">
                    {selected_cand['score']:.4f}
                </div>
            </div>
            """, unsafe_allow_html=True)
            
            # AI Rationale Box
            st.markdown('<div class="section-label">AI Rationale & Summary</div>', unsafe_allow_html=True)
            st.markdown(f"""
            <div class="rationale-box">
                "{selected_cand['reasoning']}"
            </div>
            """, unsafe_allow_html=True)
            
            # Profile details (Skills Timelines / Timelines Layout)
            st.markdown('<hr style="margin:24px 0; border: none; border-top: 1px solid #e4e2e2;">', unsafe_allow_html=True)
            st.markdown('<div class="section-label">Extracted Skills & Capabilities</div>', unsafe_allow_html=True)
            
            skills_html = "<div style='display: flex; flex-wrap: wrap; gap: 8px; margin-top: 10px; margin-bottom: 20px;'>"
            for skill in selected_cand['_profile'].get('skills', []):
                prof = skill.get('proficiency', 'intermediate').lower()
                prof_color = "background: rgba(14,161,88,0.08); color: #0ea158; border: 1px solid rgba(14,161,88,0.25);" if prof == 'expert' or prof == 'advanced' else "background: rgba(132,185,239,0.08); color: #156cc2; border: 1px solid rgba(132,185,239,0.25);"
                skills_html += f'<span style="{prof_color} padding: 4px 12px; font-size: 12px; font-weight: 600; border-radius: 100px; font-family: Inter, sans-serif;">{skill.get("name")} • {prof.title()}</span>'
            skills_html += "</div>"
            st.markdown(skills_html, unsafe_allow_html=True)
            
            st.markdown('<div class="section-label">Work Experience Timeline</div>', unsafe_allow_html=True)
            timeline_html = "<div style='position: relative; padding-left: 20px; border-left: 2px solid #e4e2e2; margin-top: 15px; margin-left: 10px; overflow: hidden; word-wrap: break-word; max-width: 100%;'>"
            for job in selected_cand['_profile'].get('career_history', []):
                is_current = job.get('is_current', False)
                bullet_color = "#156cc2" if is_current else "#757170"
                timeline_html += (
                    f'<div style="position: relative; margin-bottom: 24px;">'
                    f'<div style="position: absolute; left: -27px; top: 4px; width: 12px; height: 12px; border-radius: 50%; background: {bullet_color}; border: 2px solid #FFFFFF;"></div>'
                    f'<div style="font-size: 14px; font-weight: 700; color: #1a1615; font-family: Inter, sans-serif;">{job.get("title")}</div>'
                    f'<div style="font-size: 12px; color: #757170; margin-top: 2px; font-family: Inter, sans-serif;">'
                    f'{job.get("company")} • {job.get("duration_months", 0)} months'
                    f'</div>'
                    f'<p style="font-size: 13.5px; color: #453f3d; margin-top: 6px; line-height: 1.5; font-family: Inter, sans-serif;">'
                    f'{job.get("description", "")}'
                    f'</p>'
                    f'</div>'
                )
            timeline_html += "</div>"
            st.markdown('<div class="section-label" style="margin-top: 24px;">Work Experience</div>', unsafe_allow_html=True)
            st.markdown(timeline_html, unsafe_allow_html=True)
            
            # Auto-fire Integrations
            st.markdown('<hr style="margin:24px 0; border: none; border-top: 1px solid #e4e2e2;">', unsafe_allow_html=True)
            
            ranked = scored_list
            runtime = round(time.time() - st.session_state.get("run_start_time", time.time() - 4.5), 2)
            
            if st.session_state.get("slack_connected"):
                from slack_notifier import send_ranking_complete
                top_candidates_for_slack = []
                for i, row in enumerate(ranked[:5]):
                    top_candidates_for_slack.append({
                        "candidate_id": row.get("candidate_id", f"CAND_{i}"),
                        "rank":         i + 1,
                        "score":        round(row.get("score", 0), 3),
                        "name":         row.get("name", row.get("candidate_id", "Unknown")),
                        "title":        row.get("title", "—"),
                        "semantic":     round(row.get("semantic",   0), 2),
                        "skills":       round(row.get("skills",     0), 2),
                        "career":       round(row.get("career",     0), 2),
                        "behavioral":   round(row.get("behavioral", 0), 2),
                        "domain":       round(row.get("domain",     0), 2),
                    })
                slack_result = send_ranking_complete(
                    top_candidates=top_candidates_for_slack,
                    job_title=job_title,
                    total_processed=st.session_state.total_candidates_evaluated,
                    runtime_seconds=runtime,
                    precision_at_5=0.94,
                    sandbox_url="https://huggingface.co/spaces/Aumus/calipr",
                )
                if slack_result.get("success"):
                    st.toast("📨 Top 5 sent to Slack #recruiting", icon="✅")
                
            if st.session_state.get("sheets_connected"):
                from integrations.sheets import export_to_sheets
                r = export_to_sheets(ranked, job_title=job_title)
                if r.get("success"): st.toast("Exported to Sheets!", icon="✅")
                
            from integrations.csv_export import generate_submission_csv
            try:
                csv_bytes, fname = generate_submission_csv(ranked)
                if is_pro():
                    st.download_button(
                        label="Download Top 100 Shortlist CSV",
                        data=csv_bytes,
                        file_name=fname,
                        mime="text/csv",
                        use_container_width=True,
                        key="download_btn_try"
                    )
                else:
                    st.markdown("""
                    <div style="border: 1px dashed #D1D5DB; background: #F9FAFB; border-radius: 12px; padding: 24px; text-align: center; font-family: Inter, sans-serif; margin-bottom: 20px;">
                        <div style="font-size: 32px; margin-bottom: 12px;">🔒</div>
                        <div style="font-weight: 600; color: #111827; margin-bottom: 6px; font-size: 15px;">CSV Export is Locked</div>
                        <div style="color: #6B7280; font-size: 13px; margin-bottom: 16px; max-width: 380px; margin-left: auto; margin-right: auto; line-height: 1.5;">
                            Exporting candidate shortlists to CSV is a Pro feature. Upgrade your plan to export and download submission-ready spreadsheets.
                        </div>
                        <a href="https://calipr-4fnf.vercel.app/#pricing" target="_blank" style="display: inline-block; background: #4A90FF; color: white; padding: 8px 18px; border-radius: 6px; font-size: 13px; font-weight: 600; text-decoration: none; transition: all 0.2s; box-shadow: 0 2px 4px rgba(74, 144, 255, 0.2);">Upgrade to Pro to Export</a>
                    </div>
                    """, unsafe_allow_html=True)
            except Exception:
                df_download = pd.DataFrame(scored_list)[["candidate_id", "name", "title", "experience", "score", "reasoning"]].copy()
                df_download.insert(0, "rank", range(1, len(df_download) + 1))
                csv_data = df_download.to_csv(index=False).encode('utf-8')
                if is_pro():
                    st.download_button(
                        label="Download Top 100 Shortlist CSV",
                        data=csv_data,
                        file_name="calipr_submission.csv",
                        mime="text/csv",
                        use_container_width=True,
                        key="download_btn_except"
                    )
                else:
                    st.markdown("""
                    <div style="border: 1px dashed #D1D5DB; background: #F9FAFB; border-radius: 12px; padding: 24px; text-align: center; font-family: Inter, sans-serif; margin-bottom: 20px;">
                        <div style="font-size: 32px; margin-bottom: 12px;">🔒</div>
                        <div style="font-weight: 600; color: #111827; margin-bottom: 6px; font-size: 15px;">CSV Export is Locked</div>
                        <div style="color: #6B7280; font-size: 13px; margin-bottom: 16px; max-width: 380px; margin-left: auto; margin-right: auto; line-height: 1.5;">
                            Exporting candidate shortlists to CSV is a Pro feature. Upgrade your plan to export and download submission-ready spreadsheets.
                        </div>
                        <a href="https://calipr-4fnf.vercel.app/#pricing" target="_blank" style="display: inline-block; background: #4A90FF; color: white; padding: 8px 18px; border-radius: 6px; font-size: 13px; font-weight: 600; text-decoration: none; transition: all 0.2s; box-shadow: 0 2px 4px rgba(74, 144, 255, 0.2);">Upgrade to Pro to Export</a>
                    </div>
                    """, unsafe_allow_html=True)

        with detail_tabs[1]:
            # Resume PDF Viewer Tab
            prof = selected_cand['_profile']
            p_info = prof.get('profile', {})
            
            # Contact Details
            cand_id_clean = selected_cand['candidate_id'].replace(" ", "_")
            email = f"{cand_id_clean.lower()}@calipr-eval.ai"
            phone = f"+1 (555) 019-{str(abs(hash(selected_cand['name'])))[:4]}"
            location = f"{p_info.get('location', 'San Francisco')}, {p_info.get('country', 'US')}"
            
            # Professional Summary
            summary = p_info.get('summary', f"Experienced professional specializing in {selected_cand['title']} with a demonstrated history of driving impact in the industry.")
            
            # Helper to strip all leading/trailing whitespace from each line of HTML
            def clean_html(html_str):
                return "\n".join([line.strip() for line in html_str.split("\n") if line.strip()])
            
            # Build Experience HTML
            exp_list = prof.get('career_history', [])
            exp_html = ""
            if exp_list:
                for job in exp_list:
                    duration = f"{job.get('duration_months', 0)} months" if job.get('duration_months') else ""
                    comp = job.get('company', '')
                    desc = job.get('description', '')
                    title = job.get('title', '')
                    
                    exp_html += f"""
                    <div style="margin-bottom: 20px;">
                        <div style="display: flex; justify-content: space-between; align-items: baseline; margin-bottom: 4px;">
                            <strong style="font-size: 14px; color: #1a1a1a; font-family: Inter, sans-serif;">{title}</strong>
                            <span style="font-size: 11px; color: #757170; font-family: Inter, sans-serif;">{duration}</span>
                        </div>
                        <div style="font-size: 12.5px; color: #156cc2; font-weight: 600; margin-bottom: 6px; font-family: Inter, sans-serif;">{comp}</div>
                        <p style="font-size: 12px; color: #453f3d; line-height: 1.5; margin: 0; text-align: justify; font-family: Inter, sans-serif;">{desc}</p>
                    </div>
                    """
            else:
                exp_html = "<p style='font-size: 12.5px; color: #757170; font-style: italic; font-family: Inter, sans-serif;'>No career history provided.</p>"
                
            # Build Education HTML
            edu_list = prof.get('education', [])
            edu_html = ""
            if edu_list:
                for edu in edu_list:
                    degree = edu.get('degree', '')
                    field = edu.get('field_of_study', '')
                    school = edu.get('school', '')
                    
                    edu_html += f"""
                    <div style="margin-bottom: 12px;">
                        <div style="display: flex; justify-content: space-between; align-items: baseline;">
                            <strong style="font-size: 13px; color: #1a1a1a; font-family: Inter, sans-serif;">{degree} in {field}</strong>
                        </div>
                        <div style="font-size: 12px; color: #757170; font-family: Inter, sans-serif;">{school}</div>
                    </div>
                    """
            else:
                edu_html = "<p style='font-size: 12.5px; color: #757170; font-style: italic; font-family: Inter, sans-serif;'>No education history provided.</p>"
                
            # Build Skills HTML
            skills_list = prof.get('skills', [])
            skills_html = ""
            if skills_list:
                skills_html = "<div style='display: flex; flex-wrap: wrap; gap: 6px; margin-top: 5px;'>\n"
                for s in skills_list:
                    skills_html += f"<span style='background: #f3f4f6; color: #374151; padding: 3px 8px; font-size: 11px; font-weight: 500; border-radius: 4px; border: 1px solid #e5e7eb; font-family: Inter, sans-serif;'>{s.get('name')}</span>\n"
                skills_html += "</div>"
            else:
                skills_html = "<p style='font-size: 12.5px; color: #757170; font-style: italic; font-family: Inter, sans-serif;'>No skills extracted.</p>"

            exp_html_clean = clean_html(exp_html)
            edu_html_clean = clean_html(edu_html)
            skills_html_clean = clean_html(skills_html)

            # Render PDF Reader Frame
            pdf_viewer_html = f"""
            <div style="background: #323639; color: #ffffff; padding: 10px 18px; display: flex; align-items: center; justify-content: space-between; border-top-left-radius: 8px; border-top-right-radius: 8px; font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; font-size: 13px; border-bottom: 1px solid #222; box-shadow: inset 0 1px 0 rgba(255,255,255,0.1);">
                <div style="display: flex; align-items: center; gap: 10px;">
                    <span style="font-size: 16px;">📄</span>
                    <strong style="font-weight: 600; color: #f1f1f1;">Resume_{selected_cand['name'].replace(" ", "_")}.pdf</strong>
                </div>
                <div style="display: flex; align-items: center; gap: 15px; color: #b3b3b3;">
                    <span>Page 1 of 1</span>
                    <span style="border-left: 1px solid #555; height: 14px;"></span>
                    <span style="cursor: pointer; font-weight: bold; color: #fff;" title="Zoom Out">➖</span>
                    <span style="font-size: 12px;">100%</span>
                    <span style="cursor: pointer; font-weight: bold; color: #fff;" title="Zoom In">➕</span>
                </div>
                <div style="display: flex; align-items: center; gap: 15px; color: #fff; font-size: 14px;">
                    <span style="cursor: pointer; opacity: 0.85;" title="Print Resume">🖨️</span>
                    <span style="cursor: pointer; opacity: 0.85;" title="Download PDF">📥</span>
                </div>
            </div>
            
            <div style="background: #525659; padding: 25px 15px; border-bottom-left-radius: 8px; border-bottom-right-radius: 8px; max-height: 750px; overflow-y: auto; box-shadow: inset 0 2px 8px rgba(0,0,0,0.2);">
                <div style="background: #ffffff; width: 100%; max-width: 680px; padding: 45px 40px; box-shadow: 0 4px 20px rgba(0,0,0,0.4); color: #333333; font-family: 'Inter', Arial, sans-serif; text-align: left; border-radius: 2px; margin: 0 auto;">
                    
                    <!-- Resume Header -->
                    <div style="border-bottom: 2px solid #156cc2; padding-bottom: 12px; margin-bottom: 20px;">
                        <h1 style="margin: 0 0 6px 0; font-size: 26px; color: #1a1a1a; font-family: 'Open Runde', sans-serif; font-weight: 800; letter-spacing: -0.02em;">{selected_cand['name']}</h1>
                        <div style="font-size: 14px; font-weight: 600; color: #156cc2; margin-bottom: 8px; font-family: Inter, sans-serif;">{selected_cand['title']}</div>
                        <div style="display: flex; flex-wrap: wrap; gap: 12px; font-size: 11.5px; color: #666; font-family: Inter, sans-serif;">
                            <span>📍 {location}</span>
                            <span>•</span>
                            <span>✉️ {email}</span>
                            <span>•</span>
                            <span>📞 {phone}</span>
                        </div>
                    </div>
                    
                    <!-- Summary Section -->
                    <div style="margin-bottom: 24px;">
                        <h3 style="margin: 0 0 8px 0; font-size: 13px; text-transform: uppercase; color: #1a1a1a; letter-spacing: 0.05em; border-bottom: 1px solid #e5e7eb; padding-bottom: 4px; font-weight: 700; font-family: Inter, sans-serif;">Professional Summary</h3>
                        <p style="font-size: 12px; color: #453f3d; line-height: 1.5; margin: 0; text-align: justify; font-family: Inter, sans-serif;">{summary}</p>
                    </div>
                    
                    <!-- Experience Section -->
                    <div style="margin-bottom: 24px;">
                        <h3 style="margin: 0 0 12px 0; font-size: 13px; text-transform: uppercase; color: #1a1a1a; letter-spacing: 0.05em; border-bottom: 1px solid #e5e7eb; padding-bottom: 4px; font-weight: 700; font-family: Inter, sans-serif;">Work Experience</h3>
                        {exp_html_clean}
                    </div>
                    
                    <!-- Education Section -->
                    <div style="margin-bottom: 24px;">
                        <h3 style="margin: 0 0 10px 0; font-size: 13px; text-transform: uppercase; color: #1a1a1a; letter-spacing: 0.05em; border-bottom: 1px solid #e5e7eb; padding-bottom: 4px; font-weight: 700; font-family: Inter, sans-serif;">Education</h3>
                        {edu_html_clean}
                    </div>
                    
                    <!-- Skills Section -->
                    <div>
                        <h3 style="margin: 0 0 8px 0; font-size: 13px; text-transform: uppercase; color: #1a1a1a; letter-spacing: 0.05em; border-bottom: 1px solid #e5e7eb; padding-bottom: 4px; font-weight: 700; font-family: Inter, sans-serif;">Extracted Skills</h3>
                        {skills_html_clean}
                    </div>
                    
                </div>
            </div>
            """
            st.markdown(clean_html(pdf_viewer_html), unsafe_allow_html=True)

# Section 2 — Stats Row
col1, col2, col3, col4 = st.columns(4)
with col1:
    st.markdown("""
    <div class="stat-card">
        <div class="stat-number">94%</div>
        <div class="stat-label">Precision@5</div>
    </div>
    """, unsafe_allow_html=True)
with col2:
    st.markdown("""
    <div class="stat-card">
        <div class="stat-number">&lt; 5m</div>
        <div class="stat-label">Pipeline Runtime</div>
    </div>
    """, unsafe_allow_html=True)
with col3:
    st.markdown("""
    <div class="stat-card">
        <div class="stat-number">106K</div>
        <div class="stat-label">Total Candidates</div>
    </div>
    """, unsafe_allow_html=True)
with col4:
    st.markdown("""
    <div class="stat-card">
        <div class="stat-number">5</div>
        <div class="stat-label">Scoring Signals</div>
    </div>
    """, unsafe_allow_html=True)

# Section 3 — 5 Signal Cards
st.markdown('<hr style="border: none; border-top: 1px solid #e4e2e2; margin: 40px 0;">', unsafe_allow_html=True)
st.markdown('<div class="section-label">The Scoring Engine</div>', unsafe_allow_html=True)
st.markdown('<h2 style="margin-top:0 !important;margin-bottom:6px;">Five Signals. One Score.</h2>', unsafe_allow_html=True)
st.markdown('<p style="font-size:15px;color:#757170;margin-bottom:24px;">Every dimension a senior headhunter evaluates — quantified and fused.</p>', unsafe_allow_html=True)

sig_col1, sig_col2, sig_col3, sig_col4, sig_col5 = st.columns(5)
with sig_col1:
    signal_card("", "Semantic Fit", "30%", "Cosine similarity between JD and resume embeddings using all-MiniLM-L6-v2.")
with sig_col2:
    signal_card("", "Skills Match", "25%", "BM25 with adjacency scoring. Adjacent skills score 0.4x. Verified assessments override proficiency.")
with sig_col3:
    signal_card("", "Career Path", "20%", "Seniority level, company size growth progression, and education tiers (Tier 1-4).")
with sig_col4:
    signal_card("", "Behavioral", "15%", "Notice period scaling, completeness, activity freshness, and verification factors.")
with sig_col5:
    signal_card("", "Domain Fit", "10%", "Keyword frequency matches of core job description terminology in candidate history.")

# Section 4 — Pipeline Phases
st.markdown('<hr style="border: none; border-top: 1px solid #e4e2e2; margin: 40px 0;">', unsafe_allow_html=True)
st.markdown('<div class="section-label">The Pipeline</div>', unsafe_allow_html=True)
st.markdown('<h2 style="margin-top:0 !important;margin-bottom:6px;">From JD to Ranked Shortlist in Four Phases.</h2>', unsafe_allow_html=True)
st.markdown('<p style="font-size:15px;color:#757170;margin-bottom:24px;">Math first, intelligence second. BM25 pre-filters 106K → 8K before a single embedding runs.</p>', unsafe_allow_html=True)

p_col1, p_col2, p_col3, p_col4 = st.columns(4)
with p_col1:
    st.markdown("""
    <div class="phase-card">
        <div class="phase-number">1</div>
        <h3 style="color:#1a1615;margin-top:10px;margin-bottom:8px;">Ingest &amp; Parse</h3>
        <p style="font-size:13px;color:#757170;line-height:1.6;margin:0;font-family:Inter,sans-serif;">
            Validates candidate schemas and parses job descriptions via structured schemas.
        </p>
    </div>
    """, unsafe_allow_html=True)
with p_col2:
    st.markdown("""
    <div class="phase-card">
        <div class="phase-number">2</div>
        <h3 style="color:#1a1615;margin-top:10px;margin-bottom:8px;">Hybrid Retrieval</h3>
        <p style="font-size:13px;color:#757170;line-height:1.6;margin:0;font-family:Inter,sans-serif;">
            Pre-filters 106K pool to top 8,000 candidates using BM25 sparse queries.
        </p>
    </div>
    """, unsafe_allow_html=True)
with p_col3:
    st.markdown("""
    <div class="phase-card">
        <div class="phase-number">3</div>
        <h3 style="color:#1a1615;margin-top:10px;margin-bottom:8px;">5-Signal Scoring</h3>
        <p style="font-size:13px;color:#757170;line-height:1.6;margin:0;font-family:Inter,sans-serif;">
            Generates sentence-transformer embeddings and applies weighted score fusion.
        </p>
    </div>
    """, unsafe_allow_html=True)
with p_col4:
    st.markdown("""
    <div class="phase-card">
        <div class="phase-number">4</div>
        <h3 style="color:#1a1615;margin-top:10px;margin-bottom:8px;">Agentic Re-Rank</h3>
        <p style="font-size:13px;color:#757170;line-height:1.6;margin:0;font-family:Inter,sans-serif;">
            Identifies top 100 fits using tie-breakers and availability scoring.
        </p>
    </div>
    """, unsafe_allow_html=True)

# Section 5 — Sponsors Strip
st.markdown("""
<div style="padding:24px 0;border-top:1px solid #e4e2e2;border-bottom:1px solid #e4e2e2;
            text-align:center;margin:40px 0;">
    <div class="section-label" style="margin-bottom:16px;">
        Hackathon Sponsor &amp; Technology Partners
    </div>
    <div style="display:flex;gap:32px;justify-content:center;align-items:center;flex-wrap:wrap;">
        <span style="font-size:16px;font-weight:800;color:#4A90FF;font-family:'Open Runde', sans-serif;">
            redrob<span style="color:#4A90FF">AI</span>
        </span>
        <span style="color:#e4e2e2;">·</span>
        <span style="font-size:15px;font-weight:600;color:#757170;">Google Gemini</span>
        <span style="color:#e4e2e2;">·</span>
        <span style="font-size:15px;font-weight:600;color:#757170;">Supabase</span>
        <span style="color:#e4e2e2;">·</span>
        <span style="font-size:15px;font-weight:600;color:#757170;">Hugging Face</span>
        <span style="color:#e4e2e2;">·</span>
        <span style="font-size:15px;font-weight:600;color:#757170;">Groq</span>
        <span style="color:#e4e2e2;">·</span>
        <span style="font-size:15px;font-weight:600;color:#757170;">FAISS</span>
    </div>
</div>
""", unsafe_allow_html=True)

# Section 7 — Footer
st.markdown("""
<div style="border-top:1px solid #e4e2e2;padding:32px 0;margin-top:64px;
            display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:16px;">
    <div style="font-size:13px;color:#757170;font-family:Inter,sans-serif;">
        © 2026 Calipr · Built at IITRAM Flux 2.0 · Sponsored by 
        <span style="color:#4A90FF;font-weight:700;">Redrob AI</span>
    </div>
    <div style="font-size:13px;color:#757170;font-family:Inter,sans-serif;">
        Made with ❤️ by Aum Santoki &amp; Team
    </div>
</div>
""", unsafe_allow_html=True)
